#!/usr/bin/env python3
"""
Convert a competitive analysis markdown to a formatted DOCX.

Usage:
    python scripts/md-to-docx-competitive.py
"""
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from scripts.utils.venv import ensure_venv  # noqa: E402

ensure_venv()
from scripts.utils.docx_helpers import set_cell_shading
from scripts.utils.workspace import get_outputs_dir

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

INPUT = str(get_outputs_dir() / "documents" / "competitive-analysis-example.md")
OUTPUT = str(get_outputs_dir() / "documents" / "competitive-analysis-example.docx")

# 31C brand colors
NAVY = RGBColor(0x0A, 0x1F, 0x3C)
ACCENT_BLUE = RGBColor(0x1A, 0x73, 0xE8)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x0D, 0x7C, 0x3E)
RED = RGBColor(0xC0, 0x39, 0x2B)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
TEAL = RGBColor(0x16, 0xA0, 0x85)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = parse_xml(f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{val}" w:space="0" w:color="CCCCCC"/>')
            borders.append(border)
    tcPr.append(borders)

def add_formatted_text(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    run.font.name = "Calibri"
    return run

def parse_inline(paragraph, text, default_color=None, default_size=None):
    """Parse bold and italic markdown inline formatting."""
    # Process **bold** and *italic*
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_formatted_text(paragraph, part[2:-2], bold=True, color=default_color, size=default_size)
        elif part.startswith('*') and part.endswith('*'):
            add_formatted_text(paragraph, part[1:-1], italic=True, color=default_color, size=default_size)
        else:
            # Handle inline code
            code_parts = re.split(r'(`.*?`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = add_formatted_text(paragraph, cp[1:-1], color=ACCENT_BLUE, size=default_size)
                    run.font.name = "Consolas"
                else:
                    add_formatted_text(paragraph, cp, color=default_color, size=default_size)

def build_docx():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Define styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GRAY
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for level, size, color in [(1, 22, NAVY), (2, 16, NAVY), (3, 13, ACCENT_BLUE)]:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.size = Pt(size)
        hs.font.color.rgb = color
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(18 if level == 1 else 14 if level == 2 else 10)
        hs.paragraph_format.space_after = Pt(6)

    # Read the markdown
    with open(INPUT, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # --- COVER PAGE ---
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "COMPETITIVE ANALYSIS", bold=True, color=NAVY, size=28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "Product A vs Competitor B", bold=True, color=ACCENT_BLUE, size=20)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "Internal Use Only - CEO & Engineering Leadership", italic=True, color=MED_GRAY, size=11)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "Example Company", bold=True, color=NAVY, size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "March 2026", color=MED_GRAY, size=11)

    doc.add_page_break()

    # --- PARSE MARKDOWN ---
    i = 0
    skip_title = True  # Skip the first H1 (already on cover page)
    in_table = False
    table_rows = []
    table_headers = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip the metadata block at top (>, ---)
        if i < 10 and (line.startswith('>') or line == '---' or line == ''):
            # But capture the metadata for a small info box
            if line.startswith('>') and not line.startswith('> **'):
                pass
            i += 1
            continue

        # Horizontal rules
        if line == '---':
            # Add a thin line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            # Add a border-bottom effect
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/></w:pBdr>')
            pPr.append(pBdr)
            i += 1
            continue

        # Empty lines
        if line == '':
            i += 1
            continue

        # Tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
                table_headers = []

            cells = [c.strip() for c in line.split('|')[1:-1]]

            # Skip separator row
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue

            if not table_headers:
                table_headers = cells
            else:
                table_rows.append(cells)

            # Check if next line continues the table
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            else:
                # Render the table
                in_table = False
                num_cols = len(table_headers)
                tbl = doc.add_table(rows=1 + len(table_rows), cols=num_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.style = 'Table Grid'

                # Header row
                for j, header in enumerate(table_headers):
                    cell = tbl.rows[0].cells[j]
                    set_cell_shading(cell, "0A1F3C")
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    parse_inline(p, header, default_color=WHITE, default_size=9)
                    for run in p.runs:
                        run.bold = True

                # Data rows
                for r_idx, row_data in enumerate(table_rows):
                    bg = "FFFFFF" if r_idx % 2 == 0 else "F5F7FA"
                    for j in range(min(len(row_data), num_cols)):
                        cell = tbl.rows[r_idx + 1].cells[j]
                        set_cell_shading(cell, bg)
                        p = cell.paragraphs[0]
                        text = row_data[j]

                        # Color-code verdict cells
                        if "Product A Advantage" in text:
                            parse_inline(p, text, default_color=GREEN, default_size=9)
                            for run in p.runs:
                                run.bold = True
                        elif "Competitor B Advantage" in text:
                            parse_inline(p, text, default_color=RED, default_size=9)
                            for run in p.runs:
                                run.bold = True
                        elif "Parity" in text:
                            parse_inline(p, text, default_color=TEAL, default_size=9)
                        elif "Emerging Gap" in text:
                            parse_inline(p, text, default_color=ORANGE, default_size=9)
                        else:
                            parse_inline(p, text, default_color=DARK_GRAY, default_size=9)

                # Adjust column widths for common table types
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)

                i += 1
                continue

        # Headings
        if line.startswith('#'):
            match = re.match(r'^(#{1,3})\s+(.*)', line)
            if match:
                level = len(match.group(1))
                text = match.group(2)

                if skip_title and level == 1:
                    skip_title = False
                    i += 1
                    continue

                heading = doc.add_heading(level=level)
                parse_inline(heading, text, default_color=NAVY if level <= 2 else ACCENT_BLUE)

                i += 1
                continue

        # Blockquotes (metadata/callouts)
        if line.startswith('>'):
            text = line.lstrip('> ').strip()
            if text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                pPr = p._p.get_or_add_pPr()
                pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="4" w:color="1A73E8"/></w:pBdr>')
                pPr.append(pBdr)
                parse_inline(p, text, default_color=MED_GRAY, default_size=9)
            i += 1
            continue

        # Bullet points
        if re.match(r'^(\s*)[-*]\s', line):
            match = re.match(r'^(\s*)[-*]\s(.*)', line)
            indent_level = len(match.group(1)) // 2
            text = match.group(2)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6 + indent_level * 0.6)
            p.paragraph_format.first_line_indent = Cm(-0.3)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            # Add bullet character
            add_formatted_text(p, "- " if indent_level == 0 else "  - ", color=ACCENT_BLUE, size=10)
            parse_inline(p, text, default_color=DARK_GRAY, default_size=10)

            i += 1
            continue

        # Numbered lists
        if re.match(r'^\d+\.\s', line):
            match = re.match(r'^(\d+)\.\s(.*)', line)
            num = match.group(1)
            text = match.group(2)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            add_formatted_text(p, f"{num}. ", bold=True, color=ACCENT_BLUE, size=10)
            parse_inline(p, text, default_color=DARK_GRAY, default_size=10)

            i += 1
            continue

        # Italic-only lines (recommendations)
        if line.strip().startswith('- *') and line.strip().endswith('*'):
            text = line.strip()[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.2)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            parse_inline(p, text, default_color=MED_GRAY, default_size=9)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline(p, line, default_color=DARK_GRAY, default_size=10)

        i += 1

    # --- FOOTER ---
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, "Example Company - Confidential", italic=True, color=MED_GRAY, size=8)

    doc.save(OUTPUT)
    print(f"DOCX saved to: {OUTPUT}")
    print(f"File size: {os.path.getsize(OUTPUT) / 1024:.0f} KB")

if __name__ == "__main__":
    build_docx()
