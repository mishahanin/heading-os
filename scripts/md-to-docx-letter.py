#!/usr/bin/env python3
"""Convert the Greece visa support letter from markdown to a properly formatted DOCX.

Usage:
    python scripts/md-to-docx-letter.py
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from scripts.utils.venv import ensure_venv  # noqa: E402

ensure_venv()
from scripts.utils.workspace import get_outputs_dir


def create_letter_docx(md_path, docx_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.strip().split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Sender header block (first few lines before the date line)
        if i == 0:
            # MISHA HANIN - sender name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = line.replace('**', '')
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # Address lines (lines 1-4 of content)
        if i in [1, 2, 3, 4]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = line.replace('**', '')
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # Date line
        if line.startswith('Date:'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # Recipient block
        if line.startswith('The Consul'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            i += 1
            # Continue with address lines
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('**Subject'):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(lines[i].strip())
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                i += 1
            continue

        # Subject line
        if '**Subject:' in line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = line.replace('**', '')
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # Section headers (bold lines like **About Myself**)
        if line.startswith('**') and line.endswith('**') and not line.startswith('**Subject'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text = line.replace('**', '')
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # Salutation
        if line.startswith('Dear '):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # Closing lines (Yours respectfully, signature block)
        if line.startswith('Yours '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            i += 1
            # Signature block
            while i < len(lines):
                line = lines[i].strip()
                if not line:
                    i += 1
                    continue
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                text = line.replace('**', '')
                is_name = ('Misha Hanin' in text and len(text) < 20)
                run = p.add_run(text)
                run.bold = is_name
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                i += 1
            continue

        # Numbered list items
        if re.match(r'^\d+\.', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_formatted_text(p, line)
            i += 1
            continue

        # Regular body paragraphs
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_formatted_text(p, line)
        i += 1

    doc.save(docx_path)
    print(f"DOCX created: {docx_path}")


def add_formatted_text(paragraph, text):
    """Parse markdown bold (**text**) and add runs with proper formatting."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'


if __name__ == '__main__':
    md_path = str(get_outputs_dir() / 'documents' / 'greece-visa-support-letter-zebo.md')
    docx_path = str(get_outputs_dir() / 'documents' / 'greece-visa-support-letter-zebo.docx')
    create_letter_docx(md_path, docx_path)
