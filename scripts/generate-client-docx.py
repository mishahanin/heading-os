#!/usr/bin/env python3
"""
Generate a professional DOCX version of the ODUN.ONE Conceptual Design Template.
Uses python-docx to create a formatted Word document with images, tables, and styling.
"""

import os
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from scripts.utils.venv import ensure_venv  # noqa: E402

ensure_venv()
from scripts.utils.docx_helpers import set_cell_shading
from scripts.utils.workspace import get_outputs_dir

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# Configuration / Paths
# ============================================================
IMAGES_DIR = str(get_outputs_dir() / 'images' / 'client')
OUTPUT_PATH = str(get_outputs_dir() / 'documents' / 'ODUN-ONE-Conceptual-Design-Template.docx')

# Brand colors
DARK_NAVY = RGBColor(0x0A, 0x1A, 0x2F)
BRAND_BLUE = RGBColor(0x00, 0x6B, 0xB6)
BRAND_ORANGE = RGBColor(0xE8, 0x6C, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = '006BB6'
TABLE_ALT_BG = 'F0F6FC'


# ============================================================
# Table & Style Helpers
# ============================================================
def style_header_row(row, bg_color=TABLE_HEADER_BG):
    """Style a table header row with background color and white text."""
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GRAY
            # Alternate row shading
            if row_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(width)

    doc.add_paragraph()  # spacing after table
    return table


# ============================================================
# Content Helpers (Headings / Body / Images / Lists)
# ============================================================
def add_image_safe(doc, image_name, width=Inches(6.5)):
    """Add an image if it exists, skip gracefully if not."""
    path = str(Path(IMAGES_DIR) / image_name)
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


def add_heading_styled(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = DARK_NAVY
            run.font.size = Pt(22)
        elif level == 2:
            run.font.color.rgb = BRAND_BLUE
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = DARK_NAVY
            run.font.size = Pt(13)
        elif level == 4:
            run.font.color.rgb = BRAND_BLUE
            run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return heading


def add_body_text(doc, text, bold=False, italic=False, color=None):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = color or DARK_GRAY
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_GRAY
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_GRAY
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_GRAY
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    return p


def add_placeholder(doc, text):
    """Add a placeholder note for partner customization."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = BRAND_ORANGE
    run.italic = True
    p.paragraph_format.space_after = Pt(6)
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_separator(doc):
    """Add a visual separator line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add a bottom border to simulate a line
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="006BB6"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


# ============================================================
# Rendering / Document Builder
# ============================================================
def build_document():
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK_GRAY

    # ══════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════

    # Spacing before title
    for _ in range(4):
        doc.add_paragraph()

    # Partner logo placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[PARTNER LOGO]')
    run.font.size = Pt(14)
    run.font.color.rgb = MEDIUM_GRAY
    run.italic = True

    doc.add_paragraph()

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Conceptual Design')
    run.font.size = Pt(36)
    run.font.color.rgb = DARK_NAVY
    run.font.name = 'Calibri'
    run.bold = True

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Deep Packet Intelligence Platform')
    run.font.size = Pt(20)
    run.font.color.rgb = BRAND_BLUE
    run.font.name = 'Calibri'

    doc.add_paragraph()

    # Cover image
    add_image_safe(doc, 'client-hero.png', width=Inches(5.5))

    doc.add_paragraph()

    # Document info table on cover
    cover_table = doc.add_table(rows=6, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_data = [
        ('Prepared for:', '[CUSTOMER NAME]'),
        ('Prepared by:', '[PARTNER NAME]'),
        ('Date:', '[DATE]'),
        ('Classification:', 'CONFIDENTIAL'),
        ('Document Reference:', '[REF-NUMBER]'),
        ('Version:', '1.0'),
    ]
    for i, (label, value) in enumerate(cover_data):
        cell_l = cover_table.rows[i].cells[0]
        cell_r = cover_table.rows[i].cells[1]
        cell_l.text = ''
        cell_r.text = ''
        run_l = cell_l.paragraphs[0].add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_l.font.name = 'Calibri'
        run_l.font.color.rgb = DARK_GRAY
        run_r = cell_r.paragraphs[0].add_run(value)
        run_r.font.size = Pt(10)
        run_r.font.name = 'Calibri'
        run_r.font.color.rgb = DARK_GRAY
        cell_l.width = Inches(2)
        cell_r.width = Inches(3.5)
        if i % 2 == 0:
            set_cell_shading(cell_l, TABLE_ALT_BG)
            set_cell_shading(cell_r, TABLE_ALT_BG)

    # Remove table borders for cleaner look
    for row in cover_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # DOCUMENT CONTROL
    # ══════════════════════════════════════════════

    add_heading_styled(doc, 'Document Control', level=1)
    add_separator(doc)

    add_heading_styled(doc, 'Revision History', level=2)
    add_styled_table(doc,
        ['Version', 'Date', 'Author', 'Description'],
        [
            ['1.0', '[DATE]', '[AUTHOR]', 'Initial conceptual design'],
            ['', '', '', ''],
        ],
        col_widths=[0.8, 1.2, 1.5, 3.0]
    )

    add_heading_styled(doc, 'Distribution List', level=2)
    add_styled_table(doc,
        ['Name', 'Organization', 'Role'],
        [
            ['[NAME]', '[CUSTOMER]', '[ROLE]'],
            ['[NAME]', '[PARTNER]', '[ROLE]'],
        ],
        col_widths=[2.0, 2.5, 2.0]
    )

    add_heading_styled(doc, 'Approval', level=2)
    add_styled_table(doc,
        ['Name', 'Role', 'Signature', 'Date'],
        [
            ['', '', '', ''],
            ['', '', '', ''],
        ],
        col_widths=[2.0, 2.0, 1.5, 1.0]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # TABLE OF CONTENTS placeholder
    # ══════════════════════════════════════════════

    add_heading_styled(doc, 'Table of Contents', level=1)
    add_separator(doc)

    toc_items = [
        '1.  Executive Summary',
        '2.  Customer Requirements & Objectives',
        '3.  Solution Overview',
        '4.  Architecture Design',
        '5.  Network Integration Architecture',
        '6.  Use Case Library',
        '7.  Deployment Architecture',
        '8.  High Availability & Resilience',
        '9.  Security & Data Sovereignty',
        '10. Implementation Approach',
        '11. Support & Operations Model',
        'Appendix A: Technical Specifications',
        'Appendix B: Glossary',
        'Appendix C: Compliance & Certification Matrix',
        'Appendix D: Competitive Advantage Summary',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_NAVY if not item.startswith('Appendix') else BRAND_BLUE
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3) if item.startswith('Appendix') else Inches(0)

    add_placeholder(doc, '[Note: Update page numbers after finalizing document in Word. Use Insert > Table of Contents for automatic generation.]')

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # SECTION 1: EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════

    add_heading_styled(doc, '1. Executive Summary', level=1)
    add_separator(doc)

    add_image_safe(doc, 'client-hero.png', width=Inches(6.0))

    add_heading_styled(doc, 'The Challenge', level=2)

    add_body_text(doc, 'Modern networks face four converging forces that legacy infrastructure was never designed to handle:')

    add_bullet(doc, 'Over 90% of internet traffic is now encrypted, turning networks into blind pipes. Traditional DPI tools built for cleartext cannot provide meaningful visibility without degrading performance or violating privacy.', bold_prefix='Encryption everywhere. ')
    add_bullet(doc, 'By 2030, the majority of global network traffic will be generated or shaped by AI. New patterns, new protocols, and unpredictable workloads are overwhelming inspection systems designed for the pre-AI internet.', bold_prefix='AI-generated traffic. ')
    add_bullet(doc, 'Subscriber bases, device counts, and application volumes are exploding. Every year brings higher bandwidth demands, denser 5G/6G rollouts, and more connected devices per subscriber.', bold_prefix='Continuous growth in scale. ')
    add_bullet(doc, 'New apps and services appear daily, many encrypted end-to-end. Without a platform that adapts in real time, operators fall behind -- blind to both opportunities (monetization) and risks (security, compliance).', bold_prefix='Relentless application introductions. ')

    add_body_text(doc, 'For telecom operators, this means losing revenue, control, and relevance. For governments, it means losing visibility, speed, and sovereignty.')

    add_heading_styled(doc, 'The Solution', level=2)

    add_body_text(doc, 'This document presents the conceptual design for a next-generation Deep Packet Intelligence platform that addresses these challenges through a fundamentally different approach. Rather than upgrading legacy DPI, the proposed solution represents a paradigm shift -- from packet inspection to packet intelligence.')

    add_body_text(doc, 'The platform follows the ODUN methodology -- a continuous intelligence cycle:', bold=True)

    add_bullet(doc, 'Capture traffic at carrier scale in real time, across fixed and mobile environments, without impacting performance', bold_prefix='Observe -- ')
    add_bullet(doc, 'Classify 3,500+ applications across 14 application categories and 11 content categories with >99% accuracy, including encrypted and AI-generated traffic', bold_prefix='Decode -- ')
    add_bullet(doc, 'Extract actionable intelligence through AI-driven analytics that correlate traffic patterns, subscriber context, and anomalies in under 100 milliseconds', bold_prefix='Understand -- ')
    add_bullet(doc, 'Enforce policy, optimize performance, enable monetization, and respond to threats in real time with sub-second policy activation', bold_prefix='Navigate -- ')

    add_heading_styled(doc, 'Key Outcomes', level=2)

    add_bullet(doc, 'across all traffic types, including encrypted and obfuscated flows, with >99% classification accuracy', bold_prefix='Complete network visibility ')
    add_bullet(doc, 'processing 5M+ flow records per second with application-level classification at carrier scale', bold_prefix='Real-time traffic intelligence ')
    add_bullet(doc, 'that reduces manual operational tasks by up to 60%, with natural-language policy creation replacing weeks of manual configuration', bold_prefix='AI-powered automation ')
    add_bullet(doc, '-- under 10 milliseconds from pattern detection to alert, with 99.7% detection accuracy', bold_prefix='Threat detection in real time ')
    add_bullet(doc, 'through fully on-premises, in-country deployment with no foreign cloud dependencies', bold_prefix='Data sovereignty ')
    add_bullet(doc, 'through advanced bypass detection (tethering, VPN abuse, SNI spoofing, OTT voice bypass), preventing an estimated $2.5M+ in annual revenue leakage', bold_prefix='Revenue protection ')
    add_bullet(doc, 'via unified dashboards, proactive monitoring, and natural-language AI copilot supporting 9+ languages', bold_prefix='Operational intelligence ')

    add_heading_styled(doc, 'Deployment Approach', level=2)

    add_body_text(doc, 'The platform deploys on standard x86 hardware using a container-native architecture (Kubernetes). Integration is non-disruptive -- starting in passive monitoring mode before transitioning to inline enforcement after validation. All data remains within the customer\'s infrastructure. No external cloud dependencies exist. High availability is built in through Active-Active clustering with automatic failover and session state synchronization.')

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # SECTION 2: CUSTOMER REQUIREMENTS
    # ══════════════════════════════════════════════

    add_heading_styled(doc, '2. Customer Requirements & Objectives', level=1)
    add_separator(doc)

    add_heading_styled(doc, '2.1 Customer Context', level=2)
    add_placeholder(doc, '[PARTNER: Describe the customer\'s organization, network environment, current challenges, and objectives in 2-3 paragraphs. Include subscriber base size, network topology overview, current DPI/traffic management solutions (if any), and regulatory environment.]')

    add_heading_styled(doc, '2.2 Carrier & ISP Requirements', level=2)
    add_placeholder(doc, 'Select applicable items:')

    add_styled_table(doc,
        ['#', 'Requirement', 'Priority', 'Platform Capability', 'Module'],
        [
            ['R1', 'Real-time traffic visibility across all subscriber traffic', '[H/M/L]', 'L2-L7 DPI with 3,500+ app classification across 14 categories', 'DataONE'],
            ['R2', 'Encrypted traffic classification without decryption', '[H/M/L]', 'Patent-pending encrypted traffic intelligence -- >99% accuracy', 'DataONE'],
            ['R3', 'Quality of Experience monitoring per application', '[H/M/L]', 'AppScore QoE metrics with drill-down analytics', 'OpsONE'],
            ['R4', 'Subscriber-aware traffic management', '[H/M/L]', 'RADIUS/DIAMETER/PCRF integration with per-subscriber policy', 'ControlONE'],
            ['R5', 'Data monetization through service differentiation', '[H/M/L]', 'Usage-based monetization, QoS packs, app-based tiers', 'ControlONE + DataONE'],
            ['R6', 'Network optimization and congestion management', '[H/M/L]', 'Traffic shaping, fair use, DSCP marking, header enrichment', 'DataONE + ControlONE'],
            ['R7', 'Real-time dashboards and reporting', '[H/M/L]', 'Unified dashboard, AI-powered reports in <60 seconds', 'OpsONE + AnalyticsONE'],
            ['R8', 'AI-driven analytics and predictive insights', '[H/M/L]', 'Behavioral modeling, anomaly detection, NLP copilot', 'AnalyticsONE'],
            ['R9', 'Revenue protection and bypass detection', '[H/M/L]', 'Tethering, VPN, SNI spoofing, OTT voice bypass detection', 'DataONE + ControlONE'],
        ],
        col_widths=[0.4, 2.0, 0.6, 2.2, 1.3]
    )

    add_heading_styled(doc, '2.3 Government / Regulatory Requirements', level=2)
    add_placeholder(doc, 'Select applicable items:')

    add_styled_table(doc,
        ['#', 'Requirement', 'Priority', 'Platform Capability', 'Module'],
        [
            ['R10', 'Regulatory content control and filtering', '[H/M/L]', 'Policy-based content filtering across 11 content categories', 'ControlONE + DataONE'],
            ['R11', 'Threat detection and security monitoring', '[H/M/L]', 'Real-time anomaly detection (<10ms), 99.7% accuracy', 'AnalyticsONE + DataONE'],
            ['R12', 'Lawful intercept readiness', '[H/M/L]', 'Selective traffic diversion, ETSI LI compliant, full audit trail', 'DataONE + ControlONE'],
            ['R13', 'Full data sovereignty -- in-country, on-premises', '[H/M/L]', 'Air-gapped deployment, no foreign cloud dependencies', 'Platform-wide'],
            ['R14', 'Audit logging and RBAC', '[H/M/L]', 'Complete traceability, granular RBAC, MFA, HSM', 'OpsONE'],
            ['R15', 'Subscriber traffic monitoring', '[H/M/L]', 'Session monitoring, behavioral analysis, pattern identification', 'DataONE + AnalyticsONE'],
            ['R16', 'Signal analysis and traffic diversion (MOD)', '[H/M/L]', 'Selective traffic capture and redirection', 'DataONE + ControlONE'],
        ],
        col_widths=[0.4, 2.0, 0.6, 2.2, 1.3]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # SECTION 3: SOLUTION OVERVIEW
    # ══════════════════════════════════════════════

    add_heading_styled(doc, '3. Solution Overview', level=1)
    add_separator(doc)

    add_heading_styled(doc, '3.1 Platform Philosophy', level=2)

    philosophies = [
        ('Sovereign by design. ', 'Deployed entirely on-premises, within the customer\'s data centers, under their full control. No data leaves the country. No foreign cloud dependencies. No external access. Sovereignty is not a configuration option -- it is the architecture.'),
        ('AI-native from day one. ', 'Artificial intelligence is not a bolt-on module added to a legacy codebase. The platform was designed from the ground up with AI at its core -- from encrypted traffic classification to behavioral modeling to the natural-language operator copilot supporting 9+ languages.'),
        ('Clean-slate architecture. ', 'Zero legacy code. Zero technical debt. Built for today\'s encrypted, AI-driven, cloud-native networks. Container-native microservices running on standard x86 hardware, with API-first design enabling unlimited extensibility.'),
        ('Non-aligned and vendor-neutral. ', 'The platform operates with any core network equipment from any vendor. No geopolitical risk to buyers. No lock-in. No dependency on any single hardware manufacturer or cloud provider.'),
    ]
    for bold_part, rest in philosophies:
        p = doc.add_paragraph()
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_NAVY
        run = p.add_run(rest)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_after = Pt(8)

    add_heading_styled(doc, '3.2 The ODUN Methodology', level=2)

    add_body_text(doc, 'The platform operates through a continuous intelligence cycle -- Observe, Decode, Understand, Navigate -- where each stage feeds the next, creating a self-improving loop of network intelligence.')

    add_image_safe(doc, 'client-odun-cycle.png', width=Inches(3.5))

    add_styled_table(doc,
        ['Stage', 'Action', 'Description'],
        [
            ['Observe', 'Capture', 'Continuously monitor network traffic in real time. Carrier-scale ingestion at up to 1.2 Tbps passive or 500 Gbps inline per server. 50M+ concurrent subscribers.'],
            ['Decode', 'Classify', 'Classify encrypted and AI-driven traffic by application. 3,500+ applications across 14 app categories and 11 content categories with >99% accuracy.'],
            ['Understand', 'Analyze', 'AI-driven analytics correlate traffic patterns, subscriber context, and anomalies in <100ms. Natural-language queries in 9+ languages.'],
            ['Navigate', 'Act', 'Enforce policy, optimize performance, enable monetization, and respond to threats. Sub-second policy activation. 10+ enforcement action types.'],
        ],
        col_widths=[1.0, 0.8, 4.7]
    )

    add_heading_styled(doc, '3.3 Platform Modules', level=2)

    modules = [
        ('DataONE -- High-Performance DPI Engine', 'The core inspection engine. Performs L2-L7 deep packet inspection, classifying 3,500+ applications with >99% accuracy. Processes up to 1.2 Tbps passive or 500 Gbps inline per server. Includes HTTP header enrichment, DSCP marking, L2/L3 redirection, and advanced bypass detection.'),
        ('ControlONE -- Policy & Subscriber Management', 'The intelligent policy orchestration layer. Integrates with RADIUS/DIAMETER, PCRF/OCS for policy and charging. Sub-second policy activation with AI-powered rollback for safety.'),
        ('OpsONE -- Unified Operational Dashboard', 'Real-time visibility with sub-second granularity. QoE metrics, multi-dimensional drill-down from node to subscriber. Interactive customizable dashboards with drag-and-drop widgets.'),
        ('AnalyticsONE -- AI Intelligence Engine', 'AI Copilot with natural-language interface in 9+ languages. Generates reports in <60 seconds. Real-time anomaly detection in <10ms with 99.7% accuracy. Predictive analytics and automated policy recommendations.'),
    ]
    for title, desc in modules:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = BRAND_BLUE
        doc.add_paragraph()  # newline
        add_body_text(doc, desc)

    add_heading_styled(doc, '3.4 Key Differentiators', level=2)

    add_styled_table(doc,
        ['Differentiator', 'Description'],
        [
            ['Sovereign deployment', 'Fully on-premises, in-country, air-gapped capable. No foreign cloud dependencies.'],
            ['AI-native architecture', 'NLP copilot in 9+ languages, <100ms AI response, <60s report generation.'],
            ['Clean-slate codebase', 'Zero legacy code, zero technical debt. Container-native with zero-downtime updates.'],
            ['Non-aligned vendor', 'Vendor-neutral, works with any core network. AWS/Azure/GCP/OpenStack compatible.'],
            ['Standard hardware', 'Commodity x86 servers. Up to 1.2 Tbps passive or 500 Gbps inline per server.'],
            ['Encrypted traffic intelligence', '>99% accuracy without decryption. Patent-pending four-layer algorithms.'],
            ['Advanced bypass detection', 'Tethering, VPN, SNI spoofing, OTT voice bypass. $2.5M+ annual revenue protection.'],
            ['API-first design', 'Open RESTful APIs with OpenAPI docs. Full DevOps automation capability.'],
        ],
        col_widths=[2.0, 4.5]
    )

    add_heading_styled(doc, '3.5 Platform vs. Legacy: Why It Matters', level=2)

    add_styled_table(doc,
        ['Capability', 'This Platform', 'Legacy DPI Solutions'],
        [
            ['Architecture', 'Cloud-native microservices, Kubernetes', 'Monolithic appliance-based'],
            ['AI integration', 'Core platform capability, system-level', 'Limited ML, bolt-on analytics'],
            ['Encrypted traffic', 'AI-powered, >99% accuracy', 'Signature-based, 90-95%'],
            ['Natural language ops', '9+ languages, conversational', 'Not available'],
            ['Policy creation', 'AI-generated, minutes', 'Manual, hours to days'],
            ['Report generation', 'AI-powered, <60 seconds', 'Pre-defined templates only'],
            ['Scalability', 'Elastic auto-scaling', 'Hardware-dependent'],
            ['Updates', 'Zero-downtime rolling', 'Maintenance windows required'],
            ['Anomaly detection', 'Real-time ML, <10ms', 'Rule-based'],
            ['TCO', '20-30% lower', 'Higher (proprietary HW + manual ops)'],
        ],
        col_widths=[1.8, 2.4, 2.3]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # SECTION 4: ARCHITECTURE DESIGN
    # ══════════════════════════════════════════════

    add_heading_styled(doc, '4. Architecture Design', level=1)
    add_separator(doc)

    add_heading_styled(doc, '4.1 High-Level Architecture', level=2)

    add_body_text(doc, 'The platform follows a three-plane architecture model that cleanly separates concerns while maintaining tight integration through an API-first design.')

    add_image_safe(doc, 'client-datacenter.png', width=Inches(6.0))

    planes = [
        ('Data Plane', 'Handles all packet-level operations. L2-L7 inspection, application classification, flow tracking, session analysis, HTTP header enrichment, DSCP marking, and traffic enforcement. Processes 5M+ flow records per second with 50M+ concurrent subscriber support.'),
        ('Control Plane', 'Makes policy decisions and orchestrates enforcement. Manages subscriber awareness, policy rules, charging integration, and enforcement actions with sub-second policy activation.'),
        ('Observability Plane', 'Provides visibility, intelligence, and operational control. Real-time dashboards, QoE monitoring, AI/ML anomaly detection in <10ms, and natural-language copilot in 9+ languages.'),
        ('API Layer', 'All inter-module communication flows through open RESTful APIs with OpenAPI documentation. Enables unlimited extensibility and DevOps automation.'),
    ]
    for title, desc in planes:
        p = doc.add_paragraph()
        run = p.add_run(title + ' -- ')
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_NAVY
        run = p.add_run(desc)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_after = Pt(6)

    add_heading_styled(doc, '4.2 DataONE -- DPI Engine Capabilities', level=2)

    add_styled_table(doc,
        ['Capability', 'Description'],
        [
            ['L2-L7 Deep Packet Inspection', 'Full protocol stack analysis including HTTP/2, HTTP/3, QUIC'],
            ['Application classification', '3,500+ apps across 14 app categories, 11 content categories, >99% accuracy'],
            ['Encrypted traffic intelligence', 'Patent-pending four-layer methodology without decryption'],
            ['HTTP header enrichment', 'Inject subscriber identity, content targeting, and analytics metadata'],
            ['DSCP marking', 'Set QoS markings for downstream network treatment'],
            ['Traffic shaping & blocking', '10+ enforcement action types per-application, per-subscriber'],
            ['L2/L3 traffic redirection', 'VLAN-based (L2) and routing-based (L3) traffic steering'],
            ['Bypass detection', 'Tethering, VPN/proxy, SNI spoofing, OTT voice bypass, protocol abuse'],
            ['Tunnel inspection', 'Multi-layer analysis: GRE, GTP, IPIP, VTI, nested tunnels'],
            ['Custom signatures', 'AI-assisted creation -- describe in natural language, system converts to rules'],
            ['Flow tracking', '5M+ flow records/sec with session correlation and metadata extraction'],
        ],
        col_widths=[2.2, 4.3]
    )

    add_heading_styled(doc, '4.3 AnalyticsONE -- AI Intelligence Engine', level=2)

    add_image_safe(doc, 'client-ai-engine.png', width=Inches(5.5))

    add_styled_table(doc,
        ['Capability', 'Description'],
        [
            ['AI Copilot', 'Natural-language interface in 9+ languages. Create policies, analyze traffic, generate reports. <100ms response.'],
            ['NLP policy creation', 'Describe intent in plain language -- AI generates, tests, and implements policies. Minutes vs. hours.'],
            ['Real-time anomaly detection', 'Threats, DDoS, fraud detected in <10ms with 99.7% accuracy'],
            ['Predictive analytics', 'Forecast traffic trends, capacity needs, QoE degradation, revenue optimization'],
            ['Self-optimization', 'AI learns optimal configurations and auto-adjusts policies'],
            ['AI-assisted troubleshooting', 'Describe a problem, AI identifies root cause with recommended actions'],
            ['Automated report generation', 'Custom reports from natural language queries in <60 seconds'],
        ],
        col_widths=[2.2, 4.3]
    )

    add_heading_styled(doc, '4.4 Encrypted Traffic Intelligence', level=2)

    add_body_text(doc, 'With over 90% of internet traffic now encrypted, the platform employs a four-layer classification methodology achieving >99% accuracy without decryption:')

    add_styled_table(doc,
        ['Layer', 'Method', 'Techniques'],
        [
            ['Layer 1', 'Pattern Matching', 'TLS/SNI extraction, certificate analysis, JA3/JA4 fingerprinting, IP reputation, DNS correlation'],
            ['Layer 2', 'Behavioral Analysis', 'Packet size distribution, inter-arrival timing, flow duration, burst patterns, session behavior'],
            ['Layer 3', 'Statistical Analysis', 'Entropy analysis, protocol conformance, statistical fingerprinting of traffic shapes'],
            ['Layer 4', 'Machine Learning / Deep Learning', 'Neural network models, continuous learning, patent-pending algorithms, evasion detection'],
        ],
        col_widths=[0.8, 1.5, 4.2]
    )

    add_body_text(doc, 'Result: >99% classification accuracy on encrypted traffic. No decryption required. No performance degradation. Models improve continuously as they process new traffic.')

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # SECTION 5: NETWORK INTEGRATION
    # ══════════════════════════════════════════════

    add_heading_styled(doc, '5. Network Integration Architecture', level=1)
    add_separator(doc)

    add_heading_styled(doc, '5.1 Network Insertion Points', level=2)

    add_body_text(doc, 'The platform is positioned at the boundary between the mobile core network and external data networks:')
    add_bullet(doc, 'Gi/SGi interface (between PGW and Internet/PDN)', bold_prefix='4G EPC: ')
    add_bullet(doc, 'N6 interface (between UPF and Data Network)', bold_prefix='5G SBA: ')

    add_heading_styled(doc, 'Deployment Integration Methods', level=3)

    add_styled_table(doc,
        ['Method', 'Description', 'Use Case'],
        [
            ['Layer 1 (Inline)', 'Traffic through DPI via Silicom Intelligent Bypass Switch. Hardware bypass failsafe.', 'Full enforcement: shaping, blocking, redirect, header enrichment'],
            ['Layer 2 (VLAN/MPLS)', 'L2 integration via VLAN tagging or MPLS labels', 'Complex network topologies'],
            ['Layer 3 (DPI on Stick)', 'Routing-based traffic steering to DPI platform', 'Minimal physical changes'],
            ['Passive (TAP/Mirror)', 'Traffic copy via Silicom Fiber TAP or port mirror', 'Monitoring-only, initial phase, forensic analysis'],
            ['Hybrid', 'Passive monitoring + selective inline enforcement', 'Phased deployment approach'],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    add_heading_styled(doc, '5.2 Subscriber Awareness', level=2)

    add_heading_styled(doc, '4G Methods', level=3)
    add_styled_table(doc,
        ['Method', 'Interface', 'Mechanism'],
        [
            ['RADIUS Accounting (Primary)', 'Gi/SGi', 'PGW sends RADIUS Accounting Start with IP-to-subscriber mapping'],
            ['GTP-C Correlation', 'S5/S8, S11', 'Sniff GTP-C Create Session Request/Response for IMSI/MSISDN/IP'],
            ['Gx (Diameter) Sniffing', 'Gx/S7', 'Extract subscriber identity from Credit-Control-Answer messages'],
        ],
        col_widths=[2.0, 1.2, 3.3]
    )

    add_heading_styled(doc, '5G Methods', level=3)
    add_styled_table(doc,
        ['Method', 'Interface', 'Mechanism'],
        [
            ['SMF Event Exposure (Primary)', 'Nsmf (HTTP/2)', 'Subscribe to SMF events. Push notification with SUPI + GPSI + UE IP.'],
            ['NEF Event Exposure', 'Nnef (HTTP/2)', 'NEF as API gateway for third-party integration (3GPP preferred)'],
            ['PFCP Extensions', 'N4', 'Vendor-specific PFCP IEs for IMSI/MSISDN push during session setup'],
        ],
        col_widths=[2.0, 1.2, 3.3]
    )

    add_heading_styled(doc, '5.3 External System Integrations', level=2)

    add_styled_table(doc,
        ['System Category', 'Examples', 'Integration Purpose'],
        [
            ['Telecom Billing', 'OCS, CHF, Charging Gateway', 'Real-time/offline charging, quota management'],
            ['Policy Control', 'PCRF, PCF, PCEF', 'Policy decisions, PCC rules, QoS enforcement'],
            ['CRM / BSS', 'Customer/order management', 'Subscriber profiles, service plans'],
            ['SIEM / SOC', 'Security operations, log mgmt', 'Security event forwarding, threat intel sharing'],
            ['OSS / NMS', 'Network/fault management', 'Performance data, alarm correlation'],
            ['Data Warehouse', 'BI, big data analytics', 'Historical data export (xDR, CDR)'],
            ['Lawful Intercept', 'Government/LEA systems', 'ETSI LI compliant traffic diversion, full audit'],
            ['Cloud Platforms', 'AWS, Azure, GCP, OpenStack', 'Native cloud API integration'],
        ],
        col_widths=[1.5, 2.0, 3.0]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # SECTION 6: USE CASE LIBRARY
    # ══════════════════════════════════════════════

    add_heading_styled(doc, '6. Use Case Library', level=1)
    add_separator(doc)

    add_body_text(doc, 'Each use case is pre-built into the platform and activated through configuration. These are capabilities, not separate products.')
    add_placeholder(doc, '[PARTNER: Select the use cases relevant to this deployment. Remove sections that do not apply.]')

    # 6.1
    add_heading_styled(doc, '6.1 Data Monetization & Traffic Management', level=2)

    use_cases_monetization = [
        ('Usage-Based Monetization & Fair Use', 'Enable differentiated service tiers and enforce fair-use policies, turning network capacity into revenue.', 'DataONE + ControlONE'),
        ('Application-Based Service Tiers', 'Offer app-specific packages (social bundles, video passes, gaming) that increase ARPU.', 'DataONE + ControlONE'),
        ('Time-Based Access Controls', 'Offer time-limited packages (1-hour video pass, weekend boost) for incremental revenue.', 'ControlONE + DataONE'),
        ('Network Optimization & QoS Packs', 'Protect quality during congestion, enable premium QoS offerings with DSCP marking.', 'DataONE + ControlONE'),
        ('Tethering & Bypass Control', 'Detect tethering, VPN bypass, SNI spoofing, OTT voice bypass. $2.5M+ annual protection.', 'DataONE + ControlONE'),
    ]

    add_styled_table(doc,
        ['Use Case', 'Business Value', 'Modules'],
        [(uc[0], uc[1], uc[2]) for uc in use_cases_monetization],
        col_widths=[2.0, 3.0, 1.5]
    )

    # 6.2
    add_heading_styled(doc, '6.2 Service Quality & QoE', level=2)

    add_styled_table(doc,
        ['Use Case', 'Business Value', 'Modules'],
        [
            ('Per-Application QoE (AppScore)', 'Proactive quality management before subscribers complain. Predictive QoE degradation alerts.', 'OpsONE + DataONE + AnalyticsONE'),
            ('Location-Based Quality Heatmaps', 'Identify geographic areas with poor quality for targeted investment.', 'OpsONE + DataONE'),
            ('Proactive Subscriber Monitoring', 'Detect quality degradation before subscribers notice. AI root-cause in seconds.', 'OpsONE + AnalyticsONE'),
        ],
        col_widths=[2.0, 3.0, 1.5]
    )

    # 6.3
    add_heading_styled(doc, '6.3 Subscriber Intelligence', level=2)

    add_styled_table(doc,
        ['Use Case', 'Business Value', 'Modules'],
        [
            ('Behavioral Analytics', 'Understand subscriber behavior for marketing, planning, and product development.', 'AnalyticsONE + DataONE'),
            ('Session Monitoring & Top Users', 'Identify heavy users, abnormal sessions, usage peaks in real time.', 'OpsONE + DataONE'),
            ('ML Revenue Forecasting', 'Predict revenue per service, capacity needs, optimal upsell timing.', 'AnalyticsONE'),
        ],
        col_widths=[2.0, 3.0, 1.5]
    )

    # 6.4
    add_heading_styled(doc, '6.4 Security & Threat Awareness', level=2)

    add_styled_table(doc,
        ['Use Case', 'Business Value', 'Modules'],
        [
            ('Real-Time Anomaly Detection', 'Identify threats in <10ms with 99.7% accuracy. DDoS, fraud, abuse.', 'AnalyticsONE + DataONE + OpsONE'),
            ('Tunneling & Evasion Detection', 'Detect VPN, DNS tunneling, protocol misuse, nested tunnels.', 'DataONE + AnalyticsONE'),
            ('Content Filtering', 'Enforce regulatory policies across 11 content categories. Full audit logging.', 'ControlONE + DataONE'),
            ('Automated Threat Response', 'Auto-respond to threats with AI safety nets and KPI-based auto-rollback.', 'AnalyticsONE + ControlONE'),
        ],
        col_widths=[2.0, 3.0, 1.5]
    )

    # 6.5
    add_heading_styled(doc, '6.5 Regulatory & Compliance', level=2)

    add_styled_table(doc,
        ['Use Case', 'Business Value', 'Modules'],
        [
            ('Lawful Intercept Readiness', 'Meet LI obligations with strict governance and ETSI compliance.', 'DataONE + ControlONE + OpsONE'),
            ('Data Sovereignty Enforcement', 'All data stays in-country. Air-gapped capable. HSM key management.', 'Platform-wide'),
        ],
        col_widths=[2.0, 3.0, 1.5]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # SECTION 7: DEPLOYMENT ARCHITECTURE
    # ══════════════════════════════════════════════

    add_heading_styled(doc, '7. Deployment Architecture', level=1)
    add_separator(doc)

    add_heading_styled(doc, '7.1 Reference Hardware Configuration', level=2)

    add_styled_table(doc,
        ['Component', 'Specification'],
        [
            ['Server platform', 'HP ProLiant DL385 Gen11 (or equivalent x86)'],
            ['Processor', '2x AMD EPYC 9845 (256 cores each, 512 cores total)'],
            ['Memory', '768 GB DDR5 ECC RAM'],
            ['Network interfaces', '12x 100GbE ports'],
            ['AI acceleration', 'NVIDIA L40 GPU (optional, for enhanced AI/ML workloads)'],
            ['Bypass switch', 'Silicom 40G/10G Intelligent Bypass Switch (inline)'],
            ['Fiber TAP', 'Silicom Fiber TAP (passive deployments)'],
            ['Operating system', 'Linux-based (Kubernetes-ready)'],
        ],
        col_widths=[2.0, 4.5]
    )

    add_placeholder(doc, '[PARTNER: Adjust hardware BOM based on specific capacity requirements and customer\'s preferred server vendor. The platform is hardware-agnostic.]')

    add_heading_styled(doc, '7.2 Deployment Topologies', level=2)

    add_body_text(doc, 'INL01: Inline Deployment -- Traffic flows through DPI via Silicom Intelligent Bypass Switch. Hardware bypass ensures traffic continuity if DPI server is unavailable. Full enforcement capability.', bold=True)

    add_body_text(doc, 'OFL01: Passive Deployment -- Traffic copied via Silicom Fiber TAP. Zero risk to production traffic. Ideal for initial deployment phase or monitoring-only requirements.', bold=True)

    add_heading_styled(doc, '7.3 Scalability', level=2)

    add_styled_table(doc,
        ['Specification', 'Value'],
        [
            ['Passive capacity (per server)', 'Up to 1.2 Tbps'],
            ['Inline capacity (per server)', 'Up to 500 Gbps'],
            ['Flow records per second', '5M+'],
            ['Concurrent subscribers', '50M+'],
            ['Scaling method', 'Horizontal -- add servers for linear capacity growth'],
            ['Auto-scaling', 'Kubernetes elastic scaling based on traffic load'],
            ['Management', 'Single management plane regardless of cluster size'],
        ],
        col_widths=[2.5, 4.0]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # SECTION 8: HIGH AVAILABILITY
    # ══════════════════════════════════════════════

    add_heading_styled(doc, '8. High Availability & Resilience', level=1)
    add_separator(doc)

    add_body_text(doc, 'The platform is designed for carrier-grade availability with 99.99% uptime target.')

    add_styled_table(doc,
        ['Capability', 'Implementation'],
        [
            ['Active-Active clustering', 'Multiple instances share traffic load. No single point of failure.'],
            ['Automatic failover', 'Traffic redistributed to remaining nodes within seconds. No manual intervention.'],
            ['Session state sync', 'Active sessions synchronized across nodes. Failover transparent to subscribers.'],
            ['Hardware bypass', 'Silicom Bypass Switches pass traffic through if DPI unavailable.'],
            ['Kubernetes self-healing', 'Failed containers/pods automatically restarted. Unhealthy nodes replaced.'],
            ['Geographic redundancy', 'Instances across multiple data centers for site-level resilience.'],
            ['Disaster recovery', 'Automated DR procedures with defined RPO/RTO targets.'],
            ['Zero-downtime updates', 'Rolling updates via Kubernetes. Modules updated independently.'],
        ],
        col_widths=[2.0, 4.5]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # SECTION 9: SECURITY & DATA SOVEREIGNTY
    # ══════════════════════════════════════════════

    add_heading_styled(doc, '9. Security & Data Sovereignty', level=1)
    add_separator(doc)

    add_image_safe(doc, 'client-sovereignty.png', width=Inches(5.5))

    add_body_text(doc, 'Security and data sovereignty are foundational architectural decisions, not optional features.')

    add_heading_styled(doc, '9.1 Data Sovereignty', level=2)

    add_styled_table(doc,
        ['Principle', 'Implementation'],
        [
            ['In-country deployment', 'All components within customer\'s data center, within country borders'],
            ['No foreign cloud dependencies', 'Zero external cloud reliance. All processing local.'],
            ['Air-gapped capability', 'Operates with no internet connectivity required'],
            ['No external telemetry', 'No data sent to external servers, vendor cloud, or third parties'],
        ],
        col_widths=[2.5, 4.0]
    )

    add_heading_styled(doc, '9.2 Access Control & Cryptographic Security', level=2)

    add_styled_table(doc,
        ['Control', 'Description'],
        [
            ['Role-Based Access Control', 'Granular RBAC with fine-grained permissions'],
            ['Multi-Factor Authentication', 'Enforced MFA supporting hardware tokens, software auth, biometrics'],
            ['SSO / LDAP / OAuth2', 'Enterprise identity provider integration'],
            ['Hardware Security Module', 'HSM integration for key management and crypto operations'],
            ['Secure boot', 'Firmware and software integrity validation at boot'],
            ['Encryption at rest', 'AES-256 encryption for all stored data'],
            ['Encryption in transit', 'TLS 1.3 for all communications'],
        ],
        col_widths=[2.5, 4.0]
    )

    add_heading_styled(doc, '9.3 Compliance Standards', level=2)

    add_styled_table(doc,
        ['Standard', 'Scope'],
        [
            ['ISO 27001', 'Information security management system'],
            ['SOC 2 Type II', 'Security, availability, and confidentiality controls'],
            ['Common Criteria EAL4+', 'Security evaluation for IT products'],
            ['FIPS 140-2 Level 3', 'Cryptographic module validation'],
            ['ETSI LI', 'Lawful interception standards'],
            ['3GPP Standards', '4G/5G integration specifications'],
            ['GDPR', 'European data protection regulation'],
        ],
        col_widths=[2.5, 4.0]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # SECTION 10: IMPLEMENTATION APPROACH
    # ══════════════════════════════════════════════

    add_heading_styled(doc, '10. Implementation Approach', level=1)
    add_separator(doc)

    add_body_text(doc, 'Total time to production: 14-20 weeks. Core principle: Passive first, inline after validation.', bold=True)

    add_styled_table(doc,
        ['Phase', 'Activities', 'Duration'],
        [
            ['Phase 1: Pilot Deployment', 'Requirements validation, network assessment, hardware install, platform deploy, passive integration, initial config', '4-6 weeks'],
            ['Phase 2: Parallel Run', 'Classification validation (>99%), system integration, integration testing, performance validation, policy development, operator training', '6-8 weeks'],
            ['Phase 3: Gradual Cutover', 'Inline transition with Intelligent Bypass, staged enforcement, performance tuning, security validation, UAT', '4-6 weeks'],
            ['Phase 4: Full Production', 'All policies active, knowledge transfer, documentation, continuous AI optimization, stabilization monitoring', 'Ongoing'],
        ],
        col_widths=[1.8, 3.5, 1.2]
    )

    add_heading_styled(doc, 'Acceptance Criteria', level=2)
    add_placeholder(doc, '[PARTNER: Define specific acceptance criteria with the customer.]')

    add_styled_table(doc,
        ['KPI', 'Target', 'Measurement'],
        [
            ['Classification accuracy', '>= 99%', 'Validation against known traffic'],
            ['Platform throughput', '[X Gbps] as per design', 'Traffic generator or production'],
            ['AI response time', '< 100 ms', 'Query to response measurement'],
            ['Threat detection time', '< 10 ms', 'Pattern to alert measurement'],
            ['Platform availability', '>= 99.99%', 'Uptime monitoring'],
        ],
        col_widths=[2.0, 2.0, 2.5]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # SECTION 11: SUPPORT & OPERATIONS
    # ══════════════════════════════════════════════

    add_heading_styled(doc, '11. Support & Operations Model', level=1)
    add_separator(doc)

    add_heading_styled(doc, 'Support Framework', level=2)

    add_styled_table(doc,
        ['Element', 'Description'],
        [
            ['Support tiers', 'L1 (monitoring & triage), L2 (technical), L3 (engineering escalation)'],
            ['Critical response', '<4 hour response for critical issues'],
            ['AI-assisted support', 'Built-in AI troubleshooting reduces dependence on deep expertise'],
            ['Signature updates', 'Hot-swap, no restart, zero traffic impact'],
            ['Platform updates', 'Zero-downtime rolling updates via Kubernetes'],
        ],
        col_widths=[2.0, 4.5]
    )

    add_heading_styled(doc, 'Product Lifecycle', level=2)

    add_styled_table(doc,
        ['Period', 'Coverage'],
        [
            ['Year 1', 'Full lifecycle with updates and support. First year included with license.'],
            ['Years 2-5', 'Annual support at 20% of license fee. Software updates, signatures, support.'],
            ['Years 6-10 (optional)', 'Extended lifecycle at 30% of license fee. Continued evolution.'],
        ],
        col_widths=[2.0, 4.5]
    )

    add_heading_styled(doc, 'Business Continuity', level=2)

    add_styled_table(doc,
        ['Protection', 'Description'],
        [
            ['Source code escrow', 'Available for enterprise customers via independent escrow agent'],
            ['Open standards', 'REST APIs, Kubernetes, standard protocols -- no proprietary lock-in'],
            ['Modular architecture', 'Individual components replaceable independently'],
        ],
        col_widths=[2.0, 4.5]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # APPENDIX A: TECHNICAL SPECIFICATIONS
    # ══════════════════════════════════════════════

    add_heading_styled(doc, 'Appendix A: Technical Specifications', level=1)
    add_separator(doc)

    add_heading_styled(doc, 'Performance Specifications', level=2)

    add_styled_table(doc,
        ['Specification', 'Value'],
        [
            ['Passive capacity (per server)', 'Up to 1.2 Tbps'],
            ['Inline capacity (per server)', 'Up to 500 Gbps'],
            ['Application library', '3,500+ applications'],
            ['App/content categories', '14 application categories, 11 content categories'],
            ['Classification accuracy', '>99% (including encrypted traffic)'],
            ['Concurrent subscribers', '50M+'],
            ['Flow records/second', '5M+'],
            ['AI response time', '<100 ms'],
            ['Threat detection time', '<10 ms'],
            ['Report generation', '<60 seconds (AI-powered)'],
            ['Policy activation', 'Sub-second'],
            ['Availability target', '99.99%'],
        ],
        col_widths=[2.5, 4.0]
    )

    add_heading_styled(doc, 'Enforcement Actions', level=2)

    add_styled_table(doc,
        ['Action', 'Description'],
        [
            ['Traffic shaping', 'Per-application and per-subscriber bandwidth control'],
            ['Blocking', 'Application or content blocking based on policy'],
            ['Redirect', 'HTTP redirect to portals, captive portals, notification pages'],
            ['Mirror', 'Copy traffic to secondary destinations for analysis'],
            ['HTTP header enrichment', 'Inject subscriber identity and analytics data'],
            ['DSCP marking', 'Set QoS markings for downstream network treatment'],
            ['L2 redirection', 'VLAN-based traffic steering'],
            ['L3 redirection', 'Routing-based traffic steering'],
            ['Throttle', 'Reduce bandwidth for specific apps/subscribers'],
            ['Prioritize', 'Increase QoS for specific apps/subscribers'],
        ],
        col_widths=[2.0, 4.5]
    )

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # APPENDIX D: COMPETITIVE ADVANTAGE
    # ══════════════════════════════════════════════

    add_heading_styled(doc, 'Appendix D: Competitive Advantage Summary', level=1)
    add_separator(doc)

    add_heading_styled(doc, 'Architecture & Deployment', level=2)

    add_styled_table(doc,
        ['Capability', 'This Platform', 'Legacy Solutions'],
        [
            ['Architecture', 'Modular microservices, Docker containers', 'Monolithic appliance-based'],
            ['Orchestration', 'Native Kubernetes, auto-scaling, self-healing', 'Limited orchestration, manual scaling'],
            ['Deployment', 'Bare metal, VNF, CNF -- all supported', 'Primarily bare metal and VNF'],
            ['Scalability', 'Elastic auto-scaling to Tbps', 'Hardware-dependent, complex planning'],
            ['Updates', 'Zero-downtime rolling, modules independent', 'Maintenance windows, complex upgrades'],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    add_heading_styled(doc, 'AI & Automation', level=2)

    add_styled_table(doc,
        ['Capability', 'This Platform', 'Legacy Solutions'],
        [
            ['AI integration', 'Core platform capability', 'Limited ML, bolt-on analytics'],
            ['Natural language', '9+ languages, conversational', 'Not available'],
            ['Policy creation', 'AI-generated, minutes', 'Manual programming, hours/days'],
            ['Anomaly detection', 'Real-time ML, <10ms', 'Rule-based detection'],
            ['Self-optimization', 'AI learns and auto-adjusts', 'Manual optimization'],
            ['Reporting', 'AI-powered, <60 seconds', 'Pre-defined templates only'],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    add_heading_styled(doc, 'Operations & Management', level=2)

    add_styled_table(doc,
        ['Capability', 'This Platform', 'Legacy Solutions'],
        [
            ['Availability', '99.99%, K8s self-healing, geo-redundancy', '99.9% with proper setup'],
            ['Troubleshooting', 'AI-assisted root cause analysis', 'Manual investigation'],
            ['Backup', 'Automated scheduled backups', 'Manual procedures'],
            ['E2E tracing', 'Full request tracing through system', 'Limited tracing'],
            ['TCO', '20-30% lower through automation', 'Higher (proprietary HW + manual ops)'],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    doc.add_paragraph()
    add_separator(doc)

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('This document is confidential and intended solely for the use of the parties to whom it is addressed.')
    run.font.size = Pt(8)
    run.font.color.rgb = MEDIUM_GRAY
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[PARTNER NAME] | [DATE]')
    run.font.size = Pt(9)
    run.font.color.rgb = MEDIUM_GRAY

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f'[OK] Document saved: {OUTPUT_PATH}')
    print(f'[INFO] File size: {os.path.getsize(OUTPUT_PATH) / 1024:.0f} KB')


# ============================================================
# CLI / Main
# ============================================================
if __name__ == '__main__':
    build_document()
