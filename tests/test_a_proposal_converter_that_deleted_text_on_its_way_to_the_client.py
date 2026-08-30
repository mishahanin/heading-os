"""The markdown-to-DOCX proposal converter used to drop text on the floor.

Five measured defects in `scripts/md-to-docx-proposal.py`, all of them silent
and all of them in a document that goes to a customer:

1. Table detection was a substring test (`'---' in lines[i + 1]`), so a body
   line carrying pipes became a mangled one-cell table, or vanished outright.
2. `#` and `##` both rendered as Heading 1, collapsing two outline levels.
3. Headings stripped links but not emphasis, so `## **Scope**` kept its
   asterisks.
4. An escaped pipe (`\\|`) inside a cell split the row one field too wide, and
   the guard that noticed silently discarded the overflow.
5. `***text***` was parsed as `***text**` plus a stray `*`, which bolded a
   leading asterisk and deleted the closing pair.

Every assertion here reads the PRODUCED DOCUMENT -- paragraph text, heading
style, table cell contents, run formatting -- never the converter's source.
A test that greps the script would pass over a rewrite that reintroduced the
bug.
"""
import importlib.util
import io
import contextlib
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

SCRIPT = ROOT / "scripts" / "md-to-docx-proposal.py"

# The last line the cover page emits. Body content starts after it.
COVER_LAST = 'Classification: Confidential -- Partner Use Only'
# The closing line `build_document` appends after all body content.
CLOSING_LINE = '31 Concept | [HQ City 1] | [HQ City 2] | 31c.io'

pytest.importorskip("docx", reason="python-docx is the optional 'documents' extra")


def _load():
    """Fresh module per render: the converter keeps its paths in globals."""
    spec = importlib.util.spec_from_file_location("md_to_docx_proposal", SCRIPT)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def render(markdown, tmp_path, name="proposal"):
    """Convert `markdown` and return the parsed python-docx Document.

    Writes strictly under `tmp_path`. The converter takes its input and output
    from module-level globals, so this rebinds them rather than reaching for a
    hardcoded path in the operator's data overlay.
    """
    from docx import Document

    src = tmp_path / f"{name}.md"
    src.write_text(markdown, encoding="utf-8")
    dst = tmp_path / f"{name}.docx"

    mod = _load()
    mod.INPUT_PATH = str(src)
    mod.OUTPUT_PATH = str(dst)
    with contextlib.redirect_stdout(io.StringIO()):
        mod.build_document()

    assert dst.exists(), "converter produced no document"
    return Document(str(dst))


def body(doc):
    """Non-empty body paragraphs: cover page and closing line excluded."""
    paras = [p for p in doc.paragraphs if p.text.strip()]
    start = next(i for i, p in enumerate(paras) if p.text == COVER_LAST) + 1
    end = next(i for i, p in enumerate(paras) if p.text == CLOSING_LINE)
    return paras[start:end]


def body_text(doc):
    return [p.text for p in body(doc)]


def styled_body(doc):
    return [(p.style.name, p.text) for p in body(doc)]


def grid(doc):
    """Every table as a list of rows of cell strings."""
    return [[[c.text for c in row.cells] for row in t.rows] for t in doc.tables]


def runs_of(paragraph):
    return [(r.text, bool(r.bold), bool(r.italic)) for r in paragraph.runs]


# --------------------------------------------------------------------------
# Defect 2 -- heading levels collapsed
# --------------------------------------------------------------------------

def test_each_markdown_heading_depth_gets_its_own_word_outline_level(tmp_path):
    """`#` and `##` must not both render as Heading 1.

    Measured before the fix: h1 and h2 both produced 'Heading 1', h3 produced
    'Heading 2' and h4 'Heading 3'. Two distinct structural levels rendered
    identically, so Word's navigation pane and any generated TOC could not
    tell a section from the document title.
    """
    doc = render(
        "# Programme overview\n\n"
        "## Delivery scope\n\n"
        "### Workstream one\n\n"
        "#### Acceptance criteria\n",
        tmp_path,
    )

    assert styled_body(doc) == [
        ('Heading 1', 'Programme overview'),
        ('Heading 2', 'Delivery scope'),
        ('Heading 3', 'Workstream one'),
        ('Heading 4', 'Acceptance criteria'),
    ]


def test_the_four_heading_levels_are_visually_distinct(tmp_path):
    """A distinct level is worth nothing if two levels look identical.

    Guards the remap against being 'fixed' by pointing several depths at one
    branded style: every level the converter emits carries its own size.
    """
    doc = render(
        "# One\n\n## Two\n\n### Three\n\n#### Four\n",
        tmp_path,
    )
    sizes = [p.style.font.size for p in body(doc)]
    assert all(s is not None for s in sizes), f"an unstyled heading level: {sizes}"
    assert len(set(sizes)) == 4, f"heading sizes are not distinct: {sizes}"
    # Deeper headings are never larger than shallower ones.
    assert sizes == sorted(sizes, reverse=True), sizes


# --------------------------------------------------------------------------
# Defect 1 -- table detection was a substring test, and ate prose
# --------------------------------------------------------------------------

@pytest.mark.parametrize("second_line", [
    "Costs land between 40 --- 60 percent of the baseline.",  # em dash as ---
    "The migration path is rehost --- replatform --- refactor.",
    "--- and the rule above closes the section.",
])
def test_prose_is_not_swallowed_by_a_dash_on_the_following_line(tmp_path, second_line):
    """A body line with pipes must survive a `---` on the next line.

    Measured before the fix, for the first case: the two lines produced a
    one-cell table holding only 'b' -- 'Option a' and 'Option c' were gone
    from the document, and no warning was printed.
    """
    doc = render(f"Option a | Option b | Option c\n{second_line}\n", tmp_path)

    assert doc.tables == [], f"prose became a table: {grid(doc)}"
    assert body_text(doc) == ["Option a | Option b | Option c " + second_line]


def test_a_single_pipe_line_does_not_vanish(tmp_path):
    """The worst shape: the line left no cells, so no block was emitted at all.

    `split('|')[1:-1]` on `Scope | fixed for the term` yields an empty list,
    the row was skipped, `rows` stayed empty, and the branch `continue`d
    without appending anything. The sentence was deleted outright.
    """
    doc = render(
        "Scope | fixed for the term\nrehost --- replatform\n",
        tmp_path,
    )

    assert doc.tables == [], f"prose became a table: {grid(doc)}"
    assert "Scope | fixed for the term" in " ".join(body_text(doc))


def test_a_horizontal_rule_under_a_pipe_line_is_not_a_table(tmp_path):
    """A lone `---` is a rule, not a one-column separator for a 3-cell line.

    GFM requires the separator to match the header's column count. Without
    that check, shape-testing alone still reads this as a table because a
    horizontal rule and a single-column separator are spelled identically.
    """
    doc = render("Phase one | Phase two | Phase three\n---\nNext.\n", tmp_path)

    assert doc.tables == [], f"a rule made a table: {grid(doc)}"
    assert "Phase one | Phase two | Phase three" in " ".join(body_text(doc))


# --------------------------------------------------------------------------
# Defect 4 -- an escaped pipe split the row too wide, and the overflow was cut
# --------------------------------------------------------------------------

def test_an_escaped_pipe_stays_inside_its_cell(tmp_path):
    """`\\|` is the only way GFM has to put a pipe in a cell.

    Measured before the fix: the row split into four fields against a
    three-column header, so every value shifted one column left, the stray
    backslash survived into the first cell, and '200' -- a number in a priced
    row -- was discarded by the `c_idx < num_cols` guard.
    """
    doc = render(
        "| Line item | Year 1 | Year 2 |\n"
        "| --- | --- | --- |\n"
        "| cost model \\| fixed | 100 | 200 |\n",
        tmp_path,
    )

    assert grid(doc) == [[
        ['Line item', 'Year 1', 'Year 2'],
        ['cost model | fixed', '100', '200'],
    ]]


def test_a_row_wider_than_its_header_is_widened_not_truncated(tmp_path):
    """A genuinely malformed table must not lose its overflow silently.

    The decision recorded in the converter is to widen: an empty header cell
    is a defect the operator can see, a deleted value is not.
    """
    doc = render(
        "| Item | Value |\n"
        "| --- | --- |\n"
        "| Licence | 40 | stray fourth |\n",
        tmp_path,
    )

    assert grid(doc) == [[
        ['Item', 'Value', ''],
        ['Licence', '40', 'stray fourth'],
    ]]


def test_a_borderless_table_keeps_its_edge_columns(tmp_path):
    """`[1:-1]` deleted the first and last column of a table written without
    outer pipes, which GFM allows."""
    doc = render(
        "Phase | Weeks | Owner\n"
        "--- | --- | ---\n"
        "Design | 6 | Delivery lead\n",
        tmp_path,
    )

    assert grid(doc) == [[
        ['Phase', 'Weeks', 'Owner'],
        ['Design', '6', 'Delivery lead'],
    ]]


# --------------------------------------------------------------------------
# Defect 3 -- headings stripped links but not emphasis
# --------------------------------------------------------------------------

def test_a_heading_renders_its_emphasis_instead_of_showing_the_markers(tmp_path):
    """`## **Scope**` must not reach the client with its asterisks showing.

    Measured before the fix: the four heading branches ran the link regex
    alone, so the heading text was literally '**Scope**' and '*Timing*'. The
    table-cell path stripped `**bold**` and the paragraph path rendered it as
    a run; only headings printed the raw markers.
    """
    doc = render(
        "## **Scope**\n\n### *Timing*\n\n#### [Annex A](https://example.invalid/a)\n",
        tmp_path,
    )

    assert styled_body(doc) == [
        ('Heading 2', 'Scope'),
        ('Heading 3', 'Timing'),
        ('Heading 4', 'Annex A'),
    ]
    scope, timing, _annex = body(doc)
    assert runs_of(scope) == [('Scope', True, False)]
    assert runs_of(timing) == [('Timing', False, True)]


# --------------------------------------------------------------------------
# Defect 5 -- `***text***` was misparsed
# --------------------------------------------------------------------------

def test_triple_asterisk_is_bold_italic_and_loses_no_characters(tmp_path):
    """`re.split(r'(\\*\\*.*?\\*\\*)', ...)` matched one asterisk too few.

    Measured before the fix, mid-sentence: 'We treat this as ***critical***
    today.' rendered as 'We treat this as *critical* today.' -- the run
    ('*critical', bold) carried a LEADING asterisk and the closing pair
    migrated into the following text.
    """
    doc = render("We treat this as ***critical*** today.\n", tmp_path)

    para, = body(doc)
    assert para.text == "We treat this as critical today."
    assert runs_of(para) == [
        ('We treat this as ', False, False),
        ('critical', True, True),
        (' today.', False, False),
    ]


def test_triple_asterisk_at_the_end_of_a_line_is_not_deleted(tmp_path):
    """The same defect DELETES characters when nothing follows the emphasis.

    Measured before the fix: the whole line '***Confidential***' rendered as
    '*Confidential' -- the closing `**` became an empty italic run, so two
    characters left the document and one wrong one stayed. This is the shape
    the audit predicted; the mid-sentence case above is not, which is why both
    are pinned.
    """
    doc = render("***Confidential***\n", tmp_path)

    para, = body(doc)
    assert para.text == "Confidential"
    assert runs_of(para) == [('Confidential', True, True)]


# --------------------------------------------------------------------------
# Negative control -- the happy path must survive all five fixes
# --------------------------------------------------------------------------

ORDINARY_PROPOSAL = """## Engagement summary

We propose a phased rollout across **three regions**, starting with the
northern cluster in *week two*.

### Commercial terms

| Line item | Year 1 | Year 2 |
| --- | ---: | :--- |
| Platform licence | 128,400 | 131,900 |
| Support retainer | 46,250 | 46,250 |

#### Assumptions

- Delivery begins within 14 days of signature
- Acceptance testing runs for **21 days**
  - Two rounds of defect triage are included
- Pricing holds for 90 days

The figures above exclude travel, billed at cost.
"""


def test_an_ordinary_proposal_renders_unchanged(tmp_path):
    """Four of the five fixes touch the parser's happy path.

    A fix that mangles the normal case is worse than the defect it cures, so
    this pins the whole ordinary document: heading styles, a real table with
    an alignment row, nested bullets, and real bold and italic runs.
    """
    doc = render(ORDINARY_PROPOSAL, tmp_path)

    assert styled_body(doc) == [
        ('Heading 2', 'Engagement summary'),
        ('Normal', 'We propose a phased rollout across three regions, starting '
                   'with the northern cluster in week two.'),
        ('Heading 3', 'Commercial terms'),
        ('Heading 4', 'Assumptions'),
        ('List Bullet', 'Delivery begins within 14 days of signature'),
        ('List Bullet', 'Acceptance testing runs for 21 days'),
        ('List Bullet 2', 'Two rounds of defect triage are included'),
        ('List Bullet', 'Pricing holds for 90 days'),
        ('Normal', 'The figures above exclude travel, billed at cost.'),
    ]

    # The table keeps every column, every row and every value. The `---:` and
    # `:---` alignment cells are separators and must not become a data row.
    assert grid(doc) == [[
        ['Line item', 'Year 1', 'Year 2'],
        ['Platform licence', '128,400', '131,900'],
        ['Support retainer', '46,250', '46,250'],
    ]]

    intro = body(doc)[1]
    assert runs_of(intro) == [
        ('We propose a phased rollout across ', False, False),
        ('three regions', True, False),
        (', starting with the northern cluster in ', False, False),
        ('week two', False, True),
        ('.', False, False),
    ]

    # Emphasis inside a bullet still renders as a run, not as literal markers.
    acceptance = body(doc)[5]
    assert runs_of(acceptance) == [
        ('Acceptance testing runs for ', False, False),
        ('21 days', True, False),
    ]


def test_the_ordinary_proposal_loses_no_characters(tmp_path):
    """A blunt whole-document check, independent of the structure assertions.

    Every word of the source that is not a markdown marker must appear
    somewhere in the rendered document. This is the property the five defects
    all violated, stated once over a realistic input.
    """
    import re as _re

    doc = render(ORDINARY_PROPOSAL, tmp_path)
    rendered = " ".join(body_text(doc) + [c for t in grid(doc) for r in t for c in r])

    source_words = _re.findall(
        r"[A-Za-z0-9,.]+",
        _re.sub(r"[*|#>-]", " ", ORDINARY_PROPOSAL),
    )
    missing = [w for w in source_words if w.strip(",.") and w.strip(",.") not in rendered]
    assert missing == [], f"the converter dropped: {missing}"
