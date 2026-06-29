#!/usr/bin/env python3
"""Convert a charter confirmation letter to a professionally formatted DOCX.

Usage:
    python scripts/md-to-docx-charter.py
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from scripts.utils.venv import ensure_venv  # noqa: E402

ensure_venv()
from scripts.utils.workspace import get_outputs_dir


def add_horizontal_line(doc):
    """Add a horizontal line paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '333333')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_formatted_text(paragraph, text, default_size=Pt(11)):
    """Parse markdown bold and add runs."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.size = default_size
        run.font.name = 'Arial'


def create_charter_docx(md_path, docx_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15

    # === LETTERHEAD ===
    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('EXAMPLE CHARTERS LTD')
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('EXAMPLE CHARTERS S.A.')
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # Address line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('123 Example Street, Example District, 100 00, Athens, Greece')
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Contact line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Tel: +30 000 000 0000  |  Email: info@example-charters.com  |  www.example-charters.com')
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # VAT line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('VAT No.: EL000000000')
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_horizontal_line(doc)

    # === REFERENCE AND DATE ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Ref. No.: REF-2026/0001')
    run.font.size = Pt(10)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Date: [Date]')
    run.font.size = Pt(10)
    run.font.name = 'Arial'

    # === TO WHOM IT MAY CONCERN ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TO WHOM IT MAY CONCERN')
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Arial'

    # === SUBJECT ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Subject: Confirmation of Vessel Berthing, Maintenance and Management \u2014 S/Y EXAMPLE VESSEL (O.N. 000000)')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    # === SALUTATION ===
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run('Dear Sir/Madam,')
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    # === INTRO ===
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_formatted_text(p, 'We, **EXAMPLE CHARTERS LTD (Example Charters S.A.)**, a full-service yachting enterprise established in 2011 and operating from Athens, Greece, hereby confirm the following:')

    # === SECTION 1 ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('1. Vessel Identification')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_formatted_text(p, 'The sailing yacht **EXAMPLE VESSEL**, bearing Official Number **000000**, Port of Registry **Example Port**, is a 15.00-metre sailing pleasure craft built in 2005.')

    # === SECTION 2 ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('2. Ownership')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_formatted_text(p, 'The vessel is owned by **Mr. Example Owner**, a citizen of Example Country, currently residing in Example City.')

    # === SECTION 3 ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('3. Current Location')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_formatted_text(p, 'We confirm that S/Y EXAMPLE VESSEL is currently berthed in **Athens, Greece**, and has been under our care and supervision.')

    # === SECTION 4 ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('4. Services Provided')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_formatted_text(p, 'Example Charters S.A. provides full maintenance, technical support, and operational management services for S/Y EXAMPLE VESSEL on behalf of the owner, Mr. Owner. These services include but are not limited to:')

    services = [
        'Routine and preventive maintenance of hull, rigging, engine, and onboard systems',
        'Seasonal haul-out, antifouling, and winterisation',
        'Berth management and port liaison',
        'Pre-departure inspections and provisioning',
        'Technical repairs and emergency support',
    ]
    for svc in services:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run('\u2022  ' + svc)
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    # === SECTION 5 ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('5. Owner\u2019s Planned Visit')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_formatted_text(p, 'We have been informed by Mr. Owner that he intends to visit Greece together with his wife, **Mrs. Example Spouse**, and her daughter, **Miss Example Child**, in order to sail aboard S/Y EXAMPLE VESSEL and explore the Greek islands for a period of approximately two to four weeks. We are prepared to support the vessel\u2019s preparation for departure upon Mr. Owner\u2019s arrival.')

    # === CLOSING ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('We are available to provide any further information or documentation that may be required by the competent authorities.')
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    # Respectfully
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run('Respectfully,')
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    # Signature space
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run('_' * 40)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run('[Name of Authorised Signatory]')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    run = p.add_run('[Title/Position]')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run('EXAMPLE CHARTERS LTD \u2014 Example Charters S.A.')
    run.font.size = Pt(10)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    run = p.add_run('123 Example Street, Example District, 100 00, Athens, Greece')
    run.font.size = Pt(10)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    run = p.add_run('Tel: +30 000 000 0000  |  Email: info@example-charters.com')
    run.font.size = Pt(10)
    run.font.name = 'Arial'

    # Stamp placeholder
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Company Stamp / Seal]')
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    doc.save(docx_path)
    print(f'DOCX created: {docx_path}')


if __name__ == '__main__':
    md_path = str(get_outputs_dir() / 'documents' / 'example-charter-confirmation.md')
    docx_path = str(get_outputs_dir() / 'documents' / 'example-charter-confirmation.docx')
    create_charter_docx(md_path, docx_path)
