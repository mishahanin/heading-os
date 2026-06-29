#!/usr/bin/env python3
"""
Convert the National Programme DPI Proposal from markdown to a professional Word document.
Usage: python scripts/md-to-docx-proposal.py
"""

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from scripts.utils.venv import ensure_venv  # noqa: E402

ensure_venv()
from scripts.utils.docx_helpers import set_cell_shading
from scripts.utils.workspace import get_outputs_dir

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT_PATH = str(get_outputs_dir() / 'proposals' / '31C-National-Programme-DPI-Proposal-v1.md')
OUTPUT_PATH = str(get_outputs_dir() / 'proposals' / '31C-National-Programme-DPI-Proposal-v1.docx')

# 31C brand colors
BRAND_DARK = RGBColor(0x1A, 0x1A, 0x2E)   # Dark navy
BRAND_ACCENT = RGBColor(0x00, 0x7A, 0xCC)  # Blue accent
BRAND_LIGHT = RGBColor(0x4A, 0x4A, 0x5A)   # Body text grey
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "007ACC"
TABLE_ALT_BG = "F2F7FC"


def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = BRAND_LIGHT
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(20)
    h1.font.color.rgb = BRAND_DARK
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(15)
    h2.font.color.rgb = BRAND_ACCENT
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.color.rgb = BRAND_DARK
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True

    return doc



def format_table(table):
    """Apply professional formatting to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, TABLE_HEADER_BG)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(9.5)
                    run.font.name = 'Calibri'

    # Style data rows
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 0:
                set_cell_shading(cell, TABLE_ALT_BG)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = BRAND_LIGHT

    # Set borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_cover_page(doc):
    """Add a professional cover page."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('31 CONCEPT')
    run.font.size = Pt(36)
    run.font.color.rgb = BRAND_DARK
    run.font.bold = True
    run.font.name = 'Calibri'

    # Separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('_' * 60)
    run.font.color.rgb = BRAND_ACCENT
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ODUN.ONE Platform Response')
    run.font.size = Pt(22)
    run.font.color.rgb = BRAND_ACCENT
    run.font.bold = True
    run.font.name = 'Calibri'

    doc.add_paragraph()

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('National Programme Digital Infrastructure\n(Telecom and ICT)\nGovernance & Cyber Security Roadmap')
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_DARK
    run.font.name = 'Calibri'

    for _ in range(4):
        doc.add_paragraph()

    # Metadata
    meta_items = [
        ('Prepared by:', '31 Concept (31C) -- Platform Vendor'),
        ('Document Version:', '1.0'),
        ('Date:', 'March 2, 2026'),
        ('Classification:', 'Confidential -- Partner Use Only'),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label + ' ')
        run.font.size = Pt(10)
        run.font.color.rgb = BRAND_LIGHT
        run.font.bold = True
        run.font.name = 'Calibri'
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.color.rgb = BRAND_LIGHT
        run.font.name = 'Calibri'

    # Page break
    doc.add_page_break()


def parse_markdown(md_text):
    """Parse markdown into structured blocks."""
    blocks = []
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip the cover page metadata (already handled)
        if line.startswith('# 31 Concept') or line.startswith('## National Programme') or line.startswith('**Prepared by') or line.startswith('**Document Version') or line.startswith('**Date:') or line.startswith('**Classification'):
            i += 1
            continue

        # Skip TOC
        if line.strip().startswith('[') and '](#' in line:
            i += 1
            continue

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Skip the "Table of Contents" heading
        if line.strip() == '## Table of Contents':
            i += 1
            continue

        # Headings
        if line.startswith('## '):
            blocks.append(('h2', line[3:].strip()))
            i += 1
            continue
        if line.startswith('### '):
            blocks.append(('h3', line[4:].strip()))
            i += 1
            continue
        if line.startswith('#### '):
            blocks.append(('h4', line[5:].strip()))
            i += 1
            continue

        # Tables
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            # Parse table
            rows = []
            for tl in table_lines:
                if '---' in tl:
                    continue
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if cells:
                    rows.append(cells)
            if rows:
                blocks.append(('table', rows))
            continue

        # Bullet points
        if line.strip().startswith('- **') or line.strip().startswith('- '):
            bullet_lines = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('  ')):
                bullet_lines.append(lines[i].strip())
                i += 1
            blocks.append(('bullets', bullet_lines))
            continue

        # Regular paragraph
        if line.strip():
            para_lines = [line.strip()]
            i += 1
            # Collect continuation lines (but stop at headings, tables, bullets, blank lines)
            while i < len(lines):
                next_line = lines[i]
                if not next_line.strip():
                    break
                if next_line.startswith('#') or next_line.startswith('|') or next_line.strip().startswith('- '):
                    break
                para_lines.append(next_line.strip())
                i += 1
            blocks.append(('para', ' '.join(para_lines)))
            continue

        i += 1

    return blocks


def add_rich_text(paragraph, text):
    """Add text with inline bold/italic formatting."""
    # Split on bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italic
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sp in sub_parts:
                if sp.startswith('*') and sp.endswith('*') and not sp.startswith('**'):
                    run = paragraph.add_run(sp[1:-1])
                    run.italic = True
                else:
                    # Clean up markdown links [text](url) -> text
                    cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', sp)
                    if cleaned:
                        paragraph.add_run(cleaned)


def build_document():
    """Build the Word document from the markdown source."""
    with open(INPUT_PATH, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()
    setup_styles(doc)

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Add cover page
    add_cover_page(doc)

    # Parse and render content
    blocks = parse_markdown(md_text)

    for block_type, content in blocks:
        if block_type == 'h2':
            # Clean any anchor links
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
            doc.add_heading(clean, level=1)

        elif block_type == 'h3':
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
            doc.add_heading(clean, level=2)

        elif block_type == 'h4':
            clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
            doc.add_heading(clean, level=3)

        elif block_type == 'table':
            rows = content
            if len(rows) < 1:
                continue
            num_cols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=num_cols)
            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < num_cols:
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        # Clean markdown formatting for table cells
                        cleaned = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                        cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cleaned)
                        p.text = cleaned
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.space_before = Pt(2)
            format_table(table)
            doc.add_paragraph()  # Space after table

        elif block_type == 'bullets':
            for bullet in content:
                # Handle sub-bullets
                if bullet.startswith('  - ') or bullet.startswith('   - '):
                    text = bullet.strip().lstrip('- ')
                    p = doc.add_paragraph(style='List Bullet 2')
                    add_rich_text(p, text)
                elif bullet.startswith('- '):
                    text = bullet[2:]
                    p = doc.add_paragraph(style='List Bullet')
                    add_rich_text(p, text)
                # Handle continuation lines
                elif not bullet.startswith('- '):
                    # Append to previous paragraph if possible
                    pass

        elif block_type == 'para':
            p = doc.add_paragraph()
            add_rich_text(p, content)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('31 Concept | [HQ City 1] | [HQ City 2] | 31c.io')
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_ACCENT
    run.font.italic = True

    doc.save(OUTPUT_PATH)
    print(f"Word document saved to: {OUTPUT_PATH}")


if __name__ == '__main__':
    build_document()
