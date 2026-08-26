#!/usr/bin/env python3
"""Shard 09-p4: three document renderers, and the branches nothing could reach.

`md-to-docx-proposal.py` collected every bullet line through `.strip()` and then
decided sub-bullet versus top-level by LEADING SPACES. No stripped string has
leading spaces, so the `List Bullet 2` branch was unreachable and every nested
list in a client-facing proposal rendered flat. The same strip made the
continuation-line condition (`lines[i].strip().startswith('  ')`) impossible to
satisfy, so a wrapped bullet's second line escaped the list and came back as a
body paragraph in the middle of it.

The same file dropped any table row whose text contained `---` anywhere -- an em
dash written that way, a value like `rehost --- replatform` -- because the
separator test was a substring search over the whole line rather than a
per-cell one. Silently: no warning, the row is simply not in the document. And
it had no `# ` branch at all, so an H1 rendered as 10.5pt body text with its
`# ` still attached.

`md-to-docx-letter.py` picked its four address lines by RAW LINE INDEX. A blank
line inside the header block -- ordinary markdown -- advanced the index while
emitting nothing, so the last address line fell past the window and rendered as
a justified body paragraph. A short sender block drifted the other way and
swallowed the `Date:` line into address formatting.

`md-to-docx-competitive.py` put its recommendation style below the bullet
branch. `- *text*` matches the bullet regex too, and that branch `continue`s, so
the indented 9pt grey italic style never appeared in any output ever produced.
And a comment promised the metadata was captured "for a small info box" over a
body of `pass`.

`memory-hygiene.py --help` printed the blank line under its summary as its
description.
"""
from __future__ import annotations

import importlib.util
import subprocess
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


def _load(stem: str, name: str):
    spec = importlib.util.spec_from_file_location(name, ROOT / "scripts" / f"{stem}.py")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


prop = _load("md-to-docx-proposal", "md_to_docx_proposal_09p4")
letter = _load("md-to-docx-letter", "md_to_docx_letter_09p4")
comp = _load("md-to-docx-competitive", "md_to_docx_competitive_09p4")

docx = pytest.importorskip("docx")


def _blocks(md: str):
    return prop.parse_markdown(md)


def _render_proposal(md: str, tmp_path, monkeypatch):
    src = tmp_path / "in.md"
    src.write_text(md, encoding="utf-8")
    out = tmp_path / "out.docx"
    monkeypatch.setattr(prop, "INPUT_PATH", str(src))
    monkeypatch.setattr(prop, "OUTPUT_PATH", str(out))
    prop.build_document()
    return docx.Document(str(out))


def _paras(document):
    """(style name, text) for every non-empty paragraph."""
    return [(p.style.name, p.text) for p in document.paragraphs if p.text.strip()]


# ============================================================
# F1 -- the nesting branch no line could reach
# ============================================================
_NESTED = "## Section\n\n- parent item\n  - child item\n- second parent\n"


def test_a_nested_bullet_keeps_its_indentation_through_the_parser():
    bullets = next(c for k, c in _blocks(_NESTED) if k == "bullets")
    assert bullets == ["- parent item", "  - child item", "- second parent"]


def test_a_nested_bullet_renders_at_the_second_level(tmp_path, monkeypatch):
    styles = {text: style for style, text in _paras(
        _render_proposal(_NESTED, tmp_path, monkeypatch))}
    assert styles["child item"] == "List Bullet 2"


def test_a_top_level_bullet_still_renders_at_the_first_level(tmp_path, monkeypatch):
    styles = {text: style for style, text in _paras(
        _render_proposal(_NESTED, tmp_path, monkeypatch))}
    assert styles["parent item"] == "List Bullet"
    assert styles["second parent"] == "List Bullet"


_WRAPPED = "## Section\n\n- a bullet that wraps\n  onto a second line\n\nA real paragraph.\n"


def test_a_continuation_line_stays_inside_the_bullet_block():
    kinds = [b[0] for b in _blocks(_WRAPPED)]
    assert kinds.count("bullets") == 1
    bullets = next(c for k, c in _blocks(_WRAPPED) if k == "bullets")
    assert bullets == ["- a bullet that wraps", "  onto a second line"]


def test_a_continuation_line_is_joined_onto_its_bullet(tmp_path, monkeypatch):
    texts = [t for _, t in _paras(_render_proposal(_WRAPPED, tmp_path, monkeypatch))]
    assert "a bullet that wraps onto a second line" in texts
    assert "onto a second line" not in texts


def test_the_paragraph_after_the_list_is_still_a_paragraph(tmp_path, monkeypatch):
    paras = _paras(_render_proposal(_WRAPPED, tmp_path, monkeypatch))
    assert ("Normal", "A real paragraph.") in paras


def test_a_child_item_beginning_with_a_dash_keeps_it(tmp_path, monkeypatch):
    """`lstrip('- ')` takes a SET of characters; `-5% margin` became `5% margin`."""
    md = "## S\n\n- parent\n  - -5% margin\n"
    texts = [t for _, t in _paras(_render_proposal(md, tmp_path, monkeypatch))]
    assert "-5% margin" in texts


def test_a_blank_line_still_ends_the_bullet_block():
    blocks = _blocks("## S\n\n- one\n\n- two\n")
    assert [b[0] for b in blocks].count("bullets") == 2


def test_a_whitespace_only_line_also_ends_the_bullet_block():
    """It IS a blank line in markdown, and it starts with two spaces, so a
    naive indent test swallows it and everything after it."""
    blocks = _blocks("## S\n\n- one\n   \n- two\n")
    assert [b[0] for b in blocks].count("bullets") == 2


def test_a_deeper_indent_is_still_collected():
    bullets = next(c for k, c in _blocks("## S\n\n- a\n    - deep\n") if k == "bullets")
    assert bullets == ["- a", "    - deep"]


# ============================================================
# F2 -- a table row that vanished for containing three dashes
# ============================================================
_TABLE = ("## S\n\n| Item | Value |\n|---|---|\n"
          "| Migration | rehost --- replatform |\n| Other | fine |\n")


def test_a_data_row_containing_three_dashes_survives():
    rows = next(c for k, c in _blocks(_TABLE) if k == "table")
    assert ["Migration", "rehost --- replatform"] in rows


def test_the_separator_row_is_still_dropped():
    rows = next(c for k, c in _blocks(_TABLE) if k == "table")
    assert all(row != ["---", "---"] for row in rows)
    assert len(rows) == 3


# `|-|-|` is absent on purpose: the TABLE-START test requires `---` on the
# following line, so a single-dash separator is not recognised as a table at
# all. That predates this shard and is untouched by it.
@pytest.mark.parametrize("sep", ["|---|---|", "|:---|---:|", "| :-: | --- |"])
def test_every_alignment_separator_shape_is_dropped(sep):
    rows = next(c for k, c in _blocks(f"## S\n\n| A | B |\n{sep}\n| x | y |\n")
                if k == "table")
    assert rows == [["A", "B"], ["x", "y"]]


def test_a_row_whose_cells_merely_BEGIN_with_a_dash_is_kept():
    """`^[-:]+$` anchored, not `[-:]+`. Every cell here starts with a dash."""
    rows = next(c for k, c in _blocks(
        "## S\n\n| A | B |\n|---|---|\n| - first point | - second point |\n")
        if k == "table")
    assert ["- first point", "- second point"] in rows


def test_a_row_where_only_one_cell_is_dashes_is_kept():
    rows = next(c for k, c in _blocks(
        "## S\n\n| A | B |\n|---|---|\n| --- | real value |\n") if k == "table")
    assert ["---", "real value"] in rows


def test_the_surviving_row_reaches_the_document(tmp_path, monkeypatch):
    doc = _render_proposal(_TABLE, tmp_path, monkeypatch)
    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert "rehost --- replatform" in cells


# ============================================================
# F3 -- an H1 that rendered as body text, hash and all
# ============================================================
def test_a_level_one_heading_is_parsed_as_a_heading():
    assert _blocks("# Executive Summary\n\nbody\n")[0] == ("h1", "Executive Summary")


def test_the_hash_prefix_does_not_survive_into_the_text():
    assert "#" not in _blocks("# Executive Summary\n")[0][1]


def test_the_cover_page_heading_is_still_skipped():
    assert [b for b in _blocks("# 31 Concept\n\nbody\n") if b[0] == "h1"] == []


def test_a_level_one_heading_gets_a_heading_style(tmp_path, monkeypatch):
    paras = {t: s for s, t in _paras(
        _render_proposal("# Executive Summary\n\nbody\n", tmp_path, monkeypatch))}
    assert paras["Executive Summary"] == "Heading 1"


def test_the_other_heading_levels_are_unchanged():
    assert _blocks("## Two\n")[0] == ("h2", "Two")
    assert _blocks("### Three\n")[0] == ("h3", "Three")
    assert _blocks("#### Four\n")[0] == ("h4", "Four")


# ============================================================
# F4 -- a borders element attached to nothing
# ============================================================
def test_the_tblpr_property_raises_rather_than_returning_none():
    """The premise. `is not None` could never be False; it raised first."""
    from docx.oxml.exceptions import InvalidXmlError
    prop._ensure_docx()
    tbl = docx.Document().add_table(rows=1, cols=1)._tbl
    tbl.remove(tbl.find(prop.qn("w:tblPr")))
    with pytest.raises(InvalidXmlError):
        _ = tbl.tblPr


def test_the_fallback_branch_cannot_be_reached_through_format_table():
    """Not a gap in coverage: `table.alignment`, four lines above the branch,
    writes into `tblPr` too and raises first. The branch is kept and named as
    unreachable rather than tested into a shape that cannot pass."""
    from docx.oxml.exceptions import InvalidXmlError
    prop._ensure_docx()
    table = docx.Document().add_table(rows=2, cols=2)
    tbl = table._tbl
    tbl.remove(tbl.find(prop.qn("w:tblPr")))
    with pytest.raises(InvalidXmlError):
        prop.format_table(table)


def test_the_unreachability_is_written_down():
    src = (ROOT / "scripts" / "md-to-docx-proposal.py").read_text(encoding="utf-8")
    assert "UNREACHABLE, and more deeply than the audit" in src
    assert "tbl.find(qn('w:tblPr'))" in src


def test_the_ordinary_table_path_still_gets_borders():
    prop._ensure_docx()
    document = docx.Document()
    table = document.add_table(rows=2, cols=2)
    prop.format_table(table)
    assert b"tblBorders" in table._tbl.tblPr.xml.encode()


# ============================================================
# F5 -- a comment that named a page footer
# ============================================================
def test_the_closing_line_is_no_longer_called_a_footer():
    src = (ROOT / "scripts" / "md-to-docx-proposal.py").read_text(encoding="utf-8")
    assert "# Closing line, NOT a page footer" in src
    assert "\n    # Footer\n" not in src


def test_the_closing_line_still_reaches_the_document(tmp_path, monkeypatch):
    texts = [t for _, t in _paras(
        _render_proposal("## S\n\nbody\n", tmp_path, monkeypatch))]
    assert any("31c.io" in t for t in texts)


# ============================================================
# F6 -- address lines picked by raw index
# ============================================================
_PACKED = ("**MISHA HANIN**\nStreet 1\nCity\nCountry\nme@example.com\n"
           "Date: 1 Jan 2026\nThe Consul\nEmbassy\n**Subject: X**\nBody.\n")
_BLANKED = ("**MISHA HANIN**\n\nStreet 1\nCity\nCountry\nme@example.com\n"
            "Date: 1 Jan 2026\nThe Consul\nEmbassy\n**Subject: X**\nBody.\n")
_SHORT = ("**MISHA HANIN**\nStreet 1\nDate: 1 Jan 2026\n"
          "The Consul\nEmbassy\n**Subject: X**\nBody.\n")


def _letter_paras(md: str, tmp_path):
    tmp_path.mkdir(parents=True, exist_ok=True)
    src = tmp_path / "letter.md"
    src.write_text(md, encoding="utf-8")
    out = tmp_path / "letter.docx"
    letter.create_letter_docx(str(src), str(out))
    result = {}
    for p in docx.Document(str(out)).paragraphs:
        if not p.text.strip():
            continue
        sizes = {r.font.size.pt for r in p.runs if r.font.size}
        result[p.text] = (sizes, p.alignment)
    return result


def test_a_blank_line_in_the_header_does_not_shift_the_address_block(tmp_path):
    packed = _letter_paras(_PACKED, tmp_path / "a")
    blanked = _letter_paras(_BLANKED, tmp_path / "b")
    assert packed == blanked


def test_the_last_address_line_is_still_an_address_line(tmp_path):
    sizes, align = _letter_paras(_BLANKED, tmp_path)["me@example.com"]
    assert sizes == {11.0}
    assert align == docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT


@pytest.mark.parametrize("md", [_PACKED, _BLANKED, _SHORT])
def test_the_date_line_is_never_formatted_as_an_address(tmp_path, md):
    sizes, _ = _letter_paras(md, tmp_path)["Date: 1 Jan 2026"]
    assert sizes == {12.0}


def test_a_short_sender_block_does_not_push_the_recipient_into_the_address_style(tmp_path):
    """The counter must STOP at the date line, or the next line drifts back in."""
    sizes, _ = _letter_paras(_SHORT, tmp_path)["The Consul"]
    assert sizes == {12.0}


_LONG_HEADER = ("**MISHA HANIN**\nStreet 1\nStreet 2\nCity\nPostcode\n"
                "Country\nme@example.com\nDate: 1 Jan 2026\n"
                "The Consul\n**Subject: X**\nBody.\n")


def test_the_address_block_is_exactly_four_lines_long(tmp_path):
    """A counter that never advances would make every line an address line; a
    window one wider would take a fifth."""
    paras = _letter_paras(_LONG_HEADER, tmp_path)
    address = {text for text, (sizes, _) in paras.items() if sizes == {11.0}}
    assert address == {"Street 1", "Street 2", "City", "Postcode"}


def test_the_lines_past_the_address_block_fall_through_to_body(tmp_path):
    paras = _letter_paras(_LONG_HEADER, tmp_path)
    assert paras["Country"][0] == {12.0}
    assert paras["me@example.com"][0] == {12.0}


def test_the_sender_name_is_still_the_large_one(tmp_path):
    sizes, _ = _letter_paras(_BLANKED, tmp_path)["MISHA HANIN"]
    assert sizes == {14.0}


def test_the_body_is_still_justified(tmp_path):
    _, align = _letter_paras(_PACKED, tmp_path)["Body."]
    assert align == docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY


# ============================================================
# F7 / F8 -- a style below the branch that swallowed it
# ============================================================
def test_the_recommendation_branch_comes_before_the_bullet_branch():
    src = (ROOT / "scripts" / "md-to-docx-competitive.py").read_text(encoding="utf-8")
    assert src.index("# Italic-only lines (recommendations)") < src.index("# Bullet points")


def _render_competitive(md: str, tmp_path, monkeypatch):
    src = tmp_path / "in.md"
    src.write_text(md, encoding="utf-8")
    out = tmp_path / "out.docx"
    monkeypatch.setattr(comp, "INPUT", str(src))
    monkeypatch.setattr(comp, "OUTPUT", str(out))
    comp.build_docx()
    return docx.Document(str(out))


_REC = "# Title\n\n## Section\n\n- *Hold the anchor; trade scope, not price.*\n"


def test_a_recommendation_line_renders_without_a_bullet_marker(tmp_path, monkeypatch):
    doc = _render_competitive(_REC, tmp_path, monkeypatch)
    hit = [p for p in doc.paragraphs if "Hold the anchor" in p.text]
    assert hit, [p.text for p in doc.paragraphs]
    assert not hit[0].text.startswith("- ")


def test_a_recommendation_line_renders_at_the_smaller_size(tmp_path, monkeypatch):
    doc = _render_competitive(_REC, tmp_path, monkeypatch)
    hit = next(p for p in doc.paragraphs if "Hold the anchor" in p.text)
    assert {r.font.size.pt for r in hit.runs if r.font.size} == {9.0}


def test_an_ordinary_bullet_still_gets_its_marker(tmp_path, monkeypatch):
    doc = _render_competitive("# T\n\n## S\n\n- a plain bullet\n", tmp_path, monkeypatch)
    hit = next(p for p in doc.paragraphs if "a plain bullet" in p.text)
    assert hit.text.startswith("- ")


def test_a_bullet_with_italics_inside_is_not_a_recommendation(tmp_path, monkeypatch):
    """Only a line that is ENTIRELY italic qualifies."""
    doc = _render_competitive("# T\n\n## S\n\n- part *italic* part\n",
                              tmp_path, monkeypatch)
    hit = next(p for p in doc.paragraphs if "part" in p.text and "italic" in p.text)
    assert hit.text.startswith("- ")


def test_the_info_box_promise_is_gone():
    src = (ROOT / "scripts" / "md-to-docx-competitive.py").read_text(encoding="utf-8")
    assert "capture the metadata for a small info box" not in src
    assert "DISCARDED, deliberately" in src


# ============================================================
# F9 -- a help description that was a blank line
# ============================================================
def test_the_help_description_is_the_summary_line():
    # `sys.executable`, never a hardcoded `<root>/.venv/bin/python`. The
    # interpreter running this test already has the pinned dependencies, and a
    # checkout with no `.venv` (a fresh clone, a git worktree, CI) has no such
    # file at all, so the literal path made the subprocess die with
    # FileNotFoundError before argparse ever printed a line.
    out = subprocess.run(
        [sys.executable,
         str(ROOT / "scripts" / "memory-hygiene.py"), "--help"],
        capture_output=True, text=True, timeout=300)
    assert "objective-defect detector" in out.stdout, out.stdout + out.stderr


def test_the_description_is_not_the_blank_line_under_it():
    src = (ROOT / "scripts" / "memory-hygiene.py").read_text(encoding="utf-8")
    assert "__doc__.splitlines()[0] if __doc__" in src
    assert "__doc__.splitlines()[1]" not in src
