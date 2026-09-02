"""Shard 40: the OOXML child-order rule, written twice and applied to one file.

`w:pPr`, `w:tcPr` and `w:tblPr` each have a fixed child sequence in ECMA-376. A
child in the wrong position is schema-invalid, so the consumer may ignore it and
the border the code just wrote never appears.

`scripts/generate-odunone-docx.py` found this twice - in `add_bullet` and in
`add_table` - fixed it twice, and named the symptom in a comment both times
("the visible symptom is a table with no borders at all"). The rule never
reached its four sibling generators. Measured 2026-08-28 against the installed
python-docx 1.2.0, six raw appends existed and five emitted out of sequence:

    scripts/generate-client-docx.py:223    ['spacing', 'pBdr']
    scripts/generate-client-docx.py:333    ['tcW', 'shd', 'tcBorders']  (shaded rows only)
    scripts/md-to-docx-charter.py:59       ['spacing', 'pBdr']
    scripts/md-to-docx-competitive.py:176  ['spacing', 'pBdr']
    scripts/md-to-docx-competitive.py:284  ['ind', 'pBdr']
    scripts/md-to-docx-proposal.py:151     ['tblW', 'tblLook', 'tblBorders']

The sixth, `set_cell_shading`, was correct for every caller it had and is on the
funnel now anyway.

The load-bearing test here is `test_every_golden_document_is_in_schema_order`.
It reads the committed golden fixtures rather than one function, so it covers
every property container in every generated document at once, including ones no
author thought to check.
"""
import ast
import re
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
GOLDEN = ROOT / "tests" / "golden" / "docx"
SCRIPTS = ROOT / "scripts"

pytest.importorskip("docx", reason="python-docx is the optional 'documents' extra")

from scripts.utils.docx_helpers import (  # noqa: E402
    PBDR_SUCCESSORS,
    TBLBORDERS_SUCCESSORS,
    TCBORDERS_SUCCESSORS,
    TCSHD_SUCCESSORS,
    insert_in_order,
)
from tests.code_only import code_lines  # noqa: E402
from tests.repo_files import read_sources  # noqa: E402

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ============================================================
# The sequences, read from the library rather than remembered
# ============================================================
def _tag_seq(module_relpath: str, class_name: str) -> list[str]:
    """CT_*._tag_seq from the INSTALLED python-docx source.

    The attribute is `del`-eted at the end of each class body, so it cannot be
    read off the class. Parsing the source is the only way to ask the library
    what the order is, and asking is the point: a hand-copied tuple that nothing
    re-derives is a comment, not a constant.
    """
    import docx

    path = Path(docx.__file__).parent / "oxml" / module_relpath
    for node in ast.walk(ast.parse(path.read_text(encoding="utf-8"))):
        if isinstance(node, ast.ClassDef) and node.name == class_name:
            for stmt in node.body:
                if (isinstance(stmt, ast.Assign)
                        and getattr(stmt.targets[0], "id", "") == "_tag_seq"):
                    return [el.value for el in stmt.value.elts]
    raise AssertionError(f"{class_name}._tag_seq not found in {path}")


CT_PPR = _tag_seq("text/parfmt.py", "CT_PPr")
CT_TCPR = _tag_seq("table.py", "CT_TcPr")
CT_TBLPR = _tag_seq("table.py", "CT_TblPr")

CONTAINERS = {"pPr": CT_PPR, "tcPr": CT_TCPR, "tblPr": CT_TBLPR}


@pytest.mark.parametrize("constant,seq,tag", [
    (PBDR_SUCCESSORS, CT_PPR, "w:pBdr"),
    (TCBORDERS_SUCCESSORS, CT_TCPR, "w:tcBorders"),
    (TCSHD_SUCCESSORS, CT_TCPR, "w:shd"),
    (TBLBORDERS_SUCCESSORS, CT_TBLPR, "w:tblBorders"),
])
def test_our_successor_tuples_match_the_installed_library(constant, seq, tag):
    """A copied constant nothing re-derives goes stale on the next version bump.

    This is the guard that lets `docx_helpers` hold plain tuples instead of
    reaching into a private attribute at run time.
    """
    assert tuple(seq[seq.index(tag) + 1:]) == constant


# ============================================================
# The load-bearing check: every generated document, every container
# ============================================================
def _out_of_order(element, seq):
    """The children of `element` that break `seq`, or an empty list."""
    kids = [child.tag[len(W):] for child in element
            if isinstance(child.tag, str) and child.tag.startswith(W)]
    bare = [t.split(":", 1)[1] for t in seq]
    positions = [bare.index(k) for k in kids if k in bare]
    return [] if positions == sorted(positions) else kids


# Ten golden documents on 2026-09-01, one per generator plus variants. The floor
# sits below that so deleting a fixture deliberately is not a test failure, and
# far enough above one that a narrowed glob - which would leave a bare
# `assert found` perfectly green on a single surviving file - is red.
_MIN_GOLDEN_DOCUMENTS = 8


def _documents():
    found = sorted(GOLDEN.rglob("*__word_document.xml"))
    assert len(found) >= _MIN_GOLDEN_DOCUMENTS, (
        f"only {len(found)} golden document fixture(s) under {GOLDEN}; this is "
        f"the load-bearing check of the whole module and it is parametrized over "
        f"whatever the glob happens to return, so a narrowed glob passes silently")
    return found


@pytest.mark.parametrize("doc_path", _documents(), ids=lambda p: p.parent.name)
def test_every_golden_document_is_in_schema_order(doc_path):
    """Read the committed artifact, not one function.

    Six sites across four generators had this wrong and each was written by
    someone who had not read the other five. A per-function test would have
    needed six authors to think of it; this one asks the documents.
    """
    from lxml import etree

    # The rglob in `_documents()` runs at collection and this read runs at
    # execution -- minutes apart under `-n auto` -- so a fixture removed inside
    # that window would raise FileNotFoundError from inside the guard and report
    # a schema violation where nothing was violated.
    #
    # Skipped rather than parametrized over the text, and the reason is the
    # BYTES: `read_sources` yields `str`, and `etree.fromstring` refuses a `str`
    # carrying an XML encoding declaration ("Unicode strings with encoding
    # declaration are not supported"). Re-encoding to get back to bytes would put
    # this test's subject through a conversion the real consumer never does.
    #
    # A skip is safe here because the corpus floor is not this case: the
    # `assert len(found) >= _MIN_GOLDEN_DOCUMENTS` in `_documents()` is what
    # fails when a golden fixture is actually gone, and it fails at collection,
    # loudly, before any of these cases run. Only the deleted-and-restored race
    # reaches the skip, and a skip is visible in the report where a silent pass
    # would not be.
    try:
        raw = doc_path.read_bytes()
    except FileNotFoundError:
        pytest.skip(f"{doc_path} vanished between collection and read")

    tree = etree.fromstring(raw)
    offenders = []
    for name, seq in CONTAINERS.items():
        for element in tree.iter(f"{W}{name}"):
            bad = _out_of_order(element, seq)
            if bad:
                offenders.append(f"w:{name} {bad}")
    assert not offenders, (
        f"{doc_path.parent.name} emits {len(offenders)} property element(s) out "
        f"of the ECMA-376 child sequence: {offenders[:5]}")


def test_the_order_checker_refuses_a_known_offender():
    """A detector never shown a real offender is a claim, not a control.

    This is the exact shape `add_separator` produced before the fix.
    """
    from lxml import etree

    bad = etree.fromstring(
        f'<w:pPr xmlns:w="{W[1:-1]}">'
        '<w:spacing w:before="120"/><w:pBdr><w:bottom w:val="single"/></w:pBdr>'
        '</w:pPr>'.encode())
    assert _out_of_order(bad, CT_PPR) == ["spacing", "pBdr"]

    good = etree.fromstring(
        f'<w:pPr xmlns:w="{W[1:-1]}">'
        '<w:pBdr><w:bottom w:val="single"/></w:pBdr><w:spacing w:before="120"/>'
        '</w:pPr>'.encode())
    assert _out_of_order(good, CT_PPR) == []


def test_the_checker_ignores_tags_outside_the_sequence():
    """An unknown child must not be read as a position, or the checker fires on
    documents it has no opinion about."""
    from lxml import etree

    el = etree.fromstring(
        f'<w:pPr xmlns:w="{W[1:-1]}">'
        '<w:pBdr/><w:somethingNobodyDeclared/><w:spacing/>'
        '</w:pPr>'.encode())
    assert _out_of_order(el, CT_PPR) == []


# ============================================================
# The funnel itself
# ============================================================
def _ppr_with(*tags):
    from lxml import etree

    inner = "".join(f"<w:{t}/>" for t in tags)
    return etree.fromstring(f'<w:pPr xmlns:w="{W[1:-1]}">{inner}</w:pPr>'.encode())


def _new(tag):
    from lxml import etree

    return etree.fromstring(f'<w:{tag} xmlns:w="{W[1:-1]}"/>'.encode())


def test_insert_in_order_places_before_the_first_successor_present():
    parent = _ppr_with("spacing", "jc")
    insert_in_order(parent, _new("pBdr"), PBDR_SUCCESSORS)
    assert _out_of_order(parent, CT_PPR) == []
    assert [c.tag[len(W):] for c in parent] == ["pBdr", "spacing", "jc"]


def test_insert_in_order_appends_when_no_successor_is_present():
    """The old behaviour, kept for the case where it was right."""
    parent = _ppr_with("pStyle")
    insert_in_order(parent, _new("pBdr"), PBDR_SUCCESSORS)
    assert [c.tag[len(W):] for c in parent] == ["pStyle", "pBdr"]


def test_insert_in_order_picks_the_earliest_successor_not_the_first_listed():
    """The tuple is in sequence order, so the scan must stop at the first HIT,
    which is the earliest successor actually present."""
    parent = _ppr_with("jc")
    insert_in_order(parent, _new("pBdr"), PBDR_SUCCESSORS)
    assert [c.tag[len(W):] for c in parent] == ["pBdr", "jc"]


def test_insert_in_order_returns_the_element():
    """BOTH paths. This asserted only the append path - an empty parent - and a
    mutation that returned None from the addprevious branch SURVIVED, because
    the branch the callers actually take was never asked what it hands back."""
    empty = _ppr_with()
    el = _new("pBdr")
    assert insert_in_order(empty, el, PBDR_SUCCESSORS) is el

    with_successor = _ppr_with("spacing")
    el2 = _new("pBdr")
    assert insert_in_order(with_successor, el2, PBDR_SUCCESSORS) is el2


def test_shading_a_cell_goes_through_the_funnel():
    """`set_cell_shading` was a bare append. Every caller shaded before setting
    anything later, so nothing was wrong on that path - but the next caller is
    the one the funnel exists for."""
    from docx import Document

    from scripts.utils.docx_helpers import set_cell_shading

    cell = Document().add_table(rows=1, cols=1).rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(_new("vAlign"))
    set_cell_shading(cell, "F0F6FC")
    assert _out_of_order(tcPr, CT_TCPR) == []


# ============================================================
# The guard against the seventh copy
# ============================================================
# The containers python-docx pre-populates, so an append lands after something.
# NOT every identifier ending in "pr": the first draft of this pattern was
# `\w*[Pp]r`, which flagged `num_pr.append(ilvl)` in generate-odunone-docx.py -
# a call site building `w:numPr`'s own two children, in sequence order, into an
# element that was empty. A detector that cries on correct code gets switched
# off, so it names the six containers where the defect actually lives.
_CONTAINERS = ("ppr", "rpr", "tcpr", "trpr", "tblpr", "sectpr")
_APPEND = re.compile(r"([A-Za-z_][A-Za-z0-9_]*)\s*(?:\(\))?\s*\.\s*append\s*\(")


def _raw_property_appends(source: str, where: str = "<snippet>") -> list[int]:
    """Line numbers of raw appends onto an OOXML property container.

    Appending is correct only while the container holds nothing that belongs
    later, which is a fact about the call site rather than about the line. Seven
    authors each judged it and five judged it wrong, so the rule here is simply:
    go through the funnel.

    The comment strip is token-aware. `line.split("#", 1)[0]` stood here and cut
    at a `#` inside a string as readily as at a comment, so an append sitting
    after any `#`-bearing literal on its own line was invisible to the sweep.
    """
    offenders = []
    for i, code in enumerate(code_lines(source, where=where), start=1):
        for match in _APPEND.finditer(code):
            name = match.group(1).replace("_", "").lower()
            name = name.removeprefix("getoradd")
            if name in _CONTAINERS:
                offenders.append(i)
                break
    return offenders


def test_no_generator_appends_onto_a_property_container():
    offenders = {}
    scanned = 0
    for path, source in read_sources(sorted(SCRIPTS.rglob("*.py"))):
        scanned += 1
        lines = _raw_property_appends(source, where=str(path))
        if lines:
            offenders[str(path.relative_to(ROOT))] = lines
    # Corpus floor, added 2026-09-01 after measuring the gap: repointing SCRIPTS
    # at a directory that does not exist left this test GREEN, because
    # `offenders == {}` is trivially true over nothing. The detector below has a
    # negative case; the SWEEP had none, and a sweep is exactly the shape that
    # rots silently - the defect it guards is "the rule reached one of six
    # generators", so a walk that reaches none of them reads identically.
    # 385 Python files under scripts/ on 2026-09-01.
    assert scanned >= 200, (
        f"the sweep read only {scanned} source file(s) under {SCRIPTS}; it is "
        f"looking at the wrong tree, so the empty offender list below means "
        f"nothing")
    assert offenders == {}, (
        f"raw property append(s); use insert_in_order from "
        f"scripts.utils.docx_helpers: {offenders}")


def test_the_append_detector_refuses_a_known_offender():
    """The exact line removed from generate-client-docx.py.

    Dedented: the snippets are tokenized now, and a lone indented line is an
    IndentationError rather than Python. The detector reads the same text.
    """
    assert _raw_property_appends("pPr.append(pBdr)\n") == [1]
    assert _raw_property_appends("tcPr.append(tcBorders)\n") == [1]
    assert _raw_property_appends(
        "cell._tc.get_or_add_tcPr().append(shading)\n") == [1]


def test_the_append_detector_still_sees_past_a_hash_in_a_string():
    """The hole the token-aware strip closes.

    `line.split("#", 1)[0]` cut this line at `run#`, so the append never
    reached the pattern and the sweep reported a clean tree.
    """
    assert _raw_property_appends(
        'log("pass #2"); pPr.append(pBdr)\n') == [1]


def test_the_append_detector_does_not_fire_on_a_comment_or_an_inner_element():
    """`md-to-docx-proposal.py` carries the offending line inside a comment that
    documents an older bug, and `pBdr.append(bottom)` builds the border element
    itself, which has no sequence to violate."""
    assert _raw_property_appends("# `tblPr.append(borders)` decorated a node\n") == []
    assert _raw_property_appends("pBdr.append(bottom)\n") == []


def test_the_append_detector_does_not_fire_on_numpr():
    """`w:numPr` holds exactly `w:ilvl` then `w:numId`, and generate-odunone-
    docx.py appends them in that order into an element it just created empty.
    The first version of the pattern flagged both lines."""
    assert _raw_property_appends(
        'num_pr.append(parse_xml("<w:ilvl w:val=\'0\'/>"))\n') == []
