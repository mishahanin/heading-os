#!/usr/bin/env python3
"""Generate the Weekly Executive Team Meeting agenda as DOCX.

Builds a master template (placeholder date) and N dated weekly instances,
one per Monday, into outputs/operations/exec-meeting/.

Usage:
    python scripts/gen-exec-meeting-docx.py                 # master + next 12 Mondays from 2026-06-15
    python scripts/gen-exec-meeting-docx.py --start 2026-06-15 --weeks 12
    python scripts/gen-exec-meeting-docx.py --master-only   # just the reusable template
"""

import argparse
import json
import sys
from datetime import date, timedelta
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))


from scripts.utils.venv import ensure_venv  # noqa: E402

ensure_venv()
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from scripts.utils.workspace import get_outputs_dir, resolve_config_with_example  # noqa: E402

OUT_DIR = get_outputs_dir() / "operations" / "exec-meeting"

NAVY = RGBColor(0x1A, 0x2B, 0x4A)
GRAY = RGBColor(0x60, 0x60, 0x60)

# Attendee roster is per-instance DATA (real exec names): lives in the data overlay
# at <data-root>/config/exec-meeting-attendees.json; the engine ships
# scripts/exec-meeting-attendees.example.json as the generic fallback.
_ATTENDEES_FILE = resolve_config_with_example(
    "exec-meeting-attendees.json",
    Path(__file__).resolve().parent / "exec-meeting-attendees.example.json",
)
ATTENDEES = [tuple(a) for a in json.loads(_ATTENDEES_FILE.read_text(encoding="utf-8"))["attendees"]]

SECTIONS = [
    ("1. Quick Roundtable — 5–10 min",
     "Each exec: one win, top priority this week, one blocker.",
     "roundtable"),
    ("2. Business & Sales Pipeline — 15–20 min",
     "Deals, revenue forecast, partnerships, investor updates.",
     ["Pipeline movement (since last week):", "Forecast / commit this month:",
      "Partnerships:", "Investor / fundraising:", "Decisions needed:"]),
    ("3. Product & Technology — 15–20 min",
     "Roadmap progress, demos, blockers, compliance, patents.",
     ["Roadmap progress:", "Demos / releases:", "Technical blockers:",
      "Compliance / certifications:", "Patents / IP:"]),
    ("4. Finance & Operations — 10–15 min",
     "Cash position, burn rate, hires, HR, legal/compliance.",
     ["Cash position:", "Burn rate / runway:", "Hiring (open roles, candidates):",
      "HR / Tribe:", "Legal / compliance:"]),
    ("5. Strategic Topic of the Week — 15–20 min",
     "One deep dive. Pick before the meeting and note it here.",
     ["Topic:", "Owner:", "Discussion:", "Decision / direction:"]),
]


def _style_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _heading(doc, text, size=13, color=NAVY, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY
    return p


def _fill_table(doc, headers, n_rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for _ in range(n_rows):
        table.add_row()
    return table


def build(doc, meeting_date_label):
    _style_base(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Weekly Executive Team Meeting")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    srun = sub.add_run(meeting_date_label)
    srun.font.size = Pt(12)
    srun.font.color.rgb = GRAY
    _note(doc, "Cadence: Mondays, 16:00 (the configured timezone). Total budget: ~90 min.")

    # Meeting block
    _heading(doc, "Meeting", size=12)
    _chair = next((name for role, name in ATTENDEES if role == "CEO"), "")
    for label in ("Time: 16:00 local", f"Chair: {_chair} (CEO)", "Note-taker: ____________"):
        doc.add_paragraph(label, style="List Bullet")

    # Attendees
    _heading(doc, "Attendees", size=12)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(("✓", "Role", "Person")):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
    for role, person in ATTENDEES:
        row = t.add_row().cells
        row[0].text = "☐"
        row[1].text = role
        row[2].text = person

    # Agenda sections
    for head, note, body in SECTIONS:
        _heading(doc, head)
        _note(doc, note)
        if body == "roundtable":
            tbl = _fill_table(doc, ["Person", "Win", "Priority", "Blocker"], 4)
        elif isinstance(body, list):
            for item in body:
                p = doc.add_paragraph()
                rr = p.add_run(item)
                rr.bold = True

    # Footer blocks
    _heading(doc, "Action Items", size=12)
    at = _fill_table(doc, ["#", "Action", "Owner", "Due", "Status"], 4)
    for i, row in enumerate(at.rows[1:], start=1):
        row.cells[0].text = str(i)
        row.cells[4].text = "☐"

    _heading(doc, "Parking Lot", size=12)
    _note(doc, "Items raised but deferred.")
    for _ in range(3):
        doc.add_paragraph("", style="List Bullet")

    _heading(doc, "Decisions Log", size=12)
    _note(doc, "Binding decisions made this meeting.")
    for _ in range(3):
        doc.add_paragraph("", style="List Bullet")

    return doc


def mondays(start, weeks):
    # snap start to Monday (weekday 0)
    d = start - timedelta(days=start.weekday())
    return [d + timedelta(weeks=w) for w in range(weeks)]


def main():
    ap = argparse.ArgumentParser(description="Generate Weekly Executive Team Meeting DOCX agendas.")
    ap.add_argument("--start", default="2026-06-15", help="First Monday (YYYY-MM-DD). Snapped to Monday.")
    ap.add_argument("--weeks", type=int, default=12, help="Number of weekly instances to generate.")
    ap.add_argument("--master-only", action="store_true", help="Generate only the reusable master template.")
    args = ap.parse_args()

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    # Master template
    master = OUT_DIR / "weekly-exec-meeting-template.docx"
    build(Document(), "Master template — copy per week. Date: ____________").save(master)
    print(f"  master   {master.relative_to(OUT_DIR.parent.parent.parent)}")

    if args.master_only:
        return

    start = date.fromisoformat(args.start)
    for m in mondays(start, args.weeks):
        label = f"{m.isoformat()} (Monday)"
        path = OUT_DIR / f"{m.isoformat()}_exec-meeting_weekly.docx"
        build(Document(), label).save(path)
        print(f"  week     {path.relative_to(OUT_DIR.parent.parent.parent)}")


if __name__ == "__main__":
    main()
