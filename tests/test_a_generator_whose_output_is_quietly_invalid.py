"""Generators whose output is wrong in ways nothing reports.

Shard `scripts-07-p1` of the 2026-08-23/24 engine audit, and the only shard so
far whose input matched the working tree exactly -- so every finding here was
live when it was written.

The shape is one step further along than "it crashes": these produce a FILE.
The file opens. It is the reader who discovers the table has no borders, the
bullets are not bullets, or the routing row was generated from half a skill's
frontmatter with the CI gate green over the top.

Findings covered (numbering from `/tmp/audit_out3/scripts-07-p1.md`):

   1  tblBorders appended after tblLook; the fallback tblPr was detached
   2  numPr inserted before pStyle
   3  build_html took a page-logo parameter it never used
   4  the frontmatter split was not line-anchored
   5  `router: manual` was loaded, documented as a guardrail, never read
   6  `--check --flat` exited 0 without running the drift check
   7  a mapping was coerced into a list of its keys
   8  "preserved byte-for-byte" was false for any non-LF file
   9  the always-on index was written before the detail files it indexes
  10  escape_pipes treated any preceding backslash as an escape
  11  float EMU coordinates on every shape in one slide row
"""

import importlib.util
import subprocess
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


def _code_only(path: Path) -> str:
    return "\n".join(
        line for line in path.read_text(encoding="utf-8").splitlines()
        if not line.lstrip().startswith("#")
    )


def test_the_comment_stripper_keeps_the_code(tmp_path):
    f = tmp_path / "s.py"
    f.write_text('# text.split("---", 2)\nx = 1\n', encoding="utf-8")
    out = _code_only(f)
    assert "x = 1" in out
    assert 'text.split("---", 2)' not in out


def _load(name: str, rel: str):
    spec = importlib.util.spec_from_file_location(name, ROOT / rel)
    mod = importlib.util.module_from_spec(spec)
    sys.modules[name] = mod
    spec.loader.exec_module(mod)
    return mod


gsr = _load("generate_skill_router_p7a", "scripts/generate-skill-router.py")
god = _load("generate_odunone_docx_p7a", "scripts/generate-odunone-docx.py")


# ============================================================
# 1, 2 - the .docx is schema-valid where it claims to be styled
# ============================================================

@pytest.fixture(scope="module")
def docx_module():
    """python-docx or a SKIP, never an ERROR.

    `_ensure_docx()` was called blind until 2026-08-30, so on an environment
    without python-docx every test in sections 1 and 2 reported ERROR at
    fixture setup. The sibling shard `test_a_gate_that_shipped_what_it_never_read.py`
    guards its image tests with `pytest.importorskip("PIL.Image")`; two files
    of one campaign should not run opposite policies on the same question,
    and an ERROR that means "this package is not installed" buries the ones
    that mean something.
    """
    pytest.importorskip("docx", reason="python-docx (extra: documents) not installed")
    god._ensure_docx()
    return god


def test_table_borders_land_before_tbllook(docx_module, ):
    """CT_TblPr has a fixed child sequence and `tblBorders` belongs before
    `tblLook`, which python-docx always appends. A plain `append()` put it
    last; the visible symptom is a table with no borders at all."""
    import docx as docx_pkg

    d = docx_pkg.Document()
    table = docx_module.add_table(d, ["A", "B"], [["1", "2"]])
    kids = [c.tag.split("}")[1] for c in table._tbl.tblPr]

    assert "tblBorders" in kids, "the borders never reached the table"
    assert kids.index("tblBorders") < kids.index("tblLook"), (
        f"tblBorders after tblLook: {kids}")


def test_the_missing_tblpr_branch_reacts_to_what_actually_happens(docx_module):
    """`CT_Tbl.tblPr` RAISES when the child is missing; it does not return None.

    So the old `tbl.tblPr if tbl.tblPr is not None else ...` could never take
    its else branch -- and that branch was itself wrong, building a `tblPr` it
    never attached. Both halves are pinned here: the exception is the real
    signal, and the replacement element ends up inside the table.
    """
    import docx as docx_pkg
    from docx.oxml.exceptions import InvalidXmlError

    d = docx_pkg.Document()
    tbl = d.add_table(rows=1, cols=1)._tbl
    tbl.remove(tbl.tblPr)

    with pytest.raises(InvalidXmlError):
        _ = tbl.tblPr

    # The MODULE's helper, not a copy of it inline here. Reproducing the logic
    # in the test would have tested the test.
    pr = docx_module.get_or_add_tbl_pr(tbl)
    assert pr.getparent() is tbl, "the fallback tblPr is detached from the table"
    assert tbl.index(pr) == 0, "tblPr must be the first child of tbl"

    # And it is idempotent on a table that already has one.
    d2 = docx_pkg.Document()
    tbl2 = d2.add_table(rows=1, cols=1)._tbl
    assert docx_module.get_or_add_tbl_pr(tbl2) is tbl2.tblPr


def test_bullet_numbering_lands_after_pstyle(docx_module):
    """`insert(0, numPr)` put it BEFORE `pStyle`; the CT_PPr sequence requires
    `pStyle` first, and Word is free to drop the numbering."""
    import docx as docx_pkg

    d = docx_pkg.Document()
    p = docx_module.add_bullet(d, "hello")
    ppr = p._element.get_or_add_pPr()
    kids = [c.tag.split("}")[1] for c in ppr]

    assert "numPr" in kids, "the bullet lost its numbering"
    assert kids.index("pStyle") < kids.index("numPr"), f"numPr before pStyle: {kids}"


def test_the_bullet_still_carries_its_level_and_list_id(docx_module):
    """Reordering must not have dropped the two children that make it a bullet."""
    import docx as docx_pkg

    d = docx_pkg.Document()
    p = docx_module.add_bullet(d, "hello")
    num = p._element.get_or_add_pPr().find(docx_module.qn("w:numPr"))
    kids = [c.tag.split("}")[1] for c in num]
    assert kids == ["ilvl", "numId"]


# ============================================================
# 3 - a parameter that does nothing
# ============================================================

def test_build_html_has_no_parameter_it_ignores():
    """Read by AST, not by grep: the fix's own comment names the parameter it
    removed, and a text scan finds its own tombstone."""
    import ast

    tree = ast.parse((ROOT / "scripts" / "generate-partner-enablement.py")
                     .read_text(encoding="utf-8"))
    fn = next(n for n in ast.walk(tree)
              if isinstance(n, ast.FunctionDef) and n.name == "build_html")
    params = [a.arg for a in fn.args.args]
    assert params == ["header_logo_b64", "blue_b64", "black_b64"], params

    # And every parameter is actually referenced in the body.
    used = {n.id for n in ast.walk(fn) if isinstance(n, ast.Name)}
    unused = [p for p in params if p not in used]
    assert unused == [], f"parameters accepted and never referenced: {unused}"


# ============================================================
# 4 - the frontmatter split is line-anchored
# ============================================================

def _skill(tmp_path, frontmatter: str, body: str = "# Body\n"):
    md = tmp_path / "SKILL.md"
    md.write_text(f"---\n{frontmatter}---\n{body}", encoding="utf-8")
    return md


def test_a_scalar_containing_three_dashes_does_not_truncate_the_mapping(tmp_path):
    """THE REGRESSION. `text.split("---", 2)` matched anywhere, so a value like
    `handles drift --- state check` either failed the gate with a misleading
    YAML error or -- worse -- parsed a TRUNCATED mapping, dropped every key
    after it, generated the routing row from partial data, and had `--check`
    ratify the result."""
    md = _skill(tmp_path,
                'name: demo\n'
                'description: "handles drift --- state check"\n'
                'metadata:\n  version: "1.0"\n')
    data, err = gsr.parse_frontmatter(md)

    assert err == "", err
    assert data["description"] == "handles drift --- state check"
    assert data["metadata"]["version"] == "1.0", "keys after the embedded --- were dropped"


def test_a_normal_skill_still_parses(tmp_path):
    md = _skill(tmp_path, 'name: demo\ntriggers: ["a", "b"]\n')
    data, err = gsr.parse_frontmatter(md)
    assert err == ""
    assert data["triggers"] == ["a", "b"]


def test_a_file_with_no_closing_fence_is_still_refused(tmp_path):
    md = tmp_path / "SKILL.md"
    md.write_text("---\nname: demo\nno closing fence here\n", encoding="utf-8")
    _, err = gsr.parse_frontmatter(md)
    assert "missing closing" in err


def test_a_file_with_no_opening_fence_is_still_refused(tmp_path):
    md = tmp_path / "SKILL.md"
    md.write_text("# Just a heading\n---\nname: demo\n---\n", encoding="utf-8")
    _, err = gsr.parse_frontmatter(md)
    assert "missing opening" in err


# ============================================================
# 5 - `router` means something
# ============================================================

def test_the_live_registry_agrees_with_every_router_field(capsys):
    """23 skills declare `router: manual` and all 23 say NEVER auto-trigger.

    The field was loaded and never read, so an author who set it believed they
    had switched a control that did nothing. The generator now warns on any
    disagreement; this asserts the REAL corpus has none to warn about, which is
    what makes the warning meaningful rather than background noise.
    """
    rows, errors = gsr.load_routing_rows()
    assert errors == [], errors
    assert "note:" not in capsys.readouterr().err, "the live registry disagrees with itself"
    for row in rows:
        says_never = "never auto-trigger" in " ".join(row["triggers"]).lower()
        assert (row["router"] == "manual") == says_never, (
            f"{row['name']}: router={row['router']!r} but triggers "
            f"{'do' if says_never else 'do not'} say NEVER auto-trigger")


def test_a_manual_skill_whose_triggers_read_like_an_auto_one_is_flagged(tmp_path, monkeypatch, capsys):
    skills = tmp_path / "skills"
    (skills / "demo").mkdir(parents=True)
    (skills / "demo" / "SKILL.md").write_text(
        "---\nname: demo\nx-heading-routing:\n  category: Operations\n"
        '  triggers: ["do the thing"]\n  exclusions: ["N/A"]\n'
        '  compound: "No"\n  router: manual\n---\n# Body\n', encoding="utf-8")
    monkeypatch.setattr(gsr, "SKILLS_DIR", skills)
    monkeypatch.setattr(gsr, "ROOT", tmp_path)

    rows, errors = gsr.load_routing_rows()
    # ADVISORY, not fatal. Hard-failing on the wording of a free-form English
    # trigger phrase would reject "explicit invocation only", which is correct
    # prose; making it fatal is a change to how skills are authored and the
    # operator's call. Visible is the fix; blocking is the proposal.
    assert errors == [], errors
    assert "NEVER auto-trigger" in capsys.readouterr().err
    assert rows and rows[0]["router"] == "manual"


def test_an_auto_skill_claiming_never_auto_trigger_is_flagged(tmp_path, monkeypatch, capsys):
    skills = tmp_path / "skills"
    (skills / "demo").mkdir(parents=True)
    (skills / "demo" / "SKILL.md").write_text(
        "---\nname: demo\nx-heading-routing:\n  category: Operations\n"
        '  triggers: ["NEVER auto-trigger. Explicit /demo only."]\n'
        '  exclusions: ["N/A"]\n  compound: "No"\n  router: auto\n---\n# Body\n',
        encoding="utf-8")
    monkeypatch.setattr(gsr, "SKILLS_DIR", skills)
    monkeypatch.setattr(gsr, "ROOT", tmp_path)

    _, errors = gsr.load_routing_rows()
    assert errors == [], errors
    assert "NEVER auto-trigger" in capsys.readouterr().err


def test_a_router_value_that_is_neither_is_an_error(tmp_path, monkeypatch):
    skills = tmp_path / "skills"
    (skills / "demo").mkdir(parents=True)
    (skills / "demo" / "SKILL.md").write_text(
        "---\nname: demo\nx-heading-routing:\n  category: Operations\n"
        '  triggers: ["a"]\n  exclusions: ["N/A"]\n  compound: "No"\n'
        "  router: sometimes\n---\n# Body\n", encoding="utf-8")
    monkeypatch.setattr(gsr, "SKILLS_DIR", skills)
    monkeypatch.setattr(gsr, "ROOT", tmp_path)

    _, errors = gsr.load_routing_rows()
    assert any("must be 'auto' or 'manual'" in e for e in errors), errors


# ============================================================
# 6 - contradictory flags are refused, not silently resolved
# ============================================================

def _router_cli(*argv):
    return subprocess.run(
        [sys.executable, "scripts/generate-skill-router.py", *argv],
        cwd=str(ROOT), capture_output=True, text=True, timeout=180,
    )


@pytest.mark.parametrize("argv", [
    ("--check", "--flat"),
    ("--write", "--flat"),
    ("--check", "--split-by-category"),
])
def test_contradictory_flags_are_refused(argv):
    """`--check --flat` printed the monolith and exited 0 WITHOUT running the
    drift check: a CI invocation with a typo'd combination going green while
    checking nothing."""
    proc = _router_cli(*argv)
    assert proc.returncode == 2, proc.stdout[:200]
    assert "not allowed with" in proc.stderr


def test_check_alone_still_runs_the_gate():
    proc = _router_cli("--check")
    assert proc.returncode == 0, proc.stdout + proc.stderr
    assert "in sync" in proc.stdout


def test_flat_alone_still_prints_the_monolith():
    proc = _router_cli("--flat")
    assert proc.returncode == 0
    assert "| Skill | Triggers | Exclusions | Compound |" in proc.stdout


# ============================================================
# 7 - a mapping is not a list
# ============================================================

def test_a_mapping_where_a_list_belongs_is_refused():
    """Iterating a dict yields its string keys, so `triggers: {alpha: 1}` was
    coerced into `[alpha]` with the values dropped and both gates green."""
    with pytest.raises(ValueError, match="not a list"):
        gsr._as_list({"alpha": 1, "beta": 2}, field="triggers")


def test_a_real_list_and_a_bare_string_still_work():
    assert gsr._as_list(["a", "b"], field="triggers") == ["a", "b"]
    assert gsr._as_list("solo", field="triggers") == ["solo"]
    assert gsr._as_list(None, field="triggers") == []


def test_a_non_string_item_is_still_refused():
    """The item guard must survive the container guard."""
    with pytest.raises(ValueError, match="not a string"):
        gsr._as_list([{"a": 1}], field="triggers")


# ============================================================
# 8, 9 - what the module says about itself, and the write order
# ============================================================

def test_the_docstring_does_not_claim_byte_for_byte():
    """`read_text`/`write_text` translate newlines, so the claim was false for
    any non-LF file and the gate could not see the drift."""
    header = (ROOT / "scripts" / "generate-skill-router.py").read_text(
        encoding="utf-8").split('"""', 2)[1]
    assert "preserved byte-for-byte" not in header
    assert "LINE ENDINGS ASIDE" in header


def test_the_detail_files_are_written_before_the_always_on_index():
    """A failure mid-loop left the always-on index describing frontmatter the
    detail files had not caught up to. The index is what the model reads every
    session, so it goes last."""
    body = _code_only(ROOT / "scripts" / "generate-skill-router.py")
    write_fn = body.split("def cmd_split_write", 1)[1].split("\ndef ", 1)[0]
    detail_at = write_fn.index("CATEGORY_FILE_DIR.mkdir")
    index_at = write_fn.index("ROUTER_FILE.write_text")
    assert detail_at < index_at, "the always-on index is still written first"


# ============================================================
# 10 - escaping counts backslashes
# ============================================================

@pytest.mark.parametrize("raw,expected", [
    ("a|b", r"a\|b"),                 # a bare pipe is escaped
    (r"a\|b", r"a\|b"),               # an odd run: already escaped, left alone
    ("no pipes", "no pipes"),         # nothing to do
    ("a||b", r"a\|\|b"),              # both
])
def test_escape_pipes_counts_the_backslashes(raw, expected):
    assert gsr.escape_pipes(raw) == expected


def test_an_even_run_of_backslashes_leaves_the_pipe_unescaped_no_more():
    r"""`C:\\|foo` -- two backslashes, so the second escapes the first and the
    pipe is DATA. The plain lookbehind saw "a backslash" and skipped it,
    splitting the table cell into a spurious column."""
    assert gsr.escape_pipes("C:\\\\|foo") == "C:\\\\\\|foo"


def test_the_live_registry_still_renders_identically():
    """The parity fix must not have changed a single existing cell: the real
    frontmatter carries hand-escaped `\\|` in three cells."""
    proc = _router_cli("--check")
    assert proc.returncode == 0, proc.stdout + proc.stderr


# ============================================================
# 11 - every slide coordinate is an integer
# ============================================================

def test_every_coordinate_the_generator_places_is_an_integer(tmp_path, monkeypatch):
    """`Emu * 0.5` returns a plain float, and OOXML's coordinate type is
    xsd:long -- so every shape on the centred row was positioned with a
    non-integer EMU.

    Measured through the generator's own call sites: the three placement
    helpers are wrapped, `build()` is run for real, and every left/top/width/
    height it hands them is checked. This test used to import
    `pptx.util.Inches`, retype `Inches(0.6) + (cw + gap) // 2` inline and
    assert the result was an int - a fact about `//` and about python-pptx,
    established without importing, loading or calling the generator once. The
    script could have computed the row through any other expression and this
    stayed green; the file's own section 9 condemns exactly that
    ("Reproducing the logic in the test would have tested the test").

    NOT read off the emitted .pptx, which was the first attempt and is a
    guard that cannot refuse: measured 2026-08-30 against python-pptx in this
    venv, a float offset is coerced to an integer on serialise, so
    `ppt/slides/slideN.xml` carries `x="2377440"` whether the arithmetic was
    `// 2` or `* 0.5`. The truncation the shard warned about happens silently
    here, which is precisely why the check has to sit upstream of it.

    OUTPUT is redirected into tmp_path first. The module-level constant
    resolves through `get_outputs_dir()`, so calling `build()` without the
    monkeypatch writes a real deck into the operator's data overlay.
    """
    pytest.importorskip("pptx", reason="python-pptx (extra: documents) not installed")
    gtf = _load("generate_testing_framework_pptx_p7a",
                "scripts/generate-testing-framework-pptx.py")
    monkeypatch.setattr(gtf, "OUTPUT", str(tmp_path / "deck.pptx"))

    placed: list[tuple[str, str, object]] = []

    def record(helper_name, real):
        def wrapper(slide, left, top, w, h, *args, **kwargs):
            for label, value in (("left", left), ("top", top),
                                 ("width", w), ("height", h)):
                placed.append((helper_name, label, value))
            return real(slide, left, top, w, h, *args, **kwargs)
        return wrapper

    for helper in ("add_rect", "add_bar", "txt"):
        monkeypatch.setattr(gtf, helper, record(helper, getattr(gtf, helper)))

    gtf.build()

    assert len(placed) > 100, (
        f"only {len(placed)} coordinates recorded; the wrappers did not take")
    floats = [(h, label, v) for h, label, v in placed if isinstance(v, float)]
    assert not floats, (
        f"non-integer EMU coordinates reach the OOXML writer: {floats[:6]}")


def test_the_source_no_longer_multiplies_an_emu_by_a_float():
    src = _code_only(ROOT / "scripts" / "generate-testing-framework-pptx.py")
    assert "(cw + gap) * 0.5" not in src
    assert "(cw + gap) // 2" in src
