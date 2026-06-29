#!/usr/bin/env python3
"""Generate ODUN.ONE AI Monetization Use Cases DOCX using 31C corporate template."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import shutil, zipfile, os, sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from scripts.utils.venv import ensure_venv  # noqa: E402

ensure_venv()
from scripts.utils.workspace import get_datastore_dir, get_outputs_dir

# ============================================================
# Configuration
# ============================================================

# --- Colors ---
DARK_BLUE = RGBColor(0x0A, 0x1E, 0x3D)
ACCENT_BLUE = RGBColor(0x1A, 0x5C, 0xB0)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x66, 0x66, 0x66)

TMPL = str(get_datastore_dir() / "brand" / "templates"
           / "31C - Master Template (New Identity 2026 v1.00).dotx")
OUTPUT = str(get_outputs_dir() / "documents"
             / "ODUN.ONE - AI Monetization Use Cases for Telco Operators v2.docx")


# ============================================================
# Data Loading
# ============================================================


def load_template():
    """Load .dotx as .docx by fixing content type."""
    tmp = str(get_outputs_dir() / "_tmp_tpl.docx")
    shutil.copy2(TMPL, tmp)
    with zipfile.ZipFile(tmp, "r") as zin:
        ct = zin.read("[Content_Types].xml").decode("utf-8")
        ct = ct.replace("template.main+xml", "document.main+xml")
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}
        data["[Content_Types].xml"] = ct.encode("utf-8")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])
    doc = Document(tmp)
    os.remove(tmp)
    # Clear placeholder content
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    return doc


# ============================================================
# DOCX Styles
# ============================================================


def add_para(doc, text, style="Normal", bold=False, color=None, size=None,
             space_after=None, space_before=None, alignment=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if alignment is not None:
        p.alignment = alignment
    return p


# ============================================================
# DOCX Builders
# ============================================================


def add_impact(doc, text):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT_BLUE
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Cm(0.63)
    return p


def add_usecase(doc, number, title, body, impact):
    add_para(doc, f"{number}. {title}", style="Heading 2")
    add_para(doc, body, size=10, color=DARK_GRAY, space_after=4)
    add_impact(doc, impact)


# ============================================================
# Document Assembly
# ============================================================


def build():
    doc = load_template()

    # === TITLE PAGE ===
    for _ in range(6):
        add_para(doc, "", space_after=0)
    add_para(doc, "ODUN.ONE", style="Title", color=DARK_BLUE, size=36,
             space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "AI-Driven Monetization & Profitability", style="Subtitle",
             color=ACCENT_BLUE, size=20, space_after=8,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Use Cases for Telco Operators", style="Subtitle",
             color=MED_GRAY, size=16, space_after=24,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "From Deep Packet Inspection to Deep Packet Intelligence",
             color=MED_GRAY, size=11, space_after=4,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(4):
        add_para(doc, "", space_after=0)
    add_para(doc, "31 Concept", color=DARK_GRAY, size=11, space_after=2,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Proprietary & Confidential", color=MED_GRAY, size=9,
             space_after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "March 2026", color=MED_GRAY, size=9, space_after=0,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # === EXECUTIVE SUMMARY ===
    add_para(doc, "Executive Summary", style="Heading 1")
    add_para(doc,
        "ODUN.ONE transforms Deep Packet Inspection into Deep Packet Intelligence (DPI+). "
        "By embedding AI at the core of the platform \u2014 not as a bolt-on \u2014 ODUN.ONE enables telco operators "
        "to unlock revenue streams, reduce operational costs, and future-proof their networks against "
        "encrypted and AI-generated traffic. This document presents 20 use cases across six categories, "
        "each demonstrating how AI-driven network intelligence directly increases monetization and profitability. "
        "Sections I\u2013VI cover AI-native intelligence capabilities. Sections VII\u2013IX detail policy-driven revenue "
        "engineering, subscriber experience monetization, and vertical application bundles. Section X covers "
        "value added inline services \u2014 video/web optimization, caching, tethering control, and clean pipe "
        "security. All 41 use cases are proven telco revenue levers now supercharged by ODUN.ONE\u2019s AI engine.",
        size=10.5, color=DARK_GRAY, space_after=12)

    # === SECTION I ===
    doc.add_page_break()
    add_para(doc, "I. Intelligent Subscriber Monetization", style="Heading 1")

    add_usecase(doc, 1, "AI-Powered Dynamic Tiering & Upsell",
        "The AI engine continuously profiles subscriber behavior \u2014 identifying heavy video consumers, gamers, "
        "remote workers, social media power users. Instead of static bronze/silver/gold plans, the operator gets "
        "real-time intelligence on which subscribers are consistently exceeding their tier and would convert to a "
        "premium plan. The AI generates personalized upsell recommendations with predicted conversion probability.",
        "Revenue impact: 8-15% ARPU uplift on targeted segments.")

    add_usecase(doc, 2, "Application-Aware Charging (5G N6 Integration)",
        "ODUN.ONE sits on the N6 interface and enables per-application charging rules \u2014 separate data quotas for "
        "YouTube, WhatsApp, gaming, and enterprise apps. The AI learns usage patterns and recommends optimal quota "
        "splits per segment. Operators can sell \"YouTube Unlimited\" or \"Gaming Boost\" add-ons priced by actual "
        "demand curves the AI identifies.",
        "Revenue impact: New add-on revenue stream, 12-20% incremental data revenue.")

    add_usecase(doc, 3, "AI-Driven Zero-Rating & Sponsored Data Partnerships",
        "The platform identifies which applications dominate traffic in each subscriber segment. The AI recommends "
        "partnership targets \u2014 e.g., \"TikTok accounts for 34% of traffic among 18-25 year olds in Region B; "
        "sponsoring this traffic at $0.02/GB would generate $847,000/year.\" Operators approach content providers "
        "with data-backed proposals instead of guesswork.",
        "Revenue impact: New B2B2C revenue from content partnerships.")

    add_usecase(doc, 4, "Predictive Churn Prevention",
        "AnalyticsONE's behavioral models detect early churn indicators \u2014 degraded QoE patterns, declining usage, "
        "shift to competitor OTT services, tethering abuse. The system flags at-risk subscribers 30-60 days before "
        "churn and recommends retention actions (targeted offers, QoE prioritization, loyalty rewards).",
        "Revenue impact: 1-3% churn reduction = millions in preserved revenue annually.")

    # === SECTION II ===
    doc.add_page_break()
    add_para(doc, "II. Network Monetization & Capacity Optimization", style="Heading 1")

    add_usecase(doc, 5, "AI-Powered Peak/Off-Peak Revenue Optimization",
        "The AI identifies precise congestion windows by cell site, region, and application type. Operators can "
        "implement dynamic pricing \u2014 premium bandwidth during peak, incentivized off-peak usage through rewards or "
        "discounted data. The AI continuously adjusts thresholds based on actual network load.",
        "Profitability impact: 15-25% better capacity utilization, deferred CAPEX on network expansion.")

    add_usecase(doc, 6, "Predictive Capacity Forecasting",
        "Instead of reactive infrastructure investment, the AI forecasts traffic growth by application category, "
        "geography, and subscriber segment. It answers: \"You will need 40 Gbps additional capacity in Region C by Q3, "
        "driven primarily by video streaming growth among enterprise subscribers.\" This precision eliminates over-provisioning.",
        "Profitability impact: 20-30% reduction in unnecessary CAPEX through precision planning.")

    add_usecase(doc, 7, "Congestion-Aware Quality Tiering",
        "The AI creates real-time congestion maps and automatically enforces differentiated QoS policies \u2014 premium "
        "subscribers get guaranteed bandwidth while best-effort traffic is shaped intelligently. The conversational "
        "interface lets operators create these policies in natural language: \"During congestion in Region A, prioritize "
        "video for platinum subscribers and throttle P2P to 1 Mbps.\"",
        "Revenue impact: Justifies premium pricing tiers with measurable QoE guarantees.")

    add_usecase(doc, 8, "Fair Use Policy Automation",
        "AI monitors subscribers approaching data thresholds and automatically applies fair-use policies. More "
        "importantly, it identifies abuse patterns (tethering fraud, VPN tunneling to bypass caps, SIM-box fraud) "
        "that cost operators revenue. The system detects and acts in real-time \u2014 no manual investigation.",
        "Profitability impact: 3-7% revenue leakage recovery from abuse prevention.")

    # === SECTION III ===
    doc.add_page_break()
    add_para(doc, "III. Operational Cost Reduction (AI-Driven)", style="Heading 1")

    add_usecase(doc, 9, "Conversational Policy Management (AI Copilot)",
        "Legacy DPI platforms require specialized engineers to write complex policies. ODUN.ONE's AI Copilot lets "
        "operations staff create, test, and deploy policies in plain language. What previously took a senior engineer "
        "4-8 hours now takes any trained operator 5-15 minutes. The AI validates the policy, predicts its network "
        "impact, and monitors post-deployment.",
        "Profitability impact: 60-80% reduction in policy management OpEx, elimination of vendor PS dependency.")

    add_usecase(doc, 10, "AI-Driven Root Cause Analysis",
        "When customer complaints spike or QoE degrades, the AI automatically correlates traffic patterns, identifies "
        "root cause (application update, routing change, capacity issue, third-party CDN problem), and recommends the "
        "fix. Operators ask: \"Why did video quality drop in the Eastern region last Tuesday?\" and get a precise, "
        "evidence-backed answer in seconds.",
        "Profitability impact: 40-60% reduction in mean-time-to-resolution, lower customer care costs.")

    add_usecase(doc, 11, "Automated Report Generation",
        "Legacy vendors charge $100,000+ for custom report development. ODUN.ONE generates any report via natural "
        "language query with sub-second granularity. Board-level dashboards, regulatory submissions, partner settlement "
        "reports \u2014 all generated on demand, no professional services engagement.",
        "Profitability impact: Elimination of $100K-$500K/year in custom reporting costs.")

    add_usecase(doc, 12, "Self-Optimizing Network Policies",
        "The AI monitors policy effectiveness in real-time. If a traffic shaping rule is causing unintended QoE "
        "degradation, the system alerts, recommends adjustment, and (with approval) auto-corrects. Policies evolve "
        "based on outcomes, not static configuration.",
        "Profitability impact: Fewer incidents, fewer rollbacks, lower operational risk.")

    # === SECTION IV ===
    doc.add_page_break()
    add_para(doc, "IV. Security-as-Revenue", style="Heading 1")

    add_usecase(doc, 13, "Enterprise Security Intelligence Service",
        "The AI's anomaly detection capabilities (DDoS identification, malware traffic signatures, botnet patterns) "
        "can be packaged as a managed security service for enterprise customers. The operator becomes a security "
        "intelligence provider \u2014 not just a pipe.",
        "Revenue impact: New managed security revenue stream, $5-15/enterprise subscriber/month.")

    add_usecase(doc, 14, "Fraud Detection & Revenue Assurance",
        "Real-time AI detection of SIM-box fraud, CLI spoofing, interconnect bypass, and premium rate abuse. The "
        "system learns fraud patterns and adapts faster than fraudsters can pivot. Every detected fraud incident is "
        "recovered revenue.",
        "Profitability impact: 2-5% revenue recovery from fraud elimination.")

    add_usecase(doc, 15, "Regulatory Compliance Automation",
        "AI-powered content filtering and subscriber profiling with full audit trails. Instead of manual compliance "
        "processes, the platform automatically enforces regulatory requirements and generates compliance reports. "
        "Avoids regulatory fines and reduces compliance headcount.",
        "Profitability impact: Reduced compliance costs and zero-fine risk.")

    # === SECTION V ===
    doc.add_page_break()
    add_para(doc, "V. Data & Intelligence Monetization", style="Heading 1")

    add_usecase(doc, 16, "Anonymized Traffic Intelligence API",
        "The AI generates aggregated, anonymized insights about application usage, mobility patterns, and digital "
        "behavior trends. These intelligence APIs can be sold to: urban planners (population movement), retailers "
        "(foot traffic correlation), advertisers (audience segmentation), and app developers (market sizing).",
        "Revenue impact: Entirely new data-as-a-service revenue line.")

    add_usecase(doc, 17, "Content Provider Intelligence Reports",
        "AI-generated reports showing content providers how their services perform on the network \u2014 buffering rates, "
        "quality scores, regional performance. Valuable intelligence that CDNs and OTT providers will pay for or "
        "trade peering/caching benefits against.",
        "Revenue impact: Leverage for better peering agreements + direct intelligence sales.")

    add_usecase(doc, 18, "QoE Benchmarking & SLA Intelligence",
        "AI continuously measures per-application QoE (AppScore) across every subscriber, cell site, and time window. "
        "This intelligence enables operators to: (a) sell SLA-backed enterprise services with real data, (b) benchmark "
        "against competitors, (c) identify infrastructure investment priorities with precision.",
        "Revenue impact: Premium SLA-backed enterprise services command 30-50% price premium.")

    # === SECTION VI ===
    doc.add_page_break()
    add_para(doc, "VI. Encrypted & AI-Generated Traffic Mastery", style="Heading 1")

    add_usecase(doc, 19, "Encrypted Traffic Monetization",
        "With 95%+ of traffic now encrypted, operators are flying blind with legacy tools. ODUN.ONE's multi-layered "
        "AI classification (pattern matching + behavioral + statistical + ML + deep learning) restores visibility "
        "without breaking encryption. This means all monetization use cases above work even on encrypted traffic \u2014 "
        "a capability competitors cannot match at this accuracy.",
        "Strategic value: Without this, none of the above use cases work on modern traffic.")

    add_usecase(doc, 20, "AI-Generated Traffic Detection & Management",
        "As AI tools (ChatGPT, Copilot, AI video generation) drive new traffic patterns, ODUN.ONE's AI identifies "
        "and classifies these emerging categories before competitors even have signatures. Operators can proactively "
        "create \"AI Workspace\" bundles or enterprise AI traffic packages.",
        "Revenue impact: First-mover advantage on the fastest-growing traffic category.")

    # === SECTION VII ===
    doc.add_page_break()
    add_para(doc, "VII. Policy-Driven Revenue Engineering", style="Heading 1")

    add_usecase(doc, 21, "Location-Based Policy & Regional Bundles",
        "ODUN.ONE correlates subscriber traffic with geographic data (LAC/TAC, Cell ID, GPS) to enable "
        "precision location-based offers. Operators create regional bundles \u2014 campus zones with unlimited "
        "educational content, airport lounges with premium streaming, tourist hotspots with travel app "
        "packages. The AI identifies location-traffic correlations automatically: \"Subscribers at Mall X "
        "consume 3.2x more video between 14:00\u201318:00 \u2014 recommend a location-triggered streaming upsell.\" "
        "Geo-fenced policies update dynamically as network topology changes.",
        "Revenue impact: 5-12% incremental revenue from location-aware offers + regional pricing arbitrage.")

    add_usecase(doc, 22, "Content-Based Policy & URL Classification",
        "ControlONE classifies traffic by content category (entertainment, news, adult, education, "
        "social, e-commerce) in real-time across encrypted flows. This enables content-tier pricing \u2014 "
        "basic plans include news and messaging, premium unlocks HD streaming and gaming. The AI continuously "
        "refines classification accuracy and detects new content categories as they emerge, eliminating "
        "manual URL database maintenance that legacy systems require.",
        "Revenue impact: Enables content-tier plan differentiation, 10-18% uplift on premium content bundles.")

    add_usecase(doc, 23, "Device-Based Policy & Plans",
        "The platform identifies device model, type (smartphone, tablet, laptop, IoT), and OS in real-time. "
        "Operators create device-optimized plans: tablet-specific streaming bundles with higher bandwidth, "
        "IoT-device data caps, laptop tethering packages. The AI detects device mix shifts across the subscriber "
        "base and recommends new device-tier products before competitors react. Device intelligence also feeds "
        "into handset financing programs and manufacturer partnership negotiations.",
        "Revenue impact: Device-specific bundles command 8-15% price premium over generic plans.")

    add_usecase(doc, 24, "SIM Swap & Device Swap Detection",
        "ODUN.ONE\u2019s AI monitors IMEI/IMSI binding patterns in real-time. When a subscriber\u2019s SIM moves to "
        "a new device or a new SIM appears in a branded/financed handset, the system triggers configurable "
        "actions: fraud alerts, plan re-evaluation, device-lock enforcement, or retention offers. For "
        "operator-financed devices, this prevents subsidy fraud. The AI learns normal swap patterns (upgrades, "
        "repairs) versus suspicious patterns (theft, unauthorized resale) with increasing accuracy.",
        "Profitability impact: 1-3% reduction in device subsidy fraud + handset financing risk mitigation.")

    add_usecase(doc, 25, "Volume-Based Multi-Stage Shaping",
        "Instead of binary \"full speed / fully throttled\" behavior, ODUN.ONE applies intelligent graduated "
        "shaping based on quota consumption. At 80% usage, P2P traffic is deprioritized. At 100%, video "
        "drops to SD quality while messaging stays unrestricted. At 120%, the system presents an upsell "
        "offer. The AI optimizes these thresholds per subscriber segment based on churn correlation \u2014 "
        "finding the exact shaping profile that maximizes revenue without triggering churn.",
        "Revenue impact: 15-25% increase in data top-up purchases through intelligent step-down shaping.")

    add_usecase(doc, 26, "Time-Based Policies & Happy Hour Promotions",
        "ControlONE enforces time-aware policies: off-peak unlimited windows (\"Happy Hour\" from 01:00\u201306:00), "
        "weekend bonus data (\"X-Day\" promotions), duration-based passes (\"4-hour gaming pass for $1.99\"), "
        "and time-of-day pricing tiers. The AI analyzes network load patterns and subscriber behavior to "
        "recommend optimal promotional windows that maximize revenue without causing congestion. Happy Hour "
        "timing adapts automatically to seasonal traffic shifts.",
        "Revenue impact: 8-12% off-peak monetization + promotional bundle attach rate of 20-35%.")

    add_usecase(doc, 27, "RAT-Based Policies (3G/4G/5G Differentiation)",
        "ODUN.ONE detects the Radio Access Technology each subscriber uses and applies differentiated policies. "
        "5G subscribers get premium QoS and exclusive content bundles. 4G users receive standard tiers. "
        "3G connections are optimized for lightweight apps with compressed content. The AI identifies "
        "subscribers consistently on 4G who would benefit from 5G migration offers, feeding the operator\u2019s "
        "network upgrade ROI by accelerating voluntary migration to newer, more efficient RATs.",
        "Revenue impact: 5G premium tier commands 20-40% ARPU uplift; accelerates 3G/4G sunset timeline.")

    # === SECTION VIII ===
    doc.add_page_break()
    add_para(doc, "VIII. Subscriber Experience & Retention Monetization", style="Heading 1")

    add_usecase(doc, 28, "Bill Shock Prevention for Roaming Bundles",
        "When subscribers cross borders, ODUN.ONE\u2019s real-time policy engine detects roaming state and "
        "immediately applies configurable safeguards: automatic roaming bundle activation, spend caps, "
        "low-balance alerts, and application-specific roaming policies (e.g., allow maps and messaging, "
        "block video streaming until the subscriber opts in). The AI predicts travel patterns from "
        "historical data and proactively offers roaming bundles before departure \u2014 converting a cost "
        "protection feature into a pre-trip revenue opportunity.",
        "Revenue impact: 30-50% increase in roaming bundle attachment + near-zero bill disputes.")

    add_usecase(doc, 29, "Parental Control as Premium Service",
        "ODUN.ONE\u2019s content classification engine enables operator-branded parental control products. "
        "URL/app category filtering, screen time limits by application type, safe search enforcement, "
        "and activity reporting \u2014 all managed through a subscriber-facing portal. The AI adds behavioral "
        "intelligence: detecting new apps popular with minors before they appear in manual blocklists, "
        "and identifying risky usage patterns that static filters miss. Packaged as a $3-7/month premium "
        "add-on with strong market demand.",
        "Revenue impact: New recurring revenue line, $3-7/subscriber/month with 15-25% household attach rate.")

    add_usecase(doc, 30, "Bandwidth-on-Demand (Turbo Button)",
        "Subscribers purchase temporary speed boosts via a single tap \u2014 \"2 hours of 100 Mbps for $0.99\" "
        "or \"4K streaming pass for tonight\u2019s match.\" ODUN.ONE\u2019s SOAP/REST API integration with BSS/OSS "
        "enables real-time provisioning in under 2 seconds. The AI identifies optimal Turbo moments: "
        "when a subscriber\u2019s video is buffering, when they enter a congested cell, or when a major "
        "live event starts \u2014 triggering contextual offers at the exact moment willingness-to-pay peaks.",
        "Revenue impact: $0.50-2.00 per micro-transaction, 10-20% of active subscribers engage monthly.")

    add_usecase(doc, 31, "Shared Group Quota (Family & Enterprise)",
        "ControlONE manages shared data pools across multiple SIMs \u2014 family plans where parents control "
        "allocation per child device, or enterprise pools where departments share a corporate quota with "
        "per-user visibility. The AI optimizes allocation recommendations: \"Your family consistently "
        "under-uses the 50 GB plan \u2014 recommend 30 GB shared + Gaming Boost add-on for Device B, saving "
        "the family $4/month while increasing your net margin.\" Shared quotas have 25-40% lower churn "
        "than individual plans due to switching cost.",
        "Revenue impact: 25-40% churn reduction on group plans + higher effective ARPU per SIM.")

    add_usecase(doc, 32, "Real-Time End-User Notifications",
        "ODUN.ONE triggers subscriber-facing notifications at critical moments: quota thresholds "
        "(80%, 95%, 100%), QoE degradation events, available upsell offers, promotional activations, "
        "and roaming transitions. The AI personalizes notification timing, channel, and content per "
        "subscriber profile \u2014 learning that Subscriber A converts on SMS at 95% quota, while Subscriber B "
        "responds to in-app push at 80%. Every notification becomes a monetization touchpoint rather than "
        "a service alert.",
        "Revenue impact: 2-5x increase in top-up and add-on conversion rates versus static notifications.")

    # === SECTION IX ===
    doc.add_page_break()
    add_para(doc, "IX. Vertical Application Monetization", style="Heading 1")

    add_usecase(doc, 33, "Gaming Bundles & eSports Packages",
        "ODUN.ONE classifies gaming traffic by title and genre (mobile, cloud gaming, competitive eSports) "
        "with protocol-level precision. Operators create \"Gaming Unlimited\" bundles with QoS prioritization "
        "for latency-sensitive titles, zero-rated game updates, and tournament-day boost passes. The AI "
        "identifies the operator\u2019s gaming subscriber segment (typically 15-25% of base), their peak hours, "
        "preferred titles, and price sensitivity \u2014 enabling gaming bundles priced at maximum conversion. "
        "Partnership intelligence: \"Fortnite generates 847 GB/day across your network \u2014 negotiate a "
        "sponsored data deal with Epic Games.\"",
        "Revenue impact: Gaming bundles yield 2-3x margin versus generic data due to perceived value premium.")

    add_usecase(doc, 34, "VoIP & Video Calling Monetization",
        "Rather than treating OTT voice as a revenue threat, ODUN.ONE enables operators to monetize it "
        "directly. The platform classifies VoIP traffic by provider (WhatsApp, Telegram, Zoom, Teams) "
        "and call type (voice vs. video). Operators offer \"Unlimited VoIP\" add-ons, enterprise \"UC "
        "Optimization\" packages with guaranteed QoS for business calls, or zero-rated calling bundles "
        "in partnership with specific providers. The AI quantifies the revenue displacement: \"VoIP "
        "carries 21,000 minutes/month \u2014 a $0.02/minute bundle recovers $12,600/month.\"",
        "Revenue impact: Converts OTT voice from revenue cannibalizer to $0.50-2.00/subscriber/month add-on.")

    add_usecase(doc, 35, "Social Network Zero-Rating (SNS Zero)",
        "ODUN.ONE enables surgical zero-rating of specific social platforms \u2014 WhatsApp messaging (excluding "
        "media), Facebook Lite, Twitter/X text \u2014 as entry-level plan differentiators. The AI measures "
        "the actual cost of zero-rated traffic versus subscriber acquisition value, ensuring every free "
        "byte is an investment with measurable return. Advanced: the system detects when users tunnel "
        "non-social traffic through zero-rated apps (a common abuse vector) and applies corrective policy "
        "automatically.",
        "Revenue impact: 10-20% subscriber acquisition uplift on entry plans + abuse prevention preserves margin.")

    add_usecase(doc, 36, "Student & Campus Zone Plans",
        "Location + device + content policies combine to create university-targeted products: unlimited "
        "educational content on campus, zero-rated LMS and library access, device-specific student plans, "
        "and time-based study/entertainment splits. The AI identifies student segments by behavioral "
        "patterns (not just location) and recommends campus-specific bundles based on actual usage: "
        "\"University X students consume 73% video and 18% educational content \u2014 offer a Video + Study "
        "combo at $7.99/month.\" Campus partnerships create B2B2C co-funding opportunities.",
        "Revenue impact: Captures the 18-25 high-usage demographic with tailored plans at 15-25% higher attach rate.")

    # === SECTION X ===
    doc.add_page_break()
    add_para(doc, "X. Value Added Inline Services", style="Heading 1")

    add_usecase(doc, 37, "Video Optimization & Adaptive Streaming Control",
        "ODUN.ONE performs real-time video traffic optimization inline \u2014 detecting streaming protocols "
        "(HLS, DASH, QUIC-based) and applying intelligent quality adaptation based on subscriber tier, "
        "network congestion, and device capability. Premium subscribers get unthrottled 4K; standard tiers "
        "receive optimized HD that looks identical on a 6-inch screen but consumes 40-60% less bandwidth. "
        "The AI learns per-subscriber viewing patterns and pre-positions optimization profiles. During "
        "congestion events, the system dynamically adjusts video bitrates across thousands of sessions "
        "simultaneously \u2014 preserving QoE perception while freeing capacity.",
        "Profitability impact: 30-50% reduction in video-driven bandwidth costs + premium video tier monetization.")

    add_usecase(doc, 38, "Web Optimization & Content Acceleration",
        "Inline web optimization compresses images, minifies code, and prioritizes above-the-fold content "
        "delivery for subscribers on constrained connections. The AI selectively applies optimization "
        "intensity based on RAT type (heavier on 3G, lighter on 5G), device screen resolution, and "
        "subscriber plan tier. Enterprise subscribers get unmodified pass-through; value-tier subscribers "
        "get optimized content that loads 2-3x faster on congested cells. This directly reduces page "
        "abandonment rates and improves perceived network quality without infrastructure investment.",
        "Profitability impact: Improved NPS scores, reduced churn on value tiers, 15-25% bandwidth savings on web traffic.")

    add_usecase(doc, 39, "Intelligent Caching & Content Delivery",
        "ODUN.ONE identifies the most-consumed content across the subscriber base and enables intelligent "
        "edge caching strategies. The AI predicts content popularity \u2014 identifying viral videos, trending "
        "app updates, and recurring download patterns before they peak. Operators reduce transit costs by "
        "serving cached content locally. The platform provides cache hit-rate analytics and recommendations: "
        "\"Caching the top 50 YouTube videos in Region A would reduce 23% of your transit traffic, saving "
        "an estimated $187,000/month in interconnect fees.\"",
        "Profitability impact: 15-30% reduction in transit/peering costs through AI-guided caching strategy.")

    add_usecase(doc, 40, "Tethering Detection & Control Policy",
        "ODUN.ONE detects tethering and mobile hotspot usage through multi-signal analysis \u2014 TTL "
        "anomalies, user-agent fingerprinting, OS-level indicators, and behavioral patterns. The AI "
        "distinguishes legitimate tethering (subscriber\u2019s own laptop) from abuse (unauthorized hotspot "
        "resale). Operators can: (a) block tethering on plans that exclude it, (b) offer paid tethering "
        "add-ons auto-triggered on detection, (c) apply separate tethered-device quotas, or (d) shape "
        "tethered traffic during congestion. The system adapts as OS vendors change tethering behaviors, "
        "maintaining detection accuracy without manual signature updates.",
        "Revenue impact: Converts tethering abuse into paid add-on revenue, $1-3/subscriber/month on affected base.")

    add_usecase(doc, 41, "Clean Pipe & Anti-Virus as Managed Service",
        "ODUN.ONE\u2019s inline threat detection identifies malware downloads, phishing domains, botnet "
        "command-and-control traffic, and cryptomining patterns at the network level \u2014 protecting "
        "subscribers before threats reach their devices. Packaged as a B2C \"Safe Browsing\" premium "
        "add-on or a B2B \"Clean Pipe\" enterprise service, the operator monetizes security without "
        "requiring endpoint software. The AI continuously updates threat intelligence, detecting "
        "zero-day patterns through behavioral analysis rather than signature matching alone. "
        "Enterprise clean pipe services include DDoS mitigation, content filtering, and compliance "
        "reporting \u2014 commanding premium pricing.",
        "Revenue impact: $2-5/consumer subscriber or $8-20/enterprise subscriber as premium security add-on.")

    # === SUMMARY TABLE ===
    doc.add_page_break()
    add_para(doc, "Monetization Impact Summary", style="Heading 1")

    rows = [
        ["Category", "Use Cases", "Revenue / Savings Potential"],
        ["Subscriber Monetization", "1 \u2013 4", "8-20% ARPU uplift + churn reduction"],
        ["Network Optimization", "5 \u2013 8", "20-30% CAPEX deferral + leakage recovery"],
        ["Operational Efficiency", "9 \u2013 12", "40-80% OpEx reduction in DPI operations"],
        ["Security-as-Revenue", "13 \u2013 15", "New revenue stream + 2-5% fraud recovery"],
        ["Data Monetization", "16 \u2013 18", "New data-as-a-service business line"],
        ["Future-Proofing", "19 \u2013 20", "Prerequisite for all above on modern networks"],
        ["Policy-Driven Revenue", "21 \u2013 27", "5-40% uplift via location/device/time/RAT policies"],
        ["Subscriber Experience", "28 \u2013 32", "Roaming, parental, turbo, shared quota monetization"],
        ["Vertical App Bundles", "33 \u2013 36", "2-3x margin on gaming, VoIP, social, campus bundles"],
        ["Value Added Inline Services", "37 \u2013 41", "30-50% bandwidth savings + security add-on revenue"],
    ]

    table = doc.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # table.style = "Table Grid"  # not available in template

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9.5)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear", qn("w:color"): "auto",
                    qn("w:fill"): "0A1E3D"})
                shading.append(shd)
            else:
                run.font.color.rgb = DARK_GRAY
                if i % 2 == 0:
                    shading = cell._element.get_or_add_tcPr()
                    shd = shading.makeelement(qn("w:shd"), {
                        qn("w:val"): "clear", qn("w:color"): "auto",
                        qn("w:fill"): "F0F4F8"})
                    shading.append(shd)

    # Closing
    add_para(doc, "", space_after=12)
    p = doc.add_paragraph()
    run = p.add_run(
        "41 use cases. One platform. ODUN.ONE does not just inspect packets \u2014 it turns network traffic into an "
        "AI-powered intelligence asset that generates revenue, reduces cost, and creates entirely new business "
        "models. From dynamic subscriber tiering to location-aware bundles, from turbo buttons to campus zones "
        "\u2014 the AI is not an add-on; it is the engine that makes every use case above possible at speed and "
        "scale that manual operations cannot match.")
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_after = Pt(24)

    # Save
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


# ============================================================
# Main / CLI
# ============================================================


if __name__ == "__main__":
    build()
