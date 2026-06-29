#!/usr/bin/env python3
"""
Generate ODUN.ONE Complete Capability Document as DOCX.
Opens the 31C Master Template DOCX (v1.01) which preserves all styles,
header (logo), footer (31C branding), fonts (GT Standard M), and page setup.
Clears template body content and writes the full document using native styles.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from scripts.utils.venv import ensure_venv  # noqa: E402

ensure_venv()
from scripts.utils.docx_helpers import set_cell_shading
from scripts.utils.workspace import get_datastore_dir, get_outputs_dir

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy
import os

# ============================================================
# Configuration / Paths
# ============================================================
TEMPLATE = str(get_datastore_dir() / "brand" / "templates" /
               "31C - Master Template (New Identity 2026 v1.01).docx")
OUTPUT = str(get_outputs_dir() / "deliverables" / "documents" /
             "ODUN.ONE - AI-Powered Sovereign Intelligence Platform - Complete Capability Document.docx")

# Brand colors from template analysis
# Heading 1: GT Standard M Medium, 18pt, #747DBE (purple-blue)
# Heading 2: GT Standard M Medium, 14pt, #423BFF (blue)
# Normal: GT Standard M Light, 11pt, #000000, justified, space_after=8pt
# Bullets: List Paragraph with numId=1, orange accent #FF9235
# Title: GT Standard M Medium, 36pt, bold, centered, #747DBE
# Cover metadata: orange #FF9235
TABLE_HEADER_BG = "423BFF"   # Heading 2 blue for table headers
TABLE_ALT_BG = "F0F4FF"     # Light blue-purple tint
ORANGE = RGBColor(0xFF, 0x92, 0x35)
PURPLE = RGBColor(0x74, 0x7D, 0xBE)
BLUE = RGBColor(0x42, 0x3B, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ============================================================
# Helpers / Utilities
# ============================================================
def clear_body(doc):
    """Remove all body content from the document while preserving sections/headers/footers."""
    body = doc.element.body
    # Keep the sectPr (last child usually) but remove everything else
    sect_pr = body.findall(qn('w:sectPr'))
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'sectPr':
            body.remove(child)


def add_heading(doc, text, level=1):
    """Add heading using template's native Heading styles."""
    h = doc.add_heading(text, level=level)
    # The template styles handle font/color/size automatically
    return h


def add_normal(doc, text, bold=False, italic=False):
    """Add a Normal-style paragraph. Template defines GT Standard M Light 11pt justified."""
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return p


def add_rich_para(doc, parts):
    """Add paragraph with mixed formatting. parts = [(text, bold, italic), ...]"""
    p = doc.add_paragraph(style='Normal')
    for text, bold, italic in parts:
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
    return p


def add_bullet(doc, text, bold_prefix=None):
    """Add bullet using template's List Paragraph + numId=1 (orange bullet)."""
    p = doc.add_paragraph(style='List Paragraph')
    # Apply the numbering from the template (numId=1, ilvl=0)
    ppr = p._element.get_or_add_pPr()
    num_pr = parse_xml(f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>')
    ppr.insert(0, num_pr)

    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.color.rgb = ORANGE
        run_n = p.add_run(text)
        run_n.font.color.rgb = ORANGE
    else:
        run = p.add_run(text)
        run.font.color.rgb = ORANGE
    return p


# ============================================================
# Table & Page Helpers
# ============================================================
def add_table(doc, headers, rows, col_widths_cm=None):
    """Add a styled table matching 31C brand."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Use Normal Table (only table style in template) and add borders via XML
    table.style = 'Normal Table'
    # Add table borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>
    </w:tblBorders>'''
    tbl_pr.append(parse_xml(borders_xml))

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # Handle **bold** markers
            if "**" in cell_text:
                import re
                parts = re.split(r'\*\*', cell_text)
                for pi, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(9.5)
                        run.bold = (pi % 2 == 1)
            else:
                run = p.add_run(cell_text)
                run.font.size = Pt(9.5)
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    # Column widths
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table


def add_page_break(doc):
    """Insert a page break."""
    doc.add_page_break()


# ============================================================
# Rendering / Document Builder
# ============================================================
def build_document():
    """Build the complete DOCX from the 31C template."""
    print(f"Loading template: {TEMPLATE}")
    doc = Document(TEMPLATE)

    # Clear all template placeholder content
    clear_body(doc)

    # ============================================================
    # COVER PAGE (matching template structure)
    # ============================================================
    # Title paragraph with line breaks (matching template's P0 structure)
    title_p = doc.add_paragraph(style='Title')
    # Add vertical spacing via line breaks (template uses ~7 newlines before title)
    for _ in range(7):
        run = title_p.add_run("\n")
        run.font.color.rgb = PURPLE
    run = title_p.add_run("ODUN.ONE")
    run.font.color.rgb = PURPLE

    # Subtitle (template uses Normal with Heading 2 size/color)
    sub_p = doc.add_paragraph(style='Normal')
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("AI-Powered Sovereign Deep Packet Intelligence")
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE

    # Tagline
    tag_p = doc.add_paragraph(style='Normal')
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag_p.add_run("The Platform That Sees Everything, Understands Intent, and Acts Instantly")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph(style='Normal')

    # Divider tagline
    div_p = doc.add_paragraph(style='Normal')
    div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = div_p.add_run("From Deep Packet Inspection to Deep Packet Intelligence")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE

    # Spacing
    for _ in range(3):
        doc.add_paragraph(style='Normal')

    # Metadata (orange, matching template style)
    for text in ["Confidential -- For Authorized Recipients Only",
                 "Version: 1.0",
                 "Date: March 2026"]:
        mp = doc.add_paragraph(style='Normal')
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = mp.add_run(text)
        run.font.color.rgb = ORANGE

    add_page_break(doc)

    # ============================================================
    # TABLE OF CONTENTS
    # ============================================================
    add_heading(doc, "TABLE OF CONTENTS", 1)

    toc_items = [
        "1.  Executive Summary",
        "2.  The Market Imperative",
        "3.  Platform Overview: The ODUN Methodology",
        "4.  Core Platform Modules",
        "5.  AI & Machine Learning Architecture",
        "6.  AI Training Pipeline: Sovereign Intelligence Without Limits",
        "7.  Intelligence Modules",
        "8.  Telecom Use Cases: AI-Driven Network Intelligence",
        "9.  Law Enforcement & National Security Use Cases",
        "10. AI Analytics for Telco & Law Enforcement: Training Scenarios",
        "11. Deployment Architecture & Data Sovereignty",
        "12. Technical Specifications",
        "13. Integration Architecture",
        "14. Security, Compliance & Governance",
        "15. Why ODUN.ONE",
    ]
    for item in toc_items:
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(item)
        p.paragraph_format.space_after = Pt(3)

    add_page_break(doc)

    # ============================================================
    # SECTION 1: EXECUTIVE SUMMARY
    # ============================================================
    add_heading(doc, "EXECUTIVE SUMMARY", 1)

    add_normal(doc, "ODUN.ONE is the world's first AI-native sovereign deep packet intelligence platform. Built from a clean-slate architecture with zero legacy code, it transforms raw network traffic into actionable intelligence at carrier scale -- observing, decoding, understanding, and navigating terabits of data per second in real time.")

    add_heading(doc, "What Makes ODUN.ONE Fundamentally Different", 2)

    add_normal(doc, "Because the platform is deployed 100% on-premises, within the customer's sovereign infrastructure, with zero external dependencies -- the AI engine can be trained on the customer's actual network data without any limitation. No data leaves the country. No foreign cloud processes the intelligence. No third-party vendor has access to training data or model outputs.")

    add_normal(doc, "Full sovereignty unlocks capabilities that cloud-dependent platforms cannot offer:", bold=True)

    for b in [
        "Train AI models on national-scale traffic patterns unique to your network",
        "Build custom behavioral profiles for applications, subscribers, and threat actors specific to your operational environment",
        "Develop proprietary detection algorithms for region-specific applications and protocols",
        "Continuously improve classification accuracy using your own labeled data",
        "Create intelligence outputs tailored to your regulatory, commercial, and security requirements",
    ]:
        add_bullet(doc, b)

    add_normal(doc, "The result: a platform that doesn't just inspect packets -- it thinks, learns, predicts, and acts. And it does so entirely under your control.", bold=True)

    add_heading(doc, "Key Performance Metrics", 2)
    add_table(doc,
        ["Capability", "Specification"],
        [
            ["Passive Throughput", "Up to 1.2 Tbps per server"],
            ["Inline Enforcement", "Up to 500 Gbps per server"],
            ["Application Classification", "3,500+ apps, 98-99% accuracy"],
            ["Encrypted Traffic Intelligence", "95%+ accuracy via metadata analysis"],
            ["Concurrent Subscribers", "Tens of millions"],
            ["Flow Records", "5,000,000+ per second"],
            ["AI Query Response", "< 100ms natural language processing"],
            ["Threat Detection Latency", "< 10ms"],
            ["Report Generation", "< 60 seconds (vs. weeks with legacy DPI)"],
        ],
        col_widths_cm=[7, 9]
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 2: THE MARKET IMPERATIVE
    # ============================================================
    add_heading(doc, "THE MARKET IMPERATIVE", 1)
    add_heading(doc, "Four Unstoppable Forces", 2)

    add_normal(doc, "The network intelligence landscape is being reshaped by four forces that legacy DPI systems were never designed to handle:")

    forces = [
        ("1. Encryption Everywhere -- ", "Over 90% of global internet traffic is now encrypted. Legacy DPI tools built for cleartext inspection are blind. Networks have become opaque pipes where operators, regulators, and security agencies have lost meaningful visibility. ODUN.ONE restores this visibility through AI-powered encrypted traffic analysis that classifies content by behavioral patterns, metadata signatures, and statistical modeling -- without breaking encryption."),
        ("2. AI-Generated Traffic -- ", "By 2030, the majority of global network traffic will be generated or shaped by artificial intelligence. AI agents, automated systems, and machine-to-machine communication create traffic patterns that are unpredictable, polymorphic, and invisible to signature-based inspection. ODUN.ONE's AI engine is purpose-built to classify AI-generated traffic -- recognizing AI communication patterns as a distinct traffic category."),
        ("3. Relentless Scale -- ", "Subscriber bases, device counts, and bandwidth demands are exploding. 5G and future 6G rollouts compound the challenge. ODUN.ONE scales horizontally on commodity hardware -- adding capacity by deploying additional nodes, not by purchasing proprietary appliances. Each server processes up to 1.2 Tbps, scaling linearly with infrastructure."),
        ("4. Continuous Application Introductions -- ", "New applications appear daily, the vast majority encrypted end-to-end. Without a platform that adapts in real time, operators are perpetually blind to both revenue opportunities and security threats. ODUN.ONE's classification engine updates autonomously, identifying new applications through behavioral fingerprinting before manual signature creation is even required."),
    ]
    for bold_part, normal_part in forces:
        add_rich_para(doc, [(bold_part, True, False), (normal_part, False, False)])

    add_heading(doc, "The Incumbent Vacuum", 2)
    add_normal(doc, "The dominant legacy DPI vendor's bankruptcy removed the incumbent from [N]+ countries, creating a [$X]B addressable market opportunity. The remaining alternatives -- geopolitically-aligned infrastructure vendors -- carry baggage that sovereignty-minded governments and operators cannot accept.")
    add_normal(doc, "ODUN.ONE is the sovereign, non-aligned, AI-native answer.", bold=True)

    add_page_break(doc)

    # ============================================================
    # SECTION 3: PLATFORM OVERVIEW
    # ============================================================
    add_heading(doc, "PLATFORM OVERVIEW: THE ODUN METHODOLOGY", 1)
    add_normal(doc, "ODUN.ONE operates through a continuous four-step intelligence cycle that transforms raw network data into decisive action:")

    odun_steps = [
        ("O -- Observe",
         "Continuously monitor all network traffic in real time across fixed and mobile environments. Carrier-scale passive monitoring captures every packet, every flow, every session -- without impacting network performance.",
         ["L2-L7 deep packet inspection across all network interfaces",
          "Simultaneous monitoring of millions of concurrent subscribers",
          "Zero-copy packet processing for maximum throughput",
          "Support for 4G (Gi/SGi interface), 5G (N6 interface), and fixed broadband"]),
        ("D -- Decode",
         "Reliably classify encrypted and AI-driven traffic by application, protocol, and behavior. ODUN.ONE's multi-layered classification engine combines pattern matching, TLS/SNI metadata extraction, behavioral and statistical analysis, and deep learning models to identify 3,500+ applications with 98-99% accuracy.",
         ["3,500+ application and sub-application signatures across 14 categories",
          "TLS fingerprinting (JA3/JA4), SNI extraction, certificate analysis",
          "Behavioral classification for encrypted traffic (no decryption required)",
          "Custom signature creation for region-specific and proprietary applications",
          "Autonomous new-application detection through behavioral fingerprinting"]),
        ("U -- Understand",
         "AI-driven analytics correlate traffic patterns, subscriber context, network conditions, and anomalies -- transforming data into actionable intelligence. The platform doesn't just classify traffic, it understands what that traffic means in context.",
         ["Behavioral traffic modeling and pattern recognition",
          "Subscriber intent analysis and usage profiling",
          "Anomaly detection across multiple dimensions (volume, timing, behavior, geography)",
          "Predictive analytics for proactive threat detection and capacity planning",
          "AI-powered correlation engine linking events across time, location, and identity"]),
        ("N -- Navigate",
         "Translate intelligence into action. Enforce policies, optimize performance, enable monetization, and respond to threats -- all in real time with sub-10ms latency.",
         ["Real-time policy enforcement (blocking, shaping, prioritization, redirection)",
          "Automated threat response with configurable escalation workflows",
          "Dynamic subscriber management and service differentiation",
          "Revenue optimization through intelligent traffic handling",
          "Natural language policy creation via AI copilot"]),
    ]
    for title, desc, bullets in odun_steps:
        add_heading(doc, title, 2)
        add_normal(doc, desc)
        for b in bullets:
            add_bullet(doc, b)

    add_page_break(doc)

    # ============================================================
    # SECTION 4: CORE PLATFORM MODULES
    # ============================================================
    add_heading(doc, "CORE PLATFORM MODULES", 1)
    add_normal(doc, "ODUN.ONE is built from four modular components, each deployable independently or as an integrated suite:")

    # DataONE
    add_heading(doc, "DataONE -- High-Performance DPI Engine", 2)
    add_normal(doc, "The foundation. DataONE is the packet processing and classification powerhouse that captures, decodes, and structures network traffic at carrier scale.")
    for b in [
        "L2-L7 deep packet inspection with wire-speed processing",
        "Application and sub-application classification (3,500+ apps, 14 categories)",
        "TLS/SNI extraction and encrypted traffic metadata analysis",
        "Custom signature creation and management",
        "Traffic shaping, blocking, and redirection",
        "Flow tracking, session analysis, and record generation (5M+ flow records/second)",
        "GTP tunnel processing (GTPv1/v2) for mobile core visibility",
        "TelcoCloud support (Kubernetes-native and bare-metal deployments)",
        "HTTP header enrichment (MSISDN/IMSI insertion for subscriber-aware services)",
        "Tethering detection, SNI spoofing detection, VPN/proxy detection",
    ]:
        add_bullet(doc, b)

    # ControlONE
    add_heading(doc, "ControlONE -- Intelligent Policy & Subscriber Management", 2)
    add_normal(doc, "The enforcement engine. ControlONE manages subscriber sessions, applies policies, and integrates with carrier billing and policy infrastructure.")
    for b in [
        "RADIUS/DIAMETER integration for subscriber identity and session management",
        "4G PCRF/OCS interaction (Gx/Gy/Gz Diameter interfaces)",
        "5G PCF/CHF integration (N7/N40 HTTP/2 service-based interfaces)",
        "Whitelist and blacklist management with real-time updates",
        "Real-time charging integration (prepaid quota management, postpaid usage recording)",
        "Application-based charging (per-app quotas, zero-rating, tiered data plans)",
        "Subscriber-aware policy application (per-user rules based on plan, location, behavior)",
        "Dynamic session control with sub-second policy activation",
    ]:
        add_bullet(doc, b)

    # OpsONE
    add_heading(doc, "OpsONE -- Unified Operational Dashboard", 2)
    add_normal(doc, "The visibility layer. OpsONE provides real-time and historical operational intelligence through intuitive dashboards, reports, and drill-down analytics.")
    for b in [
        "Real-time traffic dashboards with sub-minute granularity",
        "Quality of Experience (QoE) metrics and AppScore index per application",
        "Multi-dimensional reporting: by application, traffic type, location, time, device, subscriber segment",
        "Drill-down capability: node -> category -> application -> individual subscriber",
        "Automated report generation (under 60 seconds vs. $100K-$300K and weeks with legacy vendors)",
        "Customizable executive dashboards for C-level visibility",
        "Export capabilities (PDF, CSV, API) for integration with existing BI tools",
    ]:
        add_bullet(doc, b)

    # AnalyticsONE
    add_heading(doc, "AnalyticsONE -- AI-Driven Analytics & Automation", 2)
    add_normal(doc, "The intelligence brain. AnalyticsONE is where ODUN.ONE's sovereign AI capabilities come to life -- transforming raw data into predictive insights, automated actions, and operator-guided intelligence.")

    add_normal(doc, "AI Paperclip Assistant -- Natural Language Operator Copilot:", bold=True)
    for b in [
        "Conversational interface for complex queries",
        "Automatic translation of natural language into network policies",
        "Multi-language support (English, Russian, Arabic, Chinese, Spanish, French, Turkish, and more)",
        "< 100ms query response time",
    ]:
        add_bullet(doc, b)

    for b in [
        "Behavioral Traffic Modeling -- Continuous learning from network behavior patterns",
        "Anomaly Detection -- Multi-dimensional anomaly identification",
        "Predictive Analytics -- Proactive threat detection and capacity forecasting",
        "Automatic Policy Recommendations -- AI suggests optimal policies based on observed patterns",
        "Adaptive Classification -- ML models that improve over time with network-specific data",
    ]:
        add_bullet(doc, b)

    # Bundling table
    add_heading(doc, "Module Bundling", 2)
    add_table(doc,
        ["Package", "Modules Included", "Ideal For"],
        [
            ["Essential", "DataONE", "Traffic visibility and classification"],
            ["Professional", "DataONE + ControlONE + OpsONE", "Full operational intelligence and enforcement"],
            ["Enterprise", "DataONE + ControlONE + OpsONE + AnalyticsONE", "Complete AI-powered sovereign intelligence"],
        ],
        col_widths_cm=[4, 6, 6]
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 5: AI & ML ARCHITECTURE
    # ============================================================
    add_heading(doc, "AI & MACHINE LEARNING ARCHITECTURE", 1)

    add_heading(doc, "Why AI-Native Matters", 2)
    add_normal(doc, "ODUN.ONE was designed from day one with artificial intelligence at its core -- not as an afterthought bolted onto a legacy inspection engine. Every architectural decision, from the data pipeline to the storage layer to the API surface, was made to support continuous machine learning at carrier scale.")
    add_normal(doc, "Legacy DPI vendors built their platforms in an era of cleartext traffic and static signatures. Their \"AI\" capabilities are typically bolt-on analytics modules that operate on pre-processed data, disconnected from the real-time packet processing engine. ODUN.ONE's AI is woven into the fabric of the platform -- it processes, learns, and acts in the same data pipeline as the DPI engine itself.")

    add_heading(doc, "Multi-Layered AI Classification Engine", 2)

    layers = [
        ("Layer 1: Deterministic Pattern Matching -- ", "High-speed signature matching for known applications and protocols. Handles the bulk of well-known traffic with near-zero latency using 3,500+ application signatures."),
        ("Layer 2: TLS/Encrypted Metadata Analysis -- ", "For encrypted traffic: TLS fingerprinting (JA3/JA4 hash analysis), SNI extraction, certificate chain analysis, handshake timing patterns, session resumption behavior."),
        ("Layer 3: Behavioral & Statistical Analysis -- ", "Packet size distribution analysis, inter-arrival timing patterns, flow duration and burst characteristics, connection establishment patterns, data transfer volume and directionality ratios."),
        ("Layer 4: Deep Learning Classification -- ", "Neural network models for fully obfuscated or novel applications. Temporal pattern recognition, transfer learning from known behaviors, ensemble methods, continuous model refinement using operator-labeled data."),
        ("Layer 5: Contextual Intelligence -- ", "Correlates classification with subscriber profile, geographic patterns, network topology, cross-session relationships, and threat intelligence feeds."),
    ]
    for bold_part, normal_part in layers:
        add_rich_para(doc, [(bold_part, True, False), (normal_part, False, False)])

    add_heading(doc, "AI Copilot: Natural Language Interface", 2)
    add_normal(doc, "The AnalyticsONE AI Paperclip Assistant represents a paradigm shift in how operators interact with network intelligence:")
    add_normal(doc, "Natural Language Policy Creation Examples:", bold=True)
    for e in [
        "\"Block all VoIP traffic within Telegram for prepaid subscribers\"",
        "\"Show me all suspicious traffic originating from Region X in the last 72 hours\"",
        "\"Generate a bandwidth consumption report by application for the top 1,000 subscribers\"",
        "\"Alert me when traffic from any single IP exceeds 500 Mbps sustained for more than 5 minutes\"",
        "\"Zero-rate all educational content for subscribers on Student Plans between 08:00-18:00\"",
    ]:
        add_bullet(doc, e)

    add_normal(doc, "The system translates conversational commands into precise network policies, executes them, and provides explainable confirmation of what actions were taken and why. Every AI-generated policy is auditable and reversible.")

    add_heading(doc, "Automated Report Generation", 2)
    add_table(doc,
        ["Metric", "ODUN.ONE", "Legacy DPI Vendors"],
        [
            ["Report Generation Time", "< 60 seconds", "3-4 weeks"],
            ["Cost Per Report", "Zero (included)", "$100,000 - $300,000"],
            ["Personnel Required", "Zero (automated)", "3-4 dedicated analysts"],
            ["Customization", "Unlimited (natural language)", "Fixed templates"],
            ["Real-Time Data", "Yes", "Batch processing only"],
        ],
        col_widths_cm=[5, 5.5, 5.5]
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 6: AI TRAINING PIPELINE
    # ============================================================
    add_heading(doc, "AI TRAINING PIPELINE: SOVEREIGN INTELLIGENCE WITHOUT LIMITS", 1)

    add_heading(doc, "The Sovereignty Advantage in AI Training", 2)
    add_normal(doc, "This is the single most important differentiator of ODUN.ONE's AI architecture.", bold=True)
    add_normal(doc, "Because ODUN.ONE is deployed 100% on-premises within the customer's sovereign infrastructure -- fully air-gapped if required, with zero external dependencies -- the AI training pipeline operates without the limitations imposed by cloud-based or hybrid architectures.")

    sovereignty_points = [
        ("Your data never leaves your control. ", "Every packet, every flow record, every behavioral model, every trained AI weight exists within your sovereign infrastructure. No cloud provider processes your data. No foreign jurisdiction has legal access to your intelligence."),
        ("You can train the AI on anything. ", "There are no cloud provider acceptable-use policies limiting what the AI can learn. No third-party terms of service restricting model training on sensitive data. No vendor veto over what intelligence you extract from your own network."),
        ("Your models are yours. ", "Every AI model trained on your data belongs to you. No vendor retains training data or model weights. No risk of your intelligence improving a competitor's deployment."),
        ("No capability ceiling. ", "Cloud-dependent platforms are constrained by what the vendor permits, what the cloud provider allows, and what cross-border data regulations prohibit. ODUN.ONE's sovereign deployment removes all three constraints simultaneously."),
    ]
    for bold_part, normal_part in sovereignty_points:
        add_rich_para(doc, [(bold_part, True, False), (normal_part, False, False)])

    add_heading(doc, "Sovereign AI Training Capabilities", 2)

    training_caps = [
        ("1. Custom Application Recognition Training",
         "Train the AI to recognize applications and protocols unique to your country, region, or network -- applications that no global vendor has ever seen.",
         "Operators label sample traffic from unknown applications. The ML pipeline extracts behavioral features, deep learning models are trained, and new signatures deploy to production in real time.",
         ["National banking applications with proprietary protocols",
          "Government communication platforms",
          "Regional messaging apps not covered by global signature libraries",
          "Military or defense-specific communication systems",
          "Custom enterprise applications operating on the network"]),
        ("2. Behavioral Baseline Training",
         "Teach the AI what \"normal\" looks like for your specific network, subscriber population, and traffic mix -- so that deviations are detected with extreme precision.",
         "The AI continuously observes network behavior, establishing statistical baselines. Anomaly detection models are calibrated to your network's specific characteristics with tunable sensitivity thresholds.",
         ["Every network is unique -- a baseline trained on one region's traffic is meaningless for an operator in another region",
          "Sovereign deployment means the AI learns YOUR network, not a generic global model"]),
        ("3. Threat Pattern Training",
         "Train the AI to recognize threat patterns specific to your threat landscape.",
         "Known threat samples are labeled and fed into the training pipeline. Models are continuously updated as new patterns emerge.",
         ["VPN and proxy tunneling detection (including obfuscated tunnels)",
          "DNS tunneling and covert channel identification",
          "Botnet command-and-control traffic patterns",
          "DDoS attack signature recognition and prediction",
          "Data exfiltration pattern recognition",
          "Application impersonation detection"]),
        ("4. Subscriber Behavior Modeling",
         "Build detailed behavioral models for subscriber segments, enabling predictive analytics for churn, fraud, abuse, and commercial opportunity.",
         "The AI analyzes individual and aggregate subscriber behavior over time. Clustering algorithms identify natural segments. Predictive models forecast future behavior.",
         ["Churn prediction: identify disengagement signals before subscribers leave",
          "Fraud detection: recognize SIM-swap, subscription fraud, and abuse patterns",
          "Usage forecasting: predict capacity requirements by segment and location",
          "Commercial targeting: identify upsell-responsive subscribers"]),
        ("5. Predictive Network Intelligence",
         "Train the AI to predict network events before they occur -- congestion, outages, security incidents, performance degradation.",
         "Historical event data is correlated with pre-event traffic patterns. The AI learns precursors and generates alerts with configurable lead times.",
         ["Network congestion forecasting by location, time, and application",
          "Application performance degradation prediction",
          "Security incident early warning",
          "Capacity exhaustion prediction"]),
        ("6. Custom Intelligence Model Training",
         "Train entirely custom AI models for intelligence use cases specific to your operational mandate.",
         "Operators define the intelligence objective. The platform provides data pipeline, compute, training framework, and deployment. The intelligence objectives are entirely determined by the customer.",
         ["No predefined limit to what the AI can be trained to do",
          "Supports commercial, regulatory, and security-oriented intelligence objectives"]),
    ]
    for title, desc, how, bullets in training_caps:
        add_heading(doc, title, 2)
        add_normal(doc, desc, bold=True)
        add_normal(doc, how)
        for b in bullets:
            add_bullet(doc, b)

    add_heading(doc, "Hardware Acceleration for AI Training", 2)
    for b in [
        "2x NVIDIA L40 48GB GPUs (optional, per server) for accelerated deep learning training",
        "Hardware acceleration for cryptographic operations",
        "Dedicated AI inference engines for real-time classification",
        "Scalable training infrastructure that grows with your data volume",
    ]:
        add_bullet(doc, b)

    add_page_break(doc)

    # ============================================================
    # SECTION 7: INTELLIGENCE MODULES
    # ============================================================
    add_heading(doc, "INTELLIGENCE MODULES", 1)
    add_normal(doc, "ODUN.ONE organizes its intelligence outputs into four use-case-driven modules:")

    intel_modules = [
        ("Regulatory Intelligence", "Compliance, governance, and regulatory enforcement.", [
            "Proactive blocking of illegal, prohibited, or regulated content and applications",
            "Subscriber profiling with detailed visibility into individual subscriber activity",
            "Comprehensive network/user/application visibility dashboards",
            "Complete audit trail for governance and oversight",
            "Pre-built regulatory compliance templates adaptable to national requirements"]),
        ("Application Intelligence", "Application performance monitoring and troubleshooting.", [
            "Quickly detect application-level problems invisible to basic network KPIs",
            "Real-time alerting when application performance degrades",
            "AI-driven anomaly trend analysis for emerging behavior changes",
            "AppScore (QoE Index): per-application quality scoring",
            "Continuous discovery and cataloging of all applications on the network"]),
        ("Operational Intelligence", "Network operations, subscriber care, and service assurance.", [
            "AI-powered root cause analysis identifying issue sources in seconds",
            "Real-time monitoring during outages to prioritize restoration",
            "Churn behavior analysis predicting subscriber departure before it happens",
            "Proactive issue detection reducing customer care call volume",
            "Automated detection, classification, and escalation of operational events"]),
        ("Business Intelligence", "Revenue optimization and executive visibility.", [
            "Purpose-built executive dashboards for C-level decision-making",
            "Data-driven identification of new service and monetization opportunities",
            "Subscriber segment analysis for personalized promotion benchmarking",
            "Subscriber lifetime value modeling with predictive analytics"]),
    ]
    for title, focus, bullets in intel_modules:
        add_heading(doc, title, 2)
        add_normal(doc, focus, italic=True)
        for b in bullets:
            add_bullet(doc, b)

    add_page_break(doc)

    # ============================================================
    # SECTION 8: TELECOM USE CASES
    # ============================================================
    add_heading(doc, "TELECOM USE CASES: AI-DRIVEN NETWORK INTELLIGENCE", 1)

    add_heading(doc, "Data Monetization & Revenue Optimization", 2)
    for bold_part, normal_part in [
        ("Application-Based Charging -- ", "Train the AI to classify traffic at the application level with 98%+ accuracy, enabling operators to charge differently for different applications. Zero-rate educational content, premium-charge high-bandwidth gaming, bundle social media into specific plans."),
        ("Usage-Based Fair Use Policy -- ", "AI models learn subscriber usage patterns and automatically enforce fair-use policies -- throttling heavy users during peak congestion while maintaining QoE for the majority. Dynamic thresholds adapt to real-time conditions."),
        ("Location-Based Services -- ", "Combine DPI intelligence with geographic data to offer location-specific services, pricing, and content. The AI learns geographic traffic patterns and optimizes service delivery by region."),
        ("Tethering Detection & Control -- ", "AI behavioral models detect tethering even when subscribers attempt to mask it through VPN or proxy techniques. Operators can enforce tethering policies or monetize tethering as a premium feature."),
    ]:
        add_rich_para(doc, [(bold_part, True, False), (normal_part, False, False)])

    add_heading(doc, "Quality of Experience Management", 2)
    for bold_part, normal_part in [
        ("Proactive QoE Monitoring -- ", "The AI continuously scores application quality across every subscriber, every application, and every location. When QoE drops below thresholds, operators are alerted before subscribers complain."),
        ("Video Quality Optimization -- ", "With 74% of mobile traffic being video, AI models optimize video delivery by detecting codec, resolution, and buffering patterns -- enabling dynamic quality management based on network conditions and subscriber plans."),
        ("5G Network Slicing Intelligence -- ", "ODUN.ONE's AI provides traffic intelligence for effective network slicing -- understanding which applications and subscribers should be assigned to which slices, continuously optimizing allocation based on real-time demand."),
    ]:
        add_rich_para(doc, [(bold_part, True, False), (normal_part, False, False)])

    add_heading(doc, "Network Operations Intelligence", 2)
    for bold_part, normal_part in [
        ("Automated Root Cause Analysis -- ", "When network issues occur, the AI correlates events across multiple dimensions (time, geography, application, subscriber segment, network element) to identify root causes in seconds -- not hours."),
        ("Predictive Capacity Planning -- ", "AI models trained on historical traffic patterns predict future capacity requirements by location, time period, and application category -- enabling precise infrastructure investment."),
        ("Churn Prevention -- ", "Behavioral models identify subscribers exhibiting early-stage disengagement signals and trigger retention actions before the subscriber leaves."),
    ]:
        add_rich_para(doc, [(bold_part, True, False), (normal_part, False, False)])

    add_heading(doc, "Security & Threat Intelligence", 2)
    for bold_part, normal_part in [
        ("Encrypted Threat Detection -- ", "AI models detect malicious traffic within encrypted flows through behavioral analysis -- identifying botnet communication, data exfiltration, and C2 channels without breaking encryption."),
        ("DDoS Detection & Mitigation -- ", "ML models learn normal traffic baselines and detect volumetric, protocol, and application-layer DDoS attacks within milliseconds. Automated mitigation activates instantly."),
        ("Bypass & Abuse Detection -- ", "AI continuously evolves to detect new evasion techniques -- VPN tunneling, DNS tunneling, SNI spoofing, protocol impersonation -- based on actual bypass attempts observed on your network."),
    ]:
        add_rich_para(doc, [(bold_part, True, False), (normal_part, False, False)])

    add_page_break(doc)

    # ============================================================
    # SECTION 9: LAW ENFORCEMENT & NATIONAL SECURITY
    # ============================================================
    add_heading(doc, "LAW ENFORCEMENT & NATIONAL SECURITY USE CASES", 1)
    add_normal(doc, "ODUN.ONE provides comprehensive capabilities for authorized law enforcement and national security operations. Every capability is designed with built-in governance, audit trails, and compliance frameworks.")

    # 9.1
    add_heading(doc, "Lawful Traffic Monitoring & Interception", 2)

    add_normal(doc, "Targeted Subscriber Monitoring:", bold=True)
    for b in [
        "Real-time traffic capture for targeted subscribers identified by MSISDN, IMSI, IP address, or other identifiers",
        "Full application-level classification of all subscriber activity",
        "Session reconstruction showing complete communication timelines",
        "Content category analysis revealing behavioral patterns",
        "Historical traffic retrieval for specified time periods",
        "Metadata analysis without content decryption for privacy-compliant monitoring",
    ]:
        add_bullet(doc, b)

    add_normal(doc, "Bulk Traffic Analysis for Pattern Detection:", bold=True)
    for b in [
        "Communication pattern analysis across large subscriber populations",
        "Anomalous behavior detection flagging unusual activity for investigator review",
        "Time-series analysis revealing changes in communication patterns over time",
        "Geographic correlation identifying activity clusters by location",
        "Cross-network relationship mapping showing communication links between subjects",
    ]:
        add_bullet(doc, b)

    add_normal(doc, "ETSI-Compliant Lawful Interception:", bold=True)
    for b in [
        "Standardized handover interfaces (HI1, HI2, HI3) for law enforcement delivery",
        "Mediation function for warrant management and authorization control",
        "Comprehensive audit logging of all interception activities",
        "Multi-level authorization workflows with separation of duties",
        "Tamper-evident logging for evidentiary chain of custody",
    ]:
        add_bullet(doc, b)

    # 9.2
    add_heading(doc, "Criminal Investigation Support", 2)

    add_normal(doc, "Communication Pattern Analysis:", bold=True)
    for b in [
        "Identify communication networks between subjects of interest",
        "Detect changes in communication behavior (frequency, timing, counterparties)",
        "Map social networks based on traffic metadata",
        "Identify previously unknown associates through traffic correlation",
        "Timeline reconstruction of communication activity around events of interest",
    ]:
        add_bullet(doc, b)

    add_normal(doc, "Application-Specific Monitoring:", bold=True)
    for b in [
        "Messaging application identification and session tracking (WhatsApp, Telegram, Signal, and 3,500+ others)",
        "VoIP call detection and metadata capture",
        "File transfer and cloud storage activity monitoring",
        "Social media usage profiling",
        "Dark web and anonymization tool detection (Tor, VPN, proxy services)",
    ]:
        add_bullet(doc, b)

    add_normal(doc, "Digital Evidence Collection:", bold=True)
    for b in [
        "Chain-of-custody compliant data capture and storage",
        "Timestamped, cryptographically signed evidence records",
        "Export formats compatible with major forensic analysis tools",
        "Configurable data retention policies aligned with legal requirements",
        "Role-based access control ensuring evidence integrity",
    ]:
        add_bullet(doc, b)

    # 9.3
    add_heading(doc, "National Security & Counter-Terrorism", 2)

    add_normal(doc, "Threat Actor Identification:", bold=True)
    for b in [
        "Communication patterns deviating from civilian norms",
        "Encrypted channel detection and behavioral classification",
        "Covert communication channel identification (steganography, tunneling, protocol abuse)",
        "Cross-border communication pattern analysis",
        "Periodic communication pattern detection (timed check-ins, dead drops)",
    ]:
        add_bullet(doc, b)

    add_normal(doc, "Critical Infrastructure Protection:", bold=True)
    for b in [
        "SCADA/ICS traffic anomaly detection",
        "Critical infrastructure communication pattern baselining",
        "Unauthorized access attempt detection",
        "Supply chain communication monitoring",
        "Insider threat indicator identification",
    ]:
        add_bullet(doc, b)

    add_normal(doc, "Cyber Threat Intelligence:", bold=True)
    for b in [
        "Advanced Persistent Threat (APT) traffic pattern detection",
        "Command-and-control (C2) infrastructure identification",
        "Malware communication fingerprinting",
        "Threat intelligence feed generation for SOC/SIEM integration",
        "Zero-day exploit traffic anomaly detection",
    ]:
        add_bullet(doc, b)

    # 9.4
    add_heading(doc, "Ministry of Interior (MOI) Operations", 2)

    add_normal(doc, "Traffic Control & Subscriber Monitoring:", bold=True)
    for b in [
        "Full traffic visibility across all subscriber segments",
        "Subscriber activity profiling under legal authorization",
        "Content policy enforcement (blocking prohibited content and services)",
        "Emergency traffic prioritization during crisis events",
        "Public event monitoring and crowd analytics via traffic patterns",
    ]:
        add_bullet(doc, b)

    add_normal(doc, "Governance & Accountability:", bold=True)
    for b in [
        "Multi-tier role-based access control (RBAC) with granular permissions",
        "All actions logged with operator identity, timestamp, and authorization reference",
        "Warrant/authorization management system with expiration enforcement",
        "Automatic de-activation of monitoring targets when authorization expires",
        "Complete audit trail exportable for judicial review",
        "Separation of duties between authorization, provisioning, and data access",
    ]:
        add_bullet(doc, b)

    # 9.5
    add_heading(doc, "Ministry of Defense (MOD) Operations", 2)

    add_normal(doc, "Signal Intelligence Support:", bold=True)
    for b in [
        "Selective traffic diversion to secure, segregated analysis environments",
        "Advanced signaling behavior analysis (protocol-level inspection)",
        "Communication pattern analysis across monitored zones",
        "Encrypted traffic behavioral classification without decryption",
        "Multi-protocol correlation (voice, data, signaling combined analysis)",
    ]:
        add_bullet(doc, b)

    add_normal(doc, "Operational Security:", bold=True)
    for b in [
        "Segregated processing environments with physical and logical isolation",
        "Air-gapped deployment capability with zero external connectivity",
        "Hardware Security Module (HSM) support for key management",
        "Classification-level data handling with compartmentalized access",
        "Secure destruction/purge capabilities for classified data",
        "TEMPEST-compatible deployment options",
    ]:
        add_bullet(doc, b)

    add_page_break(doc)

    # ============================================================
    # SECTION 10: AI TRAINING SCENARIOS
    # ============================================================
    add_heading(doc, "AI ANALYTICS: TELCO & LAW ENFORCEMENT TRAINING SCENARIOS", 1)
    add_normal(doc, "This section describes how ODUN.ONE's sovereign AI platform can be trained for specific use cases. Because the platform operates entirely within sovereign infrastructure, there are virtually no restrictions on what the AI can be taught to detect, classify, predict, or act upon.")

    add_heading(doc, "Telco AI Training Scenarios", 2)

    telco_scenarios = [
        ("Scenario 1: Intelligent Traffic Optimization",
         "Train the AI to automatically optimize traffic handling based on real-time network conditions, subscriber profiles, and business rules.",
         "Feed historical traffic data, congestion events, and subscriber satisfaction metrics into the training pipeline. AI learns the relationship between traffic management decisions and subscriber outcomes.",
         ["20-30% reduction in congestion-related QoE degradation",
          "Automated traffic management decisions at carrier scale",
          "Optimized bandwidth allocation across subscriber segments and applications"]),
        ("Scenario 2: Revenue Leakage Detection",
         "Train the AI to detect revenue leakage from unauthorized usage, policy bypass, and billing discrepancies.",
         "Correlate DPI traffic data with billing records. AI learns patterns associated with revenue leakage (unauthorized tethering, OTT voice bypass, data plan exploitation).",
         ["Recovery of 3-8% of lost revenue through leakage detection",
          "Real-time identification of new bypass and abuse techniques",
          "Automated enforcement of usage policies"]),
        ("Scenario 3: 5G Service Assurance",
         "Train the AI to ensure SLA compliance across 5G network slices by predicting performance degradation before it impacts subscribers.",
         "Collect per-slice performance metrics, traffic composition, and subscriber behavior data. AI learns performance characteristics and failure modes of each slice.",
         ["Proactive SLA management with predictive alerting",
          "Automated slice optimization based on real-time demand",
          "Reduced SLA penalties and improved subscriber satisfaction"]),
        ("Scenario 4: AI-Generated Traffic Classification",
         "Train the AI to identify and classify traffic generated by AI agents, automated systems, and machine-to-machine communication.",
         "Label examples of AI-generated traffic. ML models learn behavioral characteristics distinguishing AI traffic from human traffic.",
         ["Visibility into the fastest-growing traffic category",
          "Ability to monetize AI-generated traffic differently from human traffic",
          "Security monitoring of AI agent behavior on the network"]),
    ]
    for title, objective, approach, outcomes in telco_scenarios:
        add_heading(doc, title, 2)
        add_rich_para(doc, [("Objective: ", True, False), (objective, False, False)])
        add_rich_para(doc, [("Training Approach: ", True, False), (approach, False, False)])
        add_normal(doc, "Outcomes:", bold=True)
        for o in outcomes:
            add_bullet(doc, o)

    add_heading(doc, "Law Enforcement AI Training Scenarios", 2)

    le_scenarios = [
        ("Scenario 5: Behavioral Anomaly Detection for Person of Interest Monitoring",
         "Train the AI to detect significant behavioral changes in monitored subjects' communication patterns that may indicate operational activity.",
         "Establish communication baselines for monitored subjects. AI learns what constitutes significant deviation. Models trained on historical case data correlate behavioral changes with operational activity.",
         ["Reduced investigator workload through automated pattern monitoring",
          "Earlier detection of operationally significant behavioral changes",
          "Reduced false positives through AI-calibrated alerting"]),
        ("Scenario 6: Network Mapping for Criminal Organizations",
         "Train the AI to identify and map communication networks used by criminal organizations, including attempts to fragment communication across multiple platforms.",
         "Feed known organization communication patterns into the training pipeline. AI learns network-level indicators of organized communication. Graph analysis algorithms map relationships.",
         ["Automated identification of communication network structures",
          "Discovery of previously unknown associates and organizational nodes",
          "Real-time updating of network maps as patterns evolve"]),
        ("Scenario 7: Covert Channel Detection",
         "Train the AI to detect covert communication channels hidden within normal network traffic -- DNS tunneling, steganographic communication, protocol abuse.",
         "Train models on known covert channel techniques. Behavioral analysis identifies statistical anomalies. Deep learning detects subtle deviations from normal protocol behavior.",
         ["Detection of covert channels that signature-based systems miss entirely",
          "Identification of new and custom covert communication methods",
          "Sub-second alerting on detected covert channel activation"]),
        ("Scenario 8: Radicalization & Extremism Pattern Detection",
         "Train the AI to identify traffic patterns associated with radicalization processes -- progressive engagement with extremist content, migration to encrypted platforms.",
         "Define behavioral indicators associated with radicalization trajectories. AI models learn to identify these indicators. Multi-stage detection tracks progression over time.",
         ["Early identification of at-risk individuals for intervention programs",
          "Pattern-based detection independent of content decryption",
          "Reduced false positives through multi-indicator behavioral modeling"]),
        ("Scenario 9: Cross-Border Communication Intelligence",
         "Train the AI to analyze cross-border communication patterns for national security intelligence.",
         "Map normal cross-border communication patterns by subscriber segment, application, and destination. AI learns baseline patterns and flags meaningful deviations.",
         ["Automated flagging of unusual cross-border communication patterns",
          "Early warning of coordinated cross-border activities",
          "Intelligence support for border security and counter-proliferation"]),
        ("Scenario 10: Real-Time Crisis Event Response",
         "Train the AI to detect and respond to crisis events through network traffic analysis.",
         "Train on historical crisis event data. AI learns traffic signatures that precede and accompany different types of crisis events. Automated response playbooks activate.",
         ["Sub-minute detection of crisis events through traffic pattern analysis",
          "Automated traffic prioritization for emergency communications",
          "Real-time situational awareness for command and control operations"]),
    ]
    for title, objective, approach, outcomes in le_scenarios:
        add_heading(doc, title, 2)
        add_rich_para(doc, [("Objective: ", True, False), (objective, False, False)])
        add_rich_para(doc, [("Training Approach: ", True, False), (approach, False, False)])
        add_normal(doc, "Outcomes:", bold=True)
        for o in outcomes:
            add_bullet(doc, o)

    add_page_break(doc)

    # ============================================================
    # SECTION 11: DEPLOYMENT ARCHITECTURE
    # ============================================================
    add_heading(doc, "DEPLOYMENT ARCHITECTURE & DATA SOVEREIGNTY", 1)

    add_heading(doc, "Deployment Options", 2)

    for title, desc in [
        ("100% On-Premises / Air-Gapped -- ",
         "The primary deployment model for government and defense customers. Complete air-gapped infrastructure with zero external connectivity. All processing, storage, AI training, and intelligence generation occurs within sovereign infrastructure. No cloud dependencies, no external API calls, no telemetry."),
        ("Private Cloud-Native -- ",
         "Kubernetes-native containerized microservices architecture within customer data centers. Docker-based deployment with automated orchestration. Horizontal scaling by adding nodes. Rolling updates without service interruption."),
        ("Hybrid Deployment -- ",
         "Critical functions on-premises, analytics in private cloud. Flexible data residency policies per data classification level. Seamless integration between environments."),
    ]:
        add_rich_para(doc, [(title, True, False), (desc, False, False)])

    add_heading(doc, "Why 100% Sovereign Matters", 2)
    add_table(doc,
        ["Concern", "Cloud-Dependent DPI", "ODUN.ONE Sovereign"],
        [
            ["Data residency", "Data may transit foreign jurisdictions", "All data within national borders"],
            ["Legal jurisdiction", "Subject to foreign government subpoena", "Subject only to national law"],
            ["AI training data", "Shared with vendor cloud", "Exclusively owned by the customer"],
            ["Operational continuity", "Dependent on cloud provider availability", "Fully self-contained"],
            ["Capability restrictions", "Limited by cloud acceptable-use policies", "No external limitations"],
            ["Supply chain risk", "Cloud provider is an attack surface", "Minimized to hardware only"],
            ["Classification handling", "Incompatible with classified operations", "Supports all classification levels"],
        ],
        col_widths_cm=[4, 5.5, 5.5]
    )

    add_page_break(doc)

    # ============================================================
    # SECTION 12: TECHNICAL SPECIFICATIONS
    # ============================================================
    add_heading(doc, "TECHNICAL SPECIFICATIONS", 1)

    add_heading(doc, "Hardware Reference Configuration", 2)
    add_normal(doc, "HP ProLiant DL385 G11 (2U Rack-Mounted)", italic=True)
    add_table(doc,
        ["Component", "Specification"],
        [
            ["Processors", "2x AMD EPYC 9845 (256 cores, 3.7 GHz max, 768 MB L3 cache)"],
            ["Memory", "768 GB DDR5 ECC RAM"],
            ["Network Interfaces", "6x 100GbE adapters (Intel E810 / Mellanox ConnectX-6) -- 12 ports, 600 Gbps raw"],
            ["Storage", "4 TB NVMe SSD, RAID-1"],
            ["AI Accelerators", "2x NVIDIA L40 48GB GPU (optional)"],
            ["Form Factor", "2U rack-mounted"],
        ],
        col_widths_cm=[5, 11]
    )
    add_normal(doc, "ODUN.ONE is hardware-agnostic. The above is a reference configuration.", italic=True)

    add_heading(doc, "Supported Protocols", 2)
    for b in [
        "Network Layer: IPv4, IPv6, GTP (v1/v2), GRE, VXLAN, MPLS (802.1Q)",
        "Transport Layer: TCP, UDP, SCTP, QUIC",
        "Application Layer: HTTP/1.1, HTTP/2, HTTP/3, TLS 1.2, TLS 1.3, DNS, DHCP, SIP, RTP, RTCP",
        "Proprietary Protocols: 3,000+ (YouTube, WhatsApp, Telegram, TikTok, Netflix, etc.)",
        "Data Export: JSON, Protocol Buffers, Parquet, Syslog, CEF",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "Performance Summary", 2)
    add_table(doc,
        ["Metric", "Specification"],
        [
            ["Passive throughput", "Up to 1.2 Tbps per server"],
            ["Inline enforcement", "Up to 500 Gbps per server"],
            ["Flow records", "5,000,000+ per second"],
            ["Application signatures", "3,500+ (continuously updated)"],
            ["Classification accuracy", "98-99% (including encrypted traffic)"],
            ["Encrypted traffic accuracy", "95%+ via metadata analysis"],
            ["AI query response", "< 100ms"],
            ["Threat detection latency", "< 10ms"],
            ["Report generation", "< 60 seconds"],
            ["Concurrent subscribers", "Tens of millions"],
        ],
        col_widths_cm=[6, 10]
    )

    add_heading(doc, "High Availability & Redundancy", 2)
    for b in [
        "Active-Active clustering with automatic failover",
        "Session state synchronization across cluster nodes",
        "Geographic redundancy for multi-site deployments",
        "Automated configuration backup and point-in-time recovery",
        "Disaster recovery with configurable RPO/RTO objectives",
    ]:
        add_bullet(doc, b)

    add_page_break(doc)

    # ============================================================
    # SECTION 13: INTEGRATION ARCHITECTURE
    # ============================================================
    add_heading(doc, "INTEGRATION ARCHITECTURE", 1)

    add_heading(doc, "4G EPC Integration", 2)
    add_table(doc,
        ["Function", "Interface", "Protocol"],
        [
            ["Data Plane (DPI inspection)", "Gi/SGi (PGW to Internet)", "IP traffic"],
            ["Subscriber Awareness", "RADIUS / GTP-C / Gx", "RADIUS / GTP / Diameter"],
            ["Policy Enforcement", "Gx/S7 (PCRF-PCEF)", "Diameter"],
            ["Online Charging", "Gy (OCS)", "Diameter"],
            ["Offline Charging", "Gz (OFCS)", "Diameter"],
        ],
        col_widths_cm=[5, 6, 5]
    )

    add_heading(doc, "5G SBA Integration", 2)
    add_table(doc,
        ["Function", "Interface", "Protocol"],
        [
            ["Data Plane (DPI inspection)", "N6 (UPF to Data Network)", "IP traffic"],
            ["Subscriber Awareness", "Nsmf (SMF Event Exposure) / Nnef", "HTTP/2"],
            ["Policy Enforcement", "N7 (PCF-SMF) + N4 (PFCP to UPF)", "HTTP/2 + PFCP"],
            ["Converged Charging", "N40/Nchf (CHF)", "HTTP/2"],
        ],
        col_widths_cm=[5, 6, 5]
    )

    add_normal(doc, "Key Architecture Insight: In 5G, all SBI signaling is encrypted (HTTP/2 over TLS), requiring the DPI system to be an active participant in the architecture -- embedded at or alongside the UPF.", bold=True)

    add_heading(doc, "External System Integration", 2)
    for b in [
        "SIEM/SOC: Real-time alert feed to Security Information and Event Management platforms",
        "Threat Intelligence: Bi-directional integration with threat intelligence feeds (STIX/TAXII)",
        "BSS/OSS: Integration with operator billing, CRM, and operational support systems",
        "Lawful Interception: ETSI-compliant handover interfaces (HI1/HI2/HI3)",
        "SDN Controllers: Policy synchronization with software-defined networking infrastructure",
        "REST APIs: Open API surface for custom integrations and third-party platform connectivity",
    ]:
        add_bullet(doc, b)

    add_page_break(doc)

    # ============================================================
    # SECTION 14: SECURITY, COMPLIANCE & GOVERNANCE
    # ============================================================
    add_heading(doc, "SECURITY, COMPLIANCE & GOVERNANCE", 1)

    add_heading(doc, "Security Features", 2)
    add_table(doc,
        ["Capability", "Detail"],
        [
            ["Access Control", "Role-based access control (RBAC) with granular, per-function permissions"],
            ["Authentication", "Multi-factor authentication (MFA), LDAP/AD integration"],
            ["Encryption", "TLS 1.3 for all management interfaces, encrypted data at rest"],
            ["Audit Logging", "Comprehensive, tamper-evident logging of all operator actions"],
            ["Secure Boot", "Firmware validation and integrity checking"],
            ["HSM Support", "Hardware Security Module integration for cryptographic key management"],
            ["Data Purge", "Secure data destruction capabilities for classified environments"],
        ],
        col_widths_cm=[4, 12]
    )

    add_heading(doc, "Compliance Standards", 2)
    add_table(doc,
        ["Standard", "Coverage"],
        [
            ["ISO/IEC 27001", "Information Security Management System"],
            ["ETSI LI Standards", "Lawful Interception compliance"],
            ["3GPP", "Mobile network integration standards"],
            ["GDPR", "Data protection and privacy requirements"],
            ["SOC 2 Type II", "Service organization controls"],
            ["Common Criteria EAL4+", "Security evaluation assurance"],
            ["FIPS 140-2 Level 3", "Cryptographic module validation (optional)"],
        ],
        col_widths_cm=[5, 11]
    )

    add_heading(doc, "Governance Framework", 2)
    for b in [
        "Separation of Duties: Authorization, provisioning, and data access are segregated roles",
        "Warrant Management: Built-in authorization lifecycle management with automatic expiration",
        "Audit Trail: Every action logged with operator identity, timestamp, and authorization reference",
        "Judicial Export: Audit trails and evidence packages exportable for judicial review",
        "Data Retention: Configurable retention periods aligned with national legal requirements",
        "Classification Handling: Support for multi-level classification with compartmentalized access",
    ]:
        add_bullet(doc, b)

    add_page_break(doc)

    # ============================================================
    # SECTION 15: WHY ODUN.ONE
    # ============================================================
    add_heading(doc, "WHY ODUN.ONE", 1)

    add_heading(doc, "The Only Platform That Combines", 2)

    for bold_part, normal_part in [
        ("Sovereign by Design -- ", "100% on-premises, in-country, air-gapped capable. No foreign cloud dependencies. No vendor backdoors. No external data processing. Your data, your models, your intelligence."),
        ("AI-Native Architecture -- ", "Built from day one with AI at the core -- not a legacy DPI engine with analytics bolted on. The AI sees every packet, learns from every flow, and improves with every cycle."),
        ("Unlimited AI Training Capability -- ", "Because it's sovereign, the AI can be trained on anything your operational mandate requires. No cloud provider restrictions. No vendor limitations. No capability ceiling."),
        ("Carrier-Scale Performance -- ", "Terabit throughput on commodity hardware. Tens of millions of concurrent subscribers. Five million flow records per second. Reports in 60 seconds that legacy vendors charge $100K+ for."),
        ("Clean-Slate Engineering -- ", "Zero legacy code. Zero technical debt. Cloud-native microservices designed for the encrypted, AI-driven, 5G world -- not a 28-year-old codebase modernized through patches."),
        ("Non-Aligned Origin -- ", "Not Russian. Not Chinese. Not Israeli. Not American. Built specifically for nations that require sovereign technology without geopolitical strings."),
        ("Proven in Production -- ", "Commercially launched February 2026. First production deployment live. Real traffic. Real subscribers. Real intelligence."),
    ]:
        add_rich_para(doc, [(bold_part, True, False), (normal_part, False, False)])

    add_heading(doc, "vs. Legacy DPI", 2)
    add_table(doc,
        ["Capability", "ODUN.ONE", "Legacy Solutions"],
        [
            ["Architecture", "Cloud-native microservices, clean-slate", "Monolithic legacy (25+ years old)"],
            ["AI Integration", "Native AI copilot, NL, continuous ML", "Limited or bolted-on analytics"],
            ["Throughput", "Terabit scale on commodity hardware", "Sub-terabit on proprietary appliances"],
            ["Encrypted Traffic", "95%+ accuracy via multi-layered AI", "Limited encrypted traffic insight"],
            ["Deployment Time", "Days (Kubernetes)", "Months (proprietary appliances)"],
            ["Report Generation", "< 60 seconds, automated, unlimited", "3-4 weeks, $100K-$300K per report"],
            ["Sovereignty", "100% air-gapped, sovereign AI training", "Cloud-dependent, vendor-controlled"],
            ["Scalability", "Horizontal, software-defined", "Vertical, hardware-centric"],
            ["Cost Structure", "Hardware-agnostic, no lock-in", "Proprietary, vendor lock-in"],
            ["Innovation Speed", "Continuous deployment", "Quarterly major releases"],
        ],
        col_widths_cm=[4, 5.5, 5.5]
    )

    add_page_break(doc)

    # ============================================================
    # LICENSING & SUPPORT
    # ============================================================
    add_heading(doc, "LICENSING & SUPPORT", 1)

    add_heading(doc, "Licensing Models", 2)
    for b in [
        "Capacity-Based: Licensed by throughput (100 Gbps increments) and/or subscriber count",
        "Module-Based: Essential, Professional, or Enterprise bundles",
        "Perpetual or Subscription: Flexible to match customer procurement models",
    ]:
        add_bullet(doc, b)

    add_heading(doc, "Support Packages", 2)
    add_table(doc,
        ["Tier", "Coverage", "Response SLA"],
        [
            ["Standard", "8x5, email & web portal", "Next business day"],
            ["Premium", "24x7, phone & email", "4 hours"],
            ["Enterprise", "24x7, dedicated Technical Account Manager", "1 hour"],
        ],
        col_widths_cm=[4, 7, 5]
    )

    add_heading(doc, "Professional Services", 2)
    for b in [
        "Network architecture review and deployment planning",
        "Custom integration development",
        "Performance tuning and optimization",
        "Operator training (basic, advanced, and security analyst programs)",
        "Administrator certification program",
        "Managed services (24x7 monitoring, proactive threat hunting, compliance reporting)",
    ]:
        add_bullet(doc, b)

    add_page_break(doc)

    # ============================================================
    # CLOSING PAGE
    # ============================================================
    for _ in range(6):
        doc.add_paragraph(style='Normal')

    # Company name
    closing = doc.add_paragraph(style='Normal')
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run("31 Concept (31C)")
    run.font.size = Pt(20)
    run.font.color.rgb = PURPLE
    run.bold = True

    # Locations
    loc = doc.add_paragraph(style='Normal')
    loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = loc.add_run("[HQ City 1]  |  [HQ City 2]")
    run.font.size = Pt(12)

    # Website
    web = doc.add_paragraph(style='Normal')
    web.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = web.add_run("www.31c.io  |  info@31c.io")
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE

    doc.add_paragraph(style='Normal')

    # CEO quote
    quote = doc.add_paragraph(style='Normal')
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = quote.add_run("\"We're not here to wait and see where the market goes.\nWe're here to lead it.\"")
    run.font.size = Pt(14)
    run.font.color.rgb = PURPLE
    run.italic = True

    doc.add_paragraph(style='Normal')
    doc.add_paragraph(style='Normal')

    # Disclaimer
    disc = doc.add_paragraph(style='Normal')
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disc.add_run("This document contains confidential and proprietary information of 31 Concept.\nDistribution is restricted to authorized recipients under NDA.\nPatent-pending technology.")
    run.font.size = Pt(8.5)
    run.italic = True

    # Save
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"DOCX generated successfully: {OUTPUT}")


# ============================================================
# CLI / Main
# ============================================================
if __name__ == "__main__":
    build_document()
